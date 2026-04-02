"""
ForensiScan - Advanced File Intelligence & Forensic Analysis Tool
A comprehensive tool for extracting metadata from various file types
Useful for digital forensics, OSINT, and security assessments

Author: Sayan Pal
Collaborator: Soumit Santra
"""

import os
import sys
import json
import argparse
import subprocess
import logging
import platform
import zipfile
import shutil
from datetime import datetime, timezone
from pathlib import Path
from typing import Dict, Any, List, Optional, Tuple
from concurrent.futures import ThreadPoolExecutor, as_completed
from threading import Lock
import hashlib
import math
import csv
import uuid
import re
import struct

# Third-party imports
try:
    from colorama import init, Fore, Style
    from tqdm import tqdm
    COLORAMA_AVAILABLE = True
except ImportError:
    COLORAMA_AVAILABLE = False
    class MockColor:
        def __getattr__(self, name): return ""
    Fore = MockColor()
    Style = MockColor()
    def init(): pass
    def tqdm(iterable=None, **kwargs): return iterable

if COLORAMA_AVAILABLE:
    init(autoreset=True)

try:
    from PIL import Image
    from PIL.ExifTags import TAGS, GPSTAGS
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False

try:
    from hachoir.parser import createParser
    from hachoir.metadata import extractMetadata
    from hachoir.core.cmd_line import unicodeFilename
    HACHOIR_AVAILABLE = True
except ImportError:
    HACHOIR_AVAILABLE = False

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import mutagen
    from mutagen.mp3 import MP3
    from mutagen.mp4 import MP4
    from mutagen.flac import FLAC
    AUDIO_AVAILABLE = True
except ImportError:
    AUDIO_AVAILABLE = False

try:
    import magic
    MAGIC_AVAILABLE = True
except ImportError:
    MAGIC_AVAILABLE = False

try:
    import pefile
    PE_AVAILABLE = True
except ImportError:
    PE_AVAILABLE = False

try:
    import yara
    YARA_AVAILABLE = True
except ImportError:
    YARA_AVAILABLE = False

try:
    import pytsk3
    TSK_AVAILABLE = True
except ImportError:
    TSK_AVAILABLE = False

try:
    import pyewf
    EWF_AVAILABLE = True
except ImportError:
    EWF_AVAILABLE = False

try:
    import aff4
    AFF4_AVAILABLE = True
except ImportError:
    AFF4_AVAILABLE = False


# ─────────────────────────────────────────────
#  Global constants
# ─────────────────────────────────────────────
VERSION          = "4.0.1"
TOOL_NAME        = "ForensiScan"
DEFAULT_WORKERS  = 4
EXIFTOOL_TIMEOUT = 30

# ─── Known-malicious hash databases (sample IOCs – extend as needed) ──────────
KNOWN_MALICIOUS_HASHES: Dict[str, str] = {
    "44d88612fea8a8f36de82e1278abb02f": "EICAR Test File",
    "e6d290a03b70cfa5d4451da444bdea39": "Mirai Botnet Sample",
    "098f6bcd4621d373cade4e832627b4f6": "Test Hash (demo)",
    "275a021bbfb6489e54d471899f7db9d1663fc695ec2fe2a2c4538aabf651fd0f": "EICAR Test File (SHA256)",
}

# ─── Suspicious byte patterns / magic bytes ───────────────────────────────────
SUSPICIOUS_PATTERNS: List[Tuple[bytes, str]] = [
    (b"This program cannot be run in DOS mode", "Embedded PE binary"),
    (b"MZ",                                      "PE magic bytes"),
    (b"\x4d\x5a\x90\x00",                        "PE header signature"),
    (b"TVqQAAMAAAAEAAAA",                          "Base64-encoded PE"),
    (b"powershell",                               "PowerShell reference"),
    (b"cmd.exe",                                  "cmd.exe reference"),
    (b"eval(base64_decode",                       "PHP webshell pattern"),
    (b"<script>document.write(unescape(",         "JS obfuscation pattern"),
    (b"EICAR-STANDARD-ANTIVIRUS-TEST-FILE",       "EICAR test string"),
    (b"\x60\x89\xe5\x31\xc0\x31\xdb\x31\xc9\x31\xd2", "Linux shellcode pattern"),
]

# ─── Known packer signatures ──────────────────────────────────────────────────
PACKER_SIGNATURES: List[Tuple[bytes, str]] = [
    (b"UPX0",          "UPX Packer"),
    (b"UPX1",          "UPX Packer"),
    (b"UPX!",          "UPX Packer"),
    (b".MPRESS1",      "MPRESS Packer"),
    (b"PECompact2",    "PECompact Packer"),
    (b"ASPack",        "ASPack Packer"),
    (b"Themida",       "Themida Protector"),
    (b"ExECryptor",    "ExECryptor Protector"),
]

# ─── Built-in lightweight YARA rules ──────────────────────────────────────────
BUILTIN_YARA_RULES = r"""
rule Detect_EICAR {
    meta:
        description = "Detects EICAR test string"
        severity    = "low"
    strings:
        $a = "EICAR-STANDARD-ANTIVIRUS-TEST-FILE"
    condition:
        $a
}

rule Detect_UPX_Packed {
    meta:
        description = "Detects UPX-packed executables"
        severity    = "medium"
    strings:
        $a = "UPX0"
        $b = "UPX!"
    condition:
        any of them
}

rule Detect_PowerShell_Download {
    meta:
        description = "PowerShell download cradle"
        severity    = "high"
    strings:
        $a = "DownloadString" nocase
        $b = "WebClient"     nocase
        $c = "Invoke-Expression" nocase wide ascii
    condition:
        2 of them
}

rule Detect_PHP_Webshell {
    meta:
        description = "PHP webshell indicators"
        severity    = "critical"
    strings:
        $a = "eval($_POST"  nocase
        $b = "eval($_GET"   nocase
        $c = "system($_"    nocase
        $d = "passthru($_"  nocase
    condition:
        any of them
}

rule Detect_Base64_PE {
    meta:
        description = "Base64-encoded PE file"
        severity    = "high"
    strings:
        $a = "TVqQAAMAAAAEAAAA"
        $b = "TVpAAA"
    condition:
        any of them
}

rule Detect_Reverse_Shell {
    meta:
        description = "Common reverse-shell patterns"
        severity    = "critical"
    strings:
        $a = "/bin/sh"     nocase
        $b = "bash -i"     nocase
        $c = "nc -e"       nocase
        $d = "ncat -e"     nocase
        $e = "/dev/tcp/"   nocase
    condition:
        2 of them
}

rule Detect_Crypto_Mining {
    meta:
        description = "Cryptocurrency mining indicators"
        severity    = "medium"
    strings:
        $a = "stratum+tcp://" nocase
        $b = "xmrig"          nocase
        $c = "monero"         nocase
        $d = "cryptonight"    nocase
    condition:
        any of them
}
"""


# ═══════════════════════════════════════════════════════════════════════════════
#  🔐  CHAIN OF CUSTODY
# ═══════════════════════════════════════════════════════════════════════════════

class ChainOfCustody:
    """
    Maintains an immutable, append-only audit trail for evidence handling.
    Each entry is signed with its own SHA-256 hash to detect tampering.
    """

    def __init__(self, case_id: str, examiner: str = "Unknown",
                 output_path: Optional[str] = None):
        self.case_id   = case_id
        self.examiner  = examiner
        self.record_id = str(uuid.uuid4())
        self.events: List[Dict[str, Any]] = []
        self.output_path = output_path
        self._lock = Lock()

        self._record_event("CUSTODY_OPENED", {
            "tool":     f"{TOOL_NAME} v{VERSION}",
            "platform": f"{platform.system()} {platform.release()}",
            "python":   platform.python_version(),
        })

    def _now_utc_str(self) -> str:
        """Return current UTC time as ISO-8601 string (timezone-aware, no deprecation)."""
        return datetime.now(timezone.utc).isoformat()

    def _record_event(self, action: str, details: Dict[str, Any]) -> Dict[str, Any]:
        """Append an event; compute a chained hash for integrity."""
        prev_hash = self.events[-1]["entry_hash"] if self.events else "GENESIS"
        entry = {
            "seq":        len(self.events) + 1,
            "timestamp":  self._now_utc_str(),
            "case_id":    self.case_id,
            "examiner":   self.examiner,
            "action":     action,
            "details":    details,
            "prev_hash":  prev_hash,
        }
        entry_str  = json.dumps(entry, sort_keys=True, default=str)
        entry["entry_hash"] = hashlib.sha256(
            (prev_hash + entry_str).encode()
        ).hexdigest()

        with self._lock:
            self.events.append(entry)

        if self.output_path:
            self._persist()

        return entry

    def _persist(self):
        """Write the chain to disk (JSON format)."""
        try:
            with open(self.output_path, "w", encoding="utf-8") as fh:
                json.dump({
                    "record_id": self.record_id,
                    "case_id":   self.case_id,
                    "events":    self.events,
                    "valid":     self.verify_integrity(),
                }, fh, indent=2, default=str)
        except Exception:
            pass

    def log_acquisition(self, file_path: str, hashes: Dict[str, str]):
        return self._record_event("EVIDENCE_ACQUIRED", {
            "file":   file_path,
            "hashes": hashes,
            "note":   "Baseline hash recorded at acquisition",
        })

    def log_analysis(self, file_path: str, operation: str, result_summary: str):
        return self._record_event("ANALYSIS_PERFORMED", {
            "file":      file_path,
            "operation": operation,
            "result":    result_summary,
        })

    def log_hash_verification(self, file_path: str,
                               original: Dict[str, str],
                               current: Dict[str, str]) -> bool:
        integrity_ok = all(original.get(k) == current.get(k) for k in original)
        self._record_event("HASH_VERIFICATION", {
            "file":         file_path,
            "integrity_ok": integrity_ok,
            "original":     original,
            "current":      current,
            "verdict":      "PASS" if integrity_ok else "FAIL – EVIDENCE MAY BE TAMPERED",
        })
        return integrity_ok

    def log_export(self, destination: str, format_: str):
        return self._record_event("EVIDENCE_EXPORTED", {
            "destination": destination,
            "format":      format_,
        })

    def close(self):
        return self._record_event("CUSTODY_CLOSED", {
            "total_events": len(self.events),
            "final_note":   "Custody chain sealed.",
        })

    def verify_integrity(self) -> bool:
        prev_hash = "GENESIS"
        for event in self.events:
            test = dict(event)
            stored_hash = test.pop("entry_hash")
            entry_str   = json.dumps(test, sort_keys=True, default=str)
            expected    = hashlib.sha256((prev_hash + entry_str).encode()).hexdigest()
            if expected != stored_hash:
                return False
            prev_hash = stored_hash
        return True

    def print_summary(self):
        integrity = self.verify_integrity()
        status_str = (f"{Fore.GREEN}✓ INTACT{Style.RESET_ALL}"
                      if integrity else
                      f"{Fore.RED}✗ COMPROMISED{Style.RESET_ALL}")

        print(f"\n{Fore.CYAN}{'─'*70}{Style.RESET_ALL}")
        print(f"{Fore.CYAN}  CHAIN OF CUSTODY  |  Case: {self.case_id}  |  Integrity: {status_str}{Style.RESET_ALL}")
        print(f"{Fore.CYAN}{'─'*70}{Style.RESET_ALL}")
        for ev in self.events:
            ts  = ev["timestamp"]
            act = ev["action"]
            seq = ev["seq"]
            details_short = ", ".join(
                f"{k}={v}" for k, v in ev["details"].items()
                if isinstance(v, (str, int, bool)) and k != "note"
            )[:80]
            print(f"  [{seq:02d}] {ts}  {Fore.YELLOW}{act:<26}{Style.RESET_ALL}  {details_short}")
        print(f"{Fore.CYAN}{'─'*70}{Style.RESET_ALL}")
        print(f"  Record ID : {self.record_id}")
        print(f"  Chain Hash: {self.events[-1]['entry_hash'][:48]}…")
        print()

    def to_dict(self) -> Dict[str, Any]:
        return {
            "record_id":       self.record_id,
            "case_id":         self.case_id,
            "examiner":        self.examiner,
            "integrity_valid": self.verify_integrity(),
            "event_count":     len(self.events),
            "events":          self.events,
        }


# ═══════════════════════════════════════════════════════════════════════════════
#  🧾  TIMELINE RECONSTRUCTOR
# ═══════════════════════════════════════════════════════════════════════════════

class TimelineReconstructor:
    """
    Aggregates every timestamp extracted from a file into a sorted
    chronological event log.
    """

    def __init__(self):
        self._events: List[Dict[str, Any]] = []

    def add_event(self, timestamp: Optional[str], source: str,
                  description: str, category: str = "general"):
        if not timestamp:
            return
        ts = self._normalise_ts(timestamp)
        if ts:
            self._events.append({
                "timestamp":   ts,
                "source":      source,
                "description": description,
                "category":    category,
            })

    @staticmethod
    def _normalise_ts(raw: str) -> Optional[str]:
        raw = str(raw).strip()
        formats = [
            "%Y:%m:%d %H:%M:%S",
            "%Y-%m-%dT%H:%M:%S",
            "%Y-%m-%dT%H:%M:%SZ",
            "%Y-%m-%d %H:%M:%S",
            "%Y-%m-%d",
            "%d/%m/%Y %H:%M:%S",
        ]
        for fmt in formats:
            try:
                return datetime.strptime(raw[:len(fmt)], fmt).isoformat()
            except ValueError:
                continue
        if re.match(r"\d{4}-\d{2}-\d{2}T\d{2}:\d{2}", raw):
            return raw[:19]
        return None

    def build(self) -> List[Dict[str, Any]]:
        return sorted(self._events, key=lambda e: e["timestamp"])

    def print_timeline(self):
        events = self.build()
        if not events:
            print("  (No timestamps found)")
            return

        print(f"\n{Fore.CYAN}{'─'*70}{Style.RESET_ALL}")
        print(f"{Fore.CYAN}  FILE ACTIVITY TIMELINE{Style.RESET_ALL}")
        print(f"{Fore.CYAN}{'─'*70}{Style.RESET_ALL}")

        prev_date = None
        for ev in events:
            date_part = ev["timestamp"][:10]
            if date_part != prev_date:
                print(f"\n  {Fore.YELLOW}▶ {date_part}{Style.RESET_ALL}")
                prev_date = date_part

            time_part = ev["timestamp"][11:19] if len(ev["timestamp"]) > 10 else ""
            cat_color = {
                "filesystem":  Fore.GREEN,
                "exif":        Fore.CYAN,
                "document":    Fore.BLUE,
                "audio":       Fore.MAGENTA,
                "executable":  Fore.RED,
                "general":     Fore.WHITE,
            }.get(ev["category"], Fore.WHITE)

            print(f"    {time_part}  {cat_color}[{ev['category']:<11}]{Style.RESET_ALL}"
                  f"  {ev['source']:<22}  {ev['description']}")

        print(f"\n{Fore.CYAN}{'─'*70}{Style.RESET_ALL}")
        print(f"  Total timeline events: {len(events)}")
        print()

    def to_dict(self) -> Dict[str, Any]:
        events = self.build()
        return {
            "total_events": len(events),
            "earliest":     events[0]["timestamp"]  if events else None,
            "latest":       events[-1]["timestamp"] if events else None,
            "events":       events,
        }


# ═══════════════════════════════════════════════════════════════════════════════
#  🧠  INTELLIGENCE / THREAT SCORING LAYER
# ═══════════════════════════════════════════════════════════════════════════════

class ThreatIntelligence:
    """
    Analyses extracted metadata and raw file bytes to produce:
      - Risk score (0–100)
      - IOC list
      - YARA matches
      - Packer / obfuscation flags
    """

    RISK_BANDS = [
        (80, "CRITICAL",  Fore.RED    if COLORAMA_AVAILABLE else ""),
        (60, "HIGH",      Fore.RED    if COLORAMA_AVAILABLE else ""),
        (40, "MEDIUM",    Fore.YELLOW if COLORAMA_AVAILABLE else ""),
        (20, "LOW",       Fore.CYAN   if COLORAMA_AVAILABLE else ""),
        (0,  "CLEAN",     Fore.GREEN  if COLORAMA_AVAILABLE else ""),
    ]

    def __init__(self, file_path: str):
        self.file_path    = Path(file_path)
        self.iocs:         List[str] = []
        self.yara_matches: List[Dict[str, Any]] = []
        self.risk_score   = 0
        self.risk_reasons: List[str] = []
        self._yara_rules  = self._compile_yara()

    @staticmethod
    def _compile_yara() -> Optional[Any]:
        if not YARA_AVAILABLE:
            return None
        try:
            return yara.compile(source=BUILTIN_YARA_RULES)
        except Exception:
            return None

    def _run_yara(self, data: bytes) -> List[Dict[str, Any]]:
        if not self._yara_rules:
            return []
        try:
            matches = self._yara_rules.match(data=data)
            results = []
            for m in matches:
                severity = m.meta.get("severity", "unknown")
                score_bump = {"low": 10, "medium": 25, "high": 40, "critical": 60}.get(severity, 15)
                results.append({
                    "rule":        m.rule,
                    "description": m.meta.get("description", ""),
                    "severity":    severity,
                    "score_bump":  score_bump,
                    "tags":        list(m.tags),
                    "strings":     [(hex(offset), identifier)
                                    for offset, identifier, _ in m.strings][:5],
                })
                self.risk_score += score_bump
                self.risk_reasons.append(f"YARA:{m.rule} (+{score_bump})")
            return results
        except Exception:
            return []

    def _check_known_hashes(self, hashes: Dict[str, str]):
        for algo, digest in hashes.items():
            if digest.lower() in KNOWN_MALICIOUS_HASHES:
                threat = KNOWN_MALICIOUS_HASHES[digest.lower()]
                self.iocs.append(f"Known-malicious {algo.upper()}: {digest[:16]}… → {threat}")
                self.risk_score += 100
                self.risk_reasons.append(f"KnownMaliciousHash:{algo}=+100")

    def _check_packers(self, data: bytes):
        for sig, name in PACKER_SIGNATURES:
            if sig in data:
                self.iocs.append(f"Packer/Protector detected: {name}")
                self.risk_score += 30
                self.risk_reasons.append(f"Packer:{name}=+30")

    def _check_suspicious_patterns(self, data: bytes):
        ext = self.file_path.suffix.lower()
        for pattern, description in SUSPICIOUS_PATTERNS:
            if pattern in (b"MZ", b"\x4d\x5a\x90\x00") and ext in (".exe", ".dll"):
                continue
            if pattern in data:
                self.iocs.append(f"Suspicious pattern: {description}")
                self.risk_score += 20
                self.risk_reasons.append(f"Pattern:{description}=+20")

    def _score_entropy(self, entropy: float):
        LOW_ENTROPY_COMPRESSED = {
            ".jpg", ".jpeg", ".png", ".gif", ".webp",
            ".zip", ".gz", ".rar", ".7z", ".mp3", ".mp4",
            ".docx", ".xlsx", ".pptx", ".pdf",
        }
        ext = self.file_path.suffix.lower()
        if entropy > 7.95 and ext not in LOW_ENTROPY_COMPRESSED:
            self.risk_score += 30
            self.risk_reasons.append(f"HighEntropy({entropy:.2f})=+30")
            self.iocs.append(f"Very high entropy ({entropy:.2f}) – possible encryption/packing")
        elif entropy > 7.5 and ext not in LOW_ENTROPY_COMPRESSED:
            self.risk_score += 15
            self.risk_reasons.append(f"ElevatedEntropy({entropy:.2f})=+15")

    def _score_suspicious_flags(self, flags: List[str]):
        for flag in flags:
            self.iocs.append(f"Anomaly: {flag}")
            self.risk_score += 25
            self.risk_reasons.append("SuspiciousFlag=+25")

    def _check_pe_timestamp(self, pe_meta: Dict[str, Any]):
        ts_str = pe_meta.get("timestamp")
        if not ts_str:
            return
        try:
            pe_ts = datetime.fromisoformat(ts_str)
            now   = datetime.now()
            if pe_ts.year < 2000:
                self.iocs.append(
                    f"PE compile timestamp is very old ({pe_ts.year}) – possible epoch-zero tampering"
                )
                self.risk_score += 20
                self.risk_reasons.append("PETimestampAnomaly=+20")
            if pe_ts > now:
                self.iocs.append(
                    "PE compile timestamp is in the future – possible anti-forensic manipulation"
                )
                self.risk_score += 35
                self.risk_reasons.append("PEFutureTimestamp=+35")
        except Exception:
            pass

    def _check_exif_anomalies(self, exif_meta: Dict[str, Any]):
        if "GPSInfo" in exif_meta or "GPS" in str(exif_meta):
            self.iocs.append("GPS coordinates found in file – potential location privacy leak")
            self.risk_score += 5
            self.risk_reasons.append("GPSData=+5")

    def analyse(self, metadata: Dict[str, Any]) -> Dict[str, Any]:
        hashes     = metadata.get("hashes", {})
        forensic   = metadata.get("forensic_info", {})
        extracted  = metadata.get("extracted_metadata", {})
        suspicious = forensic.get("suspicious_flags", [])
        entropy    = forensic.get("entropy", 0.0)
        pe_meta    = extracted.get("pe_file", {})
        exif_meta  = extracted.get("exif", {})

        self._check_known_hashes(hashes)

        if isinstance(entropy, (int, float)):
            self._score_entropy(float(entropy))

        self._score_suspicious_flags(suspicious)

        if pe_meta:
            self._check_pe_timestamp(pe_meta)

        if exif_meta:
            self._check_exif_anomalies(exif_meta)

        file_size = metadata.get("file_info", {}).get("file_size_bytes", 0)
        if file_size and file_size < 100 * 1024 * 1024:
            try:
                with open(self.file_path, "rb") as fh:
                    data = fh.read()
                self.yara_matches = self._run_yara(data)
                self._check_packers(data)
                self._check_suspicious_patterns(data)
            except Exception:
                pass

        self.risk_score = min(100, self.risk_score)

        risk_level, risk_color = "CLEAN", Fore.GREEN if COLORAMA_AVAILABLE else ""
        for threshold, label, color in self.RISK_BANDS:
            if self.risk_score >= threshold:
                risk_level, risk_color = label, color
                break

        return {
            "risk_score":     self.risk_score,
            "risk_level":     risk_level,
            "risk_color":     risk_color,
            "risk_reasons":   self.risk_reasons,
            "iocs":           list(dict.fromkeys(self.iocs)),
            "yara_matches":   self.yara_matches,
            "yara_available": YARA_AVAILABLE,
        }

    def print_report(self, report: Dict[str, Any]):
        color = report.get("risk_color", "")
        reset = Style.RESET_ALL if COLORAMA_AVAILABLE else ""
        score = report["risk_score"]
        level = report["risk_level"]

        bar_len = 40
        filled  = int(bar_len * score / 100)
        bar     = "█" * filled + "░" * (bar_len - filled)

        print(f"\n{Fore.CYAN}{'─'*70}{Style.RESET_ALL}")
        print(f"{Fore.CYAN}  🧠  THREAT INTELLIGENCE REPORT{Style.RESET_ALL}")
        print(f"{Fore.CYAN}{'─'*70}{Style.RESET_ALL}")
        print(f"  Risk Score : {color}{score:>3}/100  [{bar}]  {level}{reset}")
        print()

        if report["iocs"]:
            print(f"  {Fore.RED}Indicators of Compromise (IOCs):{reset}")
            for ioc in report["iocs"]:
                print(f"    ⚑  {ioc}")
            print()

        if report["yara_matches"]:
            print(f"  {Fore.RED}YARA Rule Matches:{reset}")
            for m in report["yara_matches"]:
                sev_color = {
                    "critical": Fore.RED,
                    "high":     Fore.RED,
                    "medium":   Fore.YELLOW,
                    "low":      Fore.CYAN,
                }.get(m["severity"], Fore.WHITE) if COLORAMA_AVAILABLE else ""
                print(f"    ⚡ [{sev_color}{m['severity'].upper()}{reset}] "
                      f"{m['rule']} – {m['description']}")
        elif not YARA_AVAILABLE:
            print(f"  {Fore.YELLOW}⚠ YARA not installed – install with: pip install yara-python{reset}")

        if report["risk_reasons"]:
            print(f"\n  Score breakdown:")
            for r in report["risk_reasons"]:
                print(f"    + {r}")

        print(f"{Fore.CYAN}{'─'*70}{Style.RESET_ALL}\n")


# ═══════════════════════════════════════════════════════════════════════════════
#  FORENSIC LOGGER
# ═══════════════════════════════════════════════════════════════════════════════

class ForensicLogger:
    """Structured logging for forensic analysis"""

    def __init__(self, log_file: Optional[str] = None, case_id: Optional[str] = None):
        self.case_id  = case_id or f"CASE-{datetime.now().strftime('%Y%m%d-%H%M%S')}"
        self.log_file = log_file
        self.lock     = Lock()

        self.logger = logging.getLogger(TOOL_NAME)
        self.logger.setLevel(logging.DEBUG)

        # Avoid adding duplicate handlers when called multiple times
        if not self.logger.handlers:
            console_handler = logging.StreamHandler()
            console_handler.setLevel(logging.INFO)
            console_handler.setFormatter(
                logging.Formatter('%(asctime)s - %(levelname)s - %(message)s',
                                  datefmt='%Y-%m-%d %H:%M:%S')
            )
            self.logger.addHandler(console_handler)

        if log_file:
            file_handler = logging.FileHandler(log_file)
            file_handler.setLevel(logging.DEBUG)
            self.logger.addHandler(file_handler)

    def log_event(self, level: str, message: str, **kwargs):
        entry = {
            'timestamp':    datetime.now(timezone.utc).isoformat(),
            'case_id':      self.case_id,
            'level':        level,
            'message':      message,
            'tool_version': VERSION,
            **kwargs
        }
        with self.lock:
            getattr(self.logger, level.lower(), self.logger.debug)(
                json.dumps(entry)
            )

    def info(self, message, **kw):     self.log_event('INFO',     message, **kw)
    def warning(self, message, **kw):  self.log_event('WARNING',  message, **kw)
    def error(self, message, **kw):    self.log_event('ERROR',    message, **kw)
    def critical(self, message, **kw): self.log_event('CRITICAL', message, **kw)


# ═══════════════════════════════════════════════════════════════════════════════
#  EXIFTOOL WRAPPER
# ═══════════════════════════════════════════════════════════════════════════════

class ExifToolWrapper:
    def __init__(self):
        self.available = self._check_exiftool()

    def _check_exiftool(self) -> bool:
        try:
            r = subprocess.run(['exiftool', '-ver'], capture_output=True,
                               text=True, timeout=5)
            return r.returncode == 0
        except (FileNotFoundError, subprocess.TimeoutExpired):
            return False

    def extract_metadata(self, file_path: str) -> Optional[Dict[str, Any]]:
        if not self.available:
            return None
        try:
            r = subprocess.run(
                ['exiftool', '-j', '-G', '-a', '-s', str(file_path)],
                capture_output=True, text=True, timeout=EXIFTOOL_TIMEOUT
            )
            if r.returncode == 0 and r.stdout:
                data = json.loads(r.stdout)
                return data[0] if data else None
        except Exception:
            return None


# ═══════════════════════════════════════════════════════════════════════════════
#  FILE TYPE DETECTOR
# ═══════════════════════════════════════════════════════════════════════════════

class FileTypeDetector:
    def __init__(self):
        self.available = MAGIC_AVAILABLE

    def detect(self, file_path: str) -> Dict[str, str]:
        result = {
            'extension':   Path(file_path).suffix.lower(),
            'mime_type':   None,
            'description': None,
            'real_type':   None,
        }
        if self.available:
            try:
                result['mime_type']   = magic.Magic(mime=True).from_file(file_path)
                result['description'] = magic.Magic().from_file(file_path)
                result['real_type']   = self._classify_type(result['mime_type'], result['extension'])
            except Exception:
                pass
        if not result['real_type']:
            result['real_type'] = self._classify_type(None, result['extension'])
        return result

    def _classify_type(self, mime_type, extension=None):
        ext = (extension or "").lower()
        if ext in ['.e01', '.aff4', '.dd', '.img', '.iso']:
            return 'disk_image'
        if not mime_type:
            return None
        if mime_type.startswith('image/'):                       return 'image'
        if mime_type.startswith('video/'):                       return 'video'
        if mime_type.startswith('audio/'):                       return 'audio'
        if 'pdf' in mime_type:                                   return 'pdf'
        if 'word' in mime_type or 'document' in mime_type:      return 'document'
        if 'zip' in mime_type or 'compressed' in mime_type:     return 'archive'
        if 'executable' in mime_type or mime_type == 'application/x-dosexec': return 'executable'
        return 'unknown'


# ═══════════════════════════════════════════════════════════════════════════════
#  DISK IMAGE HANDLER
# ═══════════════════════════════════════════════════════════════════════════════

class DiskImageHandler:
    def __init__(self, image_path: str, logger: Optional[ForensicLogger] = None):
        self.image_path = str(image_path)
        self.logger     = logger
        self.img_info   = None
        self.partitions = []

    def open_image(self):
        if not TSK_AVAILABLE:
            raise ImportError("pytsk3 is required for disk image handling")
        ext = Path(self.image_path).suffix.lower()
        try:
            if ext in ['.e01', '.s01'] and EWF_AVAILABLE:
                filenames  = pyewf.get_filenames(self.image_path)
                ewf_handle = pyewf.handle()
                ewf_handle.open(filenames)
                self.img_info = pytsk3.Img_Info(
                    url="", type=pytsk3.TSK_IMG_TYPE_EXTERNAL,
                    external_handle=ewf_handle
                )
            else:
                self.img_info = pytsk3.Img_Info(self.image_path)
            return True
        except Exception as e:
            if self.logger:
                self.logger.error(f"Failed to open disk image: {e}")
            return False

    def list_partitions(self):
        partitions = []
        if not self.img_info:
            return partitions
        try:
            volume = pytsk3.Volume_Info(self.img_info)
            for part in volume:
                partitions.append({
                    'id':          part.addr,
                    'description': part.desc.decode('utf-8', errors='ignore'),
                    'start':       part.start,
                    'length':      part.len,
                    'flags':       part.flags,
                })
        except Exception:
            partitions.append({
                'id': 0, 'description': 'Single Partition / Raw Volume',
                'start': 0, 'length': self.img_info.get_size(),
                'flags': pytsk3.TSK_VS_PART_FLAG_ALLOC,
            })
        self.partitions = partitions
        return partitions

    def walk_filesystem(self, partition_id: int = 0) -> List[Dict[str, Any]]:
        files_metadata = []
        if not self.img_info:
            return files_metadata
        try:
            offset = 0
            for p in self.partitions:
                if p['id'] == partition_id:
                    offset = p['start'] * 512
                    break
            fs       = pytsk3.FS_Info(self.img_info, offset=offset)
            root_dir = fs.open_dir(path="/")

            def _walk(directory, current_path=""):
                for entry in directory:
                    if entry.info.name.name in [b".", b".."]:
                        continue
                    name      = entry.info.name.name.decode('utf-8', errors='ignore')
                    full_path = f"{current_path}/{name}"
                    meta      = {
                        'name':     name,
                        'path':     full_path,
                        'size':     entry.info.meta.size if entry.info.meta else 0,
                        'type':     'directory' if entry.info.meta and
                                    entry.info.meta.type == pytsk3.TSK_FS_META_TYPE_DIR
                                    else 'file',
                        'inode':    entry.info.meta.addr if entry.info.meta else 0,
                        'created':  datetime.fromtimestamp(entry.info.meta.crtime).isoformat()
                                    if entry.info.meta and hasattr(entry.info.meta, 'crtime') else None,
                        'modified': datetime.fromtimestamp(entry.info.meta.mtime).isoformat()
                                    if entry.info.meta and hasattr(entry.info.meta, 'mtime') else None,
                        'accessed': datetime.fromtimestamp(entry.info.meta.atime).isoformat()
                                    if entry.info.meta and hasattr(entry.info.meta, 'atime') else None,
                    }
                    files_metadata.append(meta)
                    if meta['type'] == 'directory':
                        try:
                            _walk(entry.as_directory(), full_path)
                        except Exception:
                            pass

            _walk(root_dir)
        except Exception as e:
            if self.logger:
                self.logger.error(f"Error walking filesystem on partition {partition_id}: {e}")
        return files_metadata


# ═══════════════════════════════════════════════════════════════════════════════
#  ENHANCED METADATA EXTRACTOR
# ═══════════════════════════════════════════════════════════════════════════════

class EnhancedMetadataExtractor:
    """Enhanced metadata extraction with forensic capabilities"""

    def __init__(self, file_path: str,
                 case_id: Optional[str] = None,
                 logger: Optional[ForensicLogger] = None,
                 chain_of_custody: Optional[ChainOfCustody] = None,
                 examiner: str = "Unknown"):
        self.file_path        = Path(file_path)
        self.case_id          = case_id
        self.logger           = logger or ForensicLogger(case_id=case_id)
        self.exiftool         = ExifToolWrapper()
        self.file_detector    = FileTypeDetector()
        self.chain_of_custody = chain_of_custody
        self.examiner         = examiner
        self.timeline         = TimelineReconstructor()

        self.metadata = {
            'forensic_info':      {},
            'file_info':          {},
            'file_type_analysis': {},
            'extracted_metadata': {},
            'hashes':             {},
            'exiftool_metadata':  {},
            'timeline':           {},
            'threat_intel':       {},
            'chain_of_custody':   {},
            'warnings':           [],
            'errors':             [],
        }

    # ── Helpers ───────────────────────────────────────────────────────────────
    @staticmethod
    def _calculate_entropy(data: bytes) -> float:
        if not data:
            return 0.0
        freq = [0] * 256
        for byte in data:
            freq[byte] += 1
        entropy = 0.0
        total   = len(data)
        for c in freq:
            if c > 0:
                p = c / total
                entropy -= p * math.log(p, 2)
        return entropy

    @staticmethod
    def _human_readable_size(size_bytes: int) -> str:
        for unit in ['B', 'KB', 'MB', 'GB', 'TB']:
            if size_bytes < 1024.0:
                return f"{size_bytes:.2f} {unit}"
            size_bytes /= 1024.0
        return f"{size_bytes:.2f} PB"

    # ── Main entry ────────────────────────────────────────────────────────────
    def extract_all(self) -> Dict[str, Any]:
        try:
            if not self.file_path.exists():
                raise FileNotFoundError(f"File not found: {self.file_path}")

            self.logger.info("Starting metadata extraction", file=str(self.file_path))

            self._add_forensic_metadata()
            self._extract_file_type()
            self._extract_file_info()
            self._extract_hashes()

            if self.chain_of_custody:
                self.chain_of_custody.log_acquisition(
                    str(self.file_path), self.metadata['hashes']
                )

            self._extract_exiftool_metadata()
            self._extract_hachoir_metadata()
            self._extract_type_specific_metadata()
            self._detect_suspicious()

            self._build_timeline()
            self.metadata['timeline'] = self.timeline.to_dict()

            ti        = ThreatIntelligence(str(self.file_path))
            ti_report = ti.analyse(self.metadata)
            self.metadata['threat_intel'] = ti_report

            if self.chain_of_custody:
                self.chain_of_custody.log_analysis(
                    str(self.file_path),
                    "Full metadata extraction",
                    f"Risk={ti_report['risk_score']}, IOCs={len(ti_report['iocs'])}, "
                    f"YARA={len(ti_report['yara_matches'])}",
                )
                self.metadata['chain_of_custody'] = self.chain_of_custody.to_dict()

            self.logger.info("Metadata extraction completed", file=str(self.file_path))

        except Exception as e:
            self.logger.error(f"Extraction failed: {e}", file=str(self.file_path))
            self.metadata['errors'].append(str(e))

        return self.metadata

    # ── Timeline builder ──────────────────────────────────────────────────────
    def _build_timeline(self):
        fi = self.metadata.get('file_info', {})

        self.timeline.add_event(fi.get('created_time'),  "Filesystem", "File created",         "filesystem")
        self.timeline.add_event(fi.get('modified_time'), "Filesystem", "File last modified",    "filesystem")
        self.timeline.add_event(fi.get('accessed_time'), "Filesystem", "File last accessed",    "filesystem")

        exif = self.metadata.get('extracted_metadata', {}).get('exif', {})
        self.timeline.add_event(exif.get('DateTimeOriginal'),  "EXIF", "Photo taken",           "exif")
        self.timeline.add_event(exif.get('DateTime'),          "EXIF", "Image file date/time",  "exif")
        self.timeline.add_event(exif.get('DateTimeDigitized'), "EXIF", "Image digitised",       "exif")

        docx = self.metadata.get('extracted_metadata', {}).get('docx', {})
        self.timeline.add_event(docx.get('created'),       "DOCX", "Document created",          "document")
        self.timeline.add_event(docx.get('modified'),      "DOCX", "Document last saved",       "document")
        self.timeline.add_event(docx.get('last_printed'),  "DOCX", "Document last printed",     "document")

        pdf_meta = self.metadata.get('extracted_metadata', {}).get('pdf', {}).get('metadata', {})
        self.timeline.add_event(pdf_meta.get('CreationDate'), "PDF", "PDF created",             "document")
        self.timeline.add_event(pdf_meta.get('ModDate'),      "PDF", "PDF modified",            "document")

        pe = self.metadata.get('extracted_metadata', {}).get('pe_file', {})
        self.timeline.add_event(pe.get('timestamp'),       "PE",  "Binary compiled",            "executable")

        et = self.metadata.get('exiftool_metadata', {})
        self.timeline.add_event(et.get('File:FileModifyDate'), "ExifTool", "File modify (ET)",  "filesystem")
        self.timeline.add_event(et.get('File:FileCreateDate'), "ExifTool", "File create (ET)",  "filesystem")

    # ── Suspicious detection ──────────────────────────────────────────────────
    def _detect_suspicious(self):
        suspicious = []
        COMPRESSED_EXT = {
            '.jpg', '.jpeg', '.png', '.gif', '.webp', '.heic',
            '.zip', '.rar', '.7z', '.gz', '.tar', '.jar', '.apk',
            '.pdf', '.docx', '.xlsx', '.pptx', '.mp3', '.mp4',
            '.avi', '.mov', '.mkv', '.flac',
        }
        try:
            file_size = self.file_path.stat().st_size
            if file_size < 100 * 1024 * 1024:
                with open(self.file_path, 'rb') as f:
                    data = f.read()
                entropy = self._calculate_entropy(data)
                self.metadata['forensic_info']['entropy'] = round(entropy, 4)

                ext = self.file_path.suffix.lower()
                if entropy > 7.95 and ext not in COMPRESSED_EXT:
                    suspicious.append(
                        f"Extremely high entropy ({entropy:.2f}) for {ext} – possible encryption/packing"
                    )
                elif entropy < 1.0 and file_size > 4096:
                    suspicious.append(f"Very low entropy ({entropy:.2f}) – large blocks of uniform data")

            if self.file_path.suffix.lower() in ['.jpg', '.jpeg']:
                with open(self.file_path, 'rb') as f:
                    content = f.read()
                eoi = content.rfind(b'\xff\xd9')
                if eoi != -1 and eoi + 2 < len(content):
                    extra = content[eoi + 2:]
                    if len(extra) > 4096 and not all(b in b'\x00\xff\r\n\t ' for b in extra):
                        suspicious.append(
                            f"Trailing data after JPEG EOF: {len(extra)} bytes (possible steganography)"
                        )

            if self.file_path.suffix.lower() == '.png':
                with open(self.file_path, 'rb') as f:
                    content = f.read()
                iend = content.rfind(b'IEND')
                if iend != -1:
                    end = iend + 8
                    if end < len(content):
                        extra = content[end:]
                        if len(extra) > 4096 and not all(b in b'\x00\xff\r\n\t ' for b in extra):
                            suspicious.append(
                                f"Trailing data after PNG IEND: {len(extra)} bytes (possible steganography)"
                            )

        except Exception as e:
            self.metadata['warnings'].append(f"Suspicious detection error: {e}")

        if suspicious:
            self.metadata['forensic_info']['suspicious_flags'] = suspicious
            for s in suspicious:
                self.metadata['warnings'].append(f"SUSPICIOUS: {s}")

    # ── Forensic metadata ─────────────────────────────────────────────────────
    def _add_forensic_metadata(self):
        self.metadata['forensic_info'] = {
            'case_id':                    self.case_id,
            'tool_name':                  TOOL_NAME,
            'tool_version':               VERSION,
            'extraction_timestamp_utc':   datetime.now(timezone.utc).isoformat(),
            'extraction_timestamp_local': datetime.now().isoformat(),
            'examiner':                   self.examiner,
            'examiner_system': {
                'platform':         platform.system(),
                'platform_release': platform.release(),
                'platform_version': platform.version(),
                'machine':          platform.machine(),
                'processor':        platform.processor(),
                'python_version':   platform.python_version(),
            },
            'capabilities': {
                'exiftool_available':    self.exiftool.available,
                'hachoir_available':     HACHOIR_AVAILABLE,
                'libmagic_available':    self.file_detector.available,
                'pil_available':         PIL_AVAILABLE,
                'pypdf2_available':      PDF_AVAILABLE,
                'python_docx_available': DOCX_AVAILABLE,
                'mutagen_available':     AUDIO_AVAILABLE,
                'pefile_available':      PE_AVAILABLE,
                'yara_available':        YARA_AVAILABLE,
            },
        }

    def _extract_file_type(self):
        self.metadata['file_type_analysis'] = self.file_detector.detect(str(self.file_path))

    def _extract_file_info(self):
        stat = self.file_path.stat()
        self.metadata['file_info'] = {
            'filename':        self.file_path.name,
            'full_path':       str(self.file_path.absolute()),
            'file_size_bytes': stat.st_size,
            'file_size_human': self._human_readable_size(stat.st_size),
            'extension':       self.file_path.suffix,
            'created_time':    datetime.fromtimestamp(stat.st_ctime).isoformat(),
            'modified_time':   datetime.fromtimestamp(stat.st_mtime).isoformat(),
            'accessed_time':   datetime.fromtimestamp(stat.st_atime).isoformat(),
            'permissions':     oct(stat.st_mode)[-3:],
        }

    def _extract_hashes(self):
        md5    = hashlib.md5()
        sha1   = hashlib.sha1()
        sha256 = hashlib.sha256()
        try:
            with open(self.file_path, 'rb') as f:
                for chunk in iter(lambda: f.read(4096), b""):
                    md5.update(chunk)
                    sha1.update(chunk)
                    sha256.update(chunk)
            self.metadata['hashes'] = {
                'md5':    md5.hexdigest(),
                'sha1':   sha1.hexdigest(),
                'sha256': sha256.hexdigest(),
            }
        except Exception as e:
            self.metadata['errors'].append(f"Hash calculation failed: {e}")

    def _extract_exiftool_metadata(self):
        if self.exiftool.available:
            data = self.exiftool.extract_metadata(str(self.file_path))
            if data:
                self.metadata['exiftool_metadata'] = data
            else:
                self.metadata['warnings'].append("ExifTool returned no data")
        else:
            self.metadata['warnings'].append("ExifTool not available")

    def _extract_hachoir_metadata(self):
        if not HACHOIR_AVAILABLE:
            return
        try:
            import contextlib
            import io as _io
            hachoir_data = {}
            with contextlib.redirect_stderr(_io.StringIO()):
                parser = createParser(str(self.file_path))
                if parser:
                    with parser:
                        meta = extractMetadata(parser)
                        if meta:
                            for item in meta:
                                if item.values:
                                    vals = [v.text for v in item.values]
                                    hachoir_data[item.key] = vals[0] if len(vals) == 1 else vals
            if hachoir_data:
                self.metadata['extracted_metadata']['hachoir'] = hachoir_data
        except Exception:
            pass

    def _extract_type_specific_metadata(self):
        ft  = self.metadata['file_type_analysis'].get('real_type')
        ext = self.file_path.suffix.lower()

        if ft == 'image' or ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp']:
            self._extract_image_metadata()
        elif ft == 'pdf' or ext == '.pdf':
            self._extract_pdf_metadata()
        elif ft == 'document' or ext in ['.docx', '.doc']:
            self._extract_docx_metadata()
        elif ft == 'audio' or ext in ['.mp3', '.mp4', '.m4a', '.flac', '.ogg', '.wav']:
            self._extract_audio_metadata()
        elif ft == 'archive' or ext in ['.zip', '.jar', '.apk']:
            self._extract_zip_metadata()
        elif ft == 'video' or ext in ['.mp4', '.avi', '.mkv', '.mov', '.wmv']:
            self._extract_video_metadata()
        elif ft == 'executable' or ext in ['.exe', '.dll']:
            self._extract_pe_metadata()
        elif ft == 'disk_image' or ext in ['.e01', '.aff4', '.dd', '.img', '.iso']:
            self._extract_disk_image_metadata()

    def _extract_disk_image_metadata(self):
        if not TSK_AVAILABLE:
            self.metadata['warnings'].append("pytsk3 not installed")
            return
        try:
            handler = DiskImageHandler(str(self.file_path), logger=self.logger)
            if handler.open_image():
                partitions = handler.list_partitions()
                disk_info  = {
                    'image_path':       str(self.file_path),
                    'total_size':       handler.img_info.get_size(),
                    'total_size_human': self._human_readable_size(handler.img_info.get_size()),
                    'partition_count':  len(partitions),
                    'partitions':       partitions,
                }
                if partitions:
                    data_part = next(
                        (p for p in partitions if p['flags'] == pytsk3.TSK_VS_PART_FLAG_ALLOC),
                        partitions[0]
                    )
                    disk_info['sample_files'] = handler.walk_filesystem(data_part['id'])[:50]
                self.metadata['extracted_metadata']['disk_image'] = disk_info
        except Exception as e:
            self.metadata['errors'].append(f"Disk image analysis failed: {e}")

    def _extract_image_metadata(self):
        if not PIL_AVAILABLE:
            self.metadata['warnings'].append("PIL/Pillow not installed")
            return
        try:
            image      = Image.open(self.file_path)
            image_info = {
                'format': image.format,
                'mode':   image.mode,
                'size':   f"{image.width}x{image.height}",
                'width':  image.width,
                'height': image.height,
            }
            exif_data = {}
            raw_exif  = None
            # Pillow ≥ 10 exposes getexif(); older versions use _getexif()
            if hasattr(image, 'getexif'):
                raw_exif = image.getexif()
            elif hasattr(image, '_getexif') and callable(image._getexif):
                raw_exif = image._getexif()

            if raw_exif:
                for tag_id, value in raw_exif.items():
                    tag = TAGS.get(tag_id, tag_id)
                    if tag == "GPSInfo":
                        gps = {}
                        for gps_id in value:
                            gps[GPSTAGS.get(gps_id, gps_id)] = value[gps_id]
                        if gps:
                            coords = self._extract_gps_coordinates(gps)
                            if coords:
                                gps['coordinates'] = coords
                        exif_data[tag] = gps
                    else:
                        if isinstance(value, bytes):
                            try:
                                value = value.decode('utf-8', errors='ignore')
                            except Exception:
                                value = str(value)
                        exif_data[tag] = value

            self.metadata['extracted_metadata']['image'] = image_info
            if exif_data:
                self.metadata['extracted_metadata']['exif'] = exif_data
        except Exception as e:
            self.metadata['errors'].append(f"Image extraction failed: {e}")

    def _extract_gps_coordinates(self, gps_data):
        try:
            def to_deg(v):
                d, m, s = v
                return float(d) + float(m) / 60.0 + float(s) / 3600.0
            lat     = gps_data.get('GPSLatitude')
            lat_ref = gps_data.get('GPSLatitudeRef')
            lon     = gps_data.get('GPSLongitude')
            lon_ref = gps_data.get('GPSLongitudeRef')
            if lat and lon and lat_ref and lon_ref:
                latitude  = to_deg(lat) * (-1 if lat_ref == 'S' else 1)
                longitude = to_deg(lon) * (-1 if lon_ref == 'W' else 1)
                return {
                    'latitude':        latitude,
                    'longitude':       longitude,
                    'google_maps_url': f"https://www.google.com/maps?q={latitude},{longitude}",
                }
        except Exception as e:
            self.metadata['warnings'].append(f"GPS extraction failed: {e}")
        return None

    def _extract_pdf_metadata(self):
        if not PDF_AVAILABLE:
            self.metadata['warnings'].append("PyPDF2 not installed")
            return
        try:
            with open(self.file_path, 'rb') as f:
                reader   = PyPDF2.PdfReader(f)
                pdf_info = {
                    'num_pages':    len(reader.pages),
                    'is_encrypted': reader.is_encrypted,
                }
                if reader.metadata:
                    pdf_info['metadata'] = {
                        k.lstrip('/'): v for k, v in reader.metadata.items()
                    }
                self.metadata['extracted_metadata']['pdf'] = pdf_info
        except Exception as e:
            self.metadata['errors'].append(f"PDF extraction failed: {e}")

    def _extract_docx_metadata(self):
        if not DOCX_AVAILABLE:
            self.metadata['warnings'].append("python-docx not installed")
            return
        try:
            doc = Document(self.file_path)
            cp  = doc.core_properties
            self.metadata['extracted_metadata']['docx'] = {
                'author':           cp.author,
                'category':         cp.category,
                'comments':         cp.comments,
                'content_status':   cp.content_status,
                'created':          cp.created.isoformat()      if cp.created      else None,
                'identifier':       cp.identifier,
                'keywords':         cp.keywords,
                'language':         cp.language,
                'last_modified_by': cp.last_modified_by,
                'last_printed':     cp.last_printed.isoformat() if cp.last_printed else None,
                'modified':         cp.modified.isoformat()     if cp.modified     else None,
                'revision':         cp.revision,
                'subject':          cp.subject,
                'title':            cp.title,
                'version':          cp.version,
                'num_paragraphs':   len(doc.paragraphs),
                'num_tables':       len(doc.tables),
            }
        except Exception as e:
            self.metadata['errors'].append(f"DOCX extraction failed: {e}")

    def _extract_audio_metadata(self):
        if not AUDIO_AVAILABLE:
            self.metadata['warnings'].append("mutagen not installed")
            return
        try:
            audio = mutagen.File(self.file_path)
            if audio is None:
                self.metadata['warnings'].append("Could not read audio file")
                return
            audio_info = {
                'length_seconds': audio.info.length      if hasattr(audio.info, 'length')      else None,
                'bitrate':        audio.info.bitrate     if hasattr(audio.info, 'bitrate')     else None,
                'sample_rate':    audio.info.sample_rate if hasattr(audio.info, 'sample_rate') else None,
                'channels':       audio.info.channels    if hasattr(audio.info, 'channels')    else None,
            }
            if audio.tags:
                audio_info['tags'] = {k: str(v) for k, v in audio.tags.items()}
            self.metadata['extracted_metadata']['audio'] = audio_info
        except Exception as e:
            self.metadata['errors'].append(f"Audio extraction failed: {e}")

    def _extract_zip_metadata(self):
        try:
            if not zipfile.is_zipfile(self.file_path):
                return
            with zipfile.ZipFile(self.file_path, 'r') as zf:
                zip_info = {
                    'file_count':             len(zf.namelist()),
                    'files':                  [],
                    'total_uncompressed_size': 0,
                    'total_compressed_size':   0,
                }
                for info in zf.infolist():
                    zip_info['files'].append({
                        'filename':          info.filename,
                        'compressed_size':   info.compress_size,
                        'uncompressed_size': info.file_size,
                        'compression_ratio': (
                            f"{(1 - info.compress_size / info.file_size) * 100:.1f}%"
                            if info.file_size > 0 else "0%"
                        ),
                        'date_time': datetime(*info.date_time).isoformat(),
                        'crc':       hex(info.CRC),
                    })
                    zip_info['total_uncompressed_size'] += info.file_size
                    zip_info['total_compressed_size']   += info.compress_size
                zip_info['total_uncompressed_size_human'] = self._human_readable_size(
                    zip_info['total_uncompressed_size']
                )
                zip_info['total_compressed_size_human'] = self._human_readable_size(
                    zip_info['total_compressed_size']
                )
                self.metadata['extracted_metadata']['zip'] = zip_info
        except Exception as e:
            self.metadata['errors'].append(f"ZIP extraction failed: {e}")

    def _extract_video_metadata(self):
        try:
            r = subprocess.run(
                ['ffprobe', '-v', 'quiet', '-print_format', 'json',
                 '-show_format', '-show_streams', str(self.file_path)],
                capture_output=True, text=True, timeout=30
            )
            if r.returncode == 0 and r.stdout:
                vd  = json.loads(r.stdout)
                fmt = vd.get('format', {})
                self.metadata['extracted_metadata']['video'] = {
                    'format':      vd.get('format', {}),
                    'streams':     vd.get('streams', []),
                    'duration':    fmt.get('duration'),
                    'size':        fmt.get('size'),
                    'bit_rate':    fmt.get('bit_rate'),
                    'format_name': fmt.get('format_name'),
                }
            else:
                self.metadata['warnings'].append("ffprobe not available or failed")
        except Exception as e:
            self.metadata['warnings'].append(f"Video extraction failed: {e}")

    def _extract_pe_metadata(self):
        if not PE_AVAILABLE:
            self.metadata['warnings'].append("pefile not installed")
            return
        try:
            pe      = pefile.PE(str(self.file_path))
            pe_info = {
                'machine':            hex(pe.FILE_HEADER.Machine),
                'timestamp':          datetime.fromtimestamp(
                                          pe.FILE_HEADER.TimeDateStamp
                                      ).isoformat(),
                'number_of_sections': pe.FILE_HEADER.NumberOfSections,
                'characteristics':    hex(pe.FILE_HEADER.Characteristics),
                'sections': [{
                    'name':            s.Name.decode('utf-8', errors='ignore').strip('\x00'),
                    'virtual_address': hex(s.VirtualAddress),
                    'virtual_size':    s.Misc_VirtualSize,
                    'raw_size':        s.SizeOfRawData,
                    'characteristics': hex(s.Characteristics),
                } for s in pe.sections],
            }
            if hasattr(pe, 'DIRECTORY_ENTRY_IMPORT'):
                pe_info['imports'] = [{
                    'dll':          e.dll.decode('utf-8', errors='ignore'),
                    'import_count': len(e.imports),
                } for e in pe.DIRECTORY_ENTRY_IMPORT]
            if hasattr(pe, 'VS_VERSIONINFO') and hasattr(pe, 'FileInfo'):
                for fi in pe.FileInfo:
                    if hasattr(fi, 'StringTable'):
                        for st in fi.StringTable:
                            pe_info['version_info'] = {
                                k.decode('utf-8', errors='ignore'): v.decode('utf-8', errors='ignore')
                                for k, v in st.entries.items()
                            }
            self.metadata['extracted_metadata']['pe_file'] = pe_info
        except Exception as e:
            self.metadata['errors'].append(f"PE extraction failed: {e}")


# ═══════════════════════════════════════════════════════════════════════════════
#  PARALLEL SCANNER
# ═══════════════════════════════════════════════════════════════════════════════

class ParallelScanner:
    def __init__(self, max_workers: int = DEFAULT_WORKERS,
                 case_id: Optional[str] = None,
                 logger: Optional[ForensicLogger] = None,
                 chain_of_custody: Optional[ChainOfCustody] = None,
                 examiner: str = "Unknown"):
        self.max_workers      = max_workers
        self.case_id          = case_id
        self.logger           = logger or ForensicLogger(case_id=case_id)
        self.chain_of_custody = chain_of_custody
        self.examiner         = examiner
        self.processed_count  = 0
        self.failed_count     = 0
        self.lock             = Lock()

    def scan_directory(self, directory: Path, recursive: bool = False) -> List[Dict[str, Any]]:
        pattern   = '**/*' if recursive else '*'
        file_list = [f for f in directory.glob(pattern) if f.is_file()]
        total     = len(file_list)

        self.logger.info("Starting parallel scan",
                         directory=str(directory), total_files=total, workers=self.max_workers)
        results = []

        with tqdm(total=total, desc=f"{Fore.CYAN}Scanning{Style.RESET_ALL}", unit="file") as pbar:
            with ThreadPoolExecutor(max_workers=self.max_workers) as ex:
                futures = {ex.submit(self._process_file, fp): fp for fp in file_list}
                for future in as_completed(futures):
                    try:
                        results.append(future.result())
                        with self.lock:
                            self.processed_count += 1
                    except Exception as e:
                        with self.lock:
                            self.failed_count += 1
                            self.processed_count += 1
                        self.logger.error("File processing failed",
                                          file=str(futures[future]), error=str(e))
                    finally:
                        pbar.update(1)
                        pbar.set_postfix(failed=f"{Fore.RED}{self.failed_count}{Style.RESET_ALL}")

        self.logger.info("Scan completed",
                         processed=self.processed_count, failed=self.failed_count)
        return results

    def _process_file(self, file_path: Path) -> Dict[str, Any]:
        extractor = EnhancedMetadataExtractor(
            str(file_path),
            case_id=self.case_id,
            logger=self.logger,
            chain_of_custody=self.chain_of_custody,
            examiner=self.examiner,
        )
        return extractor.extract_all()


# ═══════════════════════════════════════════════════════════════════════════════
#  OUTPUT / REPORTING
# ═══════════════════════════════════════════════════════════════════════════════

def print_metadata(metadata: Dict[str, Any], verbose: bool = False):
    print("\n" + Fore.BLUE + "=" * 80 + Style.RESET_ALL)
    print(f"{Fore.GREEN}ENHANCED METADATA EXTRACTION REPORT  v{VERSION}{Style.RESET_ALL}")
    print(Fore.BLUE + "=" * 80 + Style.RESET_ALL)

    if metadata.get('forensic_info'):
        print(f"\n{Fore.YELLOW}[FORENSIC INFORMATION]{Style.RESET_ALL}")
        fi = metadata['forensic_info']
        print(f"  Case ID:          {fi.get('case_id')}")
        print(f"  Examiner:         {fi.get('examiner')}")
        print(f"  Tool:             {fi.get('tool_name')} v{fi.get('tool_version')}")
        print(f"  Extraction (UTC): {fi.get('extraction_timestamp_utc')}")
        sys_info = fi.get('examiner_system', {})
        print(f"  System:           {sys_info.get('platform')} {sys_info.get('platform_release')}")

    if metadata.get('file_type_analysis'):
        print(f"\n{Fore.YELLOW}[FILE TYPE ANALYSIS]{Style.RESET_ALL}")
        fta = metadata['file_type_analysis']
        print(f"  Extension:     {fta.get('extension')}")
        print(f"  MIME Type:     {fta.get('mime_type')}")
        print(f"  Description:   {fta.get('description')}")
        print(f"  Classified As: {fta.get('real_type')}")

    if metadata.get('file_info'):
        print(f"\n{Fore.YELLOW}[FILE INFORMATION]{Style.RESET_ALL}")
        for k, v in metadata['file_info'].items():
            print(f"  {k.replace('_', ' ').title()}: {v}")

    if metadata.get('hashes'):
        print(f"\n{Fore.YELLOW}[FILE HASHES]{Style.RESET_ALL}")
        for k, v in metadata['hashes'].items():
            print(f"  {k.upper()}: {v}")

    if metadata.get('threat_intel'):
        ti      = metadata['threat_intel']
        color   = ti.get('risk_color', '')
        reset   = Style.RESET_ALL if COLORAMA_AVAILABLE else ''
        score   = ti.get('risk_score', 0)
        level   = ti.get('risk_level', 'UNKNOWN')
        bar_len = 40
        filled  = int(bar_len * score / 100)
        bar     = "█" * filled + "░" * (bar_len - filled)
        print(f"\n{Fore.YELLOW}[🧠 THREAT INTELLIGENCE]{Style.RESET_ALL}")
        print(f"  Risk Score : {color}{score:>3}/100  [{bar}]  {level}{reset}")
        if ti.get('iocs'):
            print(f"  {Fore.RED}IOCs:{reset}")
            for ioc in ti['iocs']:
                print(f"    ⚑  {ioc}")
        if ti.get('yara_matches'):
            print(f"  {Fore.RED}YARA Matches:{reset}")
            for m in ti['yara_matches']:
                print(f"    ⚡ [{m['severity'].upper()}] {m['rule']} – {m['description']}")

    if metadata.get('timeline', {}).get('events'):
        tl = metadata['timeline']
        print(f"\n{Fore.YELLOW}[🧾 ACTIVITY TIMELINE]{Style.RESET_ALL}")
        print(f"  Events: {tl['total_events']}  |  "
              f"Earliest: {tl['earliest']}  |  Latest: {tl['latest']}")
        for ev in tl['events'][:15]:
            print(f"    {ev['timestamp'][:19]}  [{ev['category']:<11}]  "
                  f"{ev['source']:<18}  {ev['description']}")
        if tl['total_events'] > 15:
            print(f"    … and {tl['total_events'] - 15} more (see JSON for full timeline)")

    if metadata.get('chain_of_custody'):
        coc    = metadata['chain_of_custody']
        valid  = coc.get('integrity_valid', False)
        status = (f"{Fore.GREEN}✓ INTACT{Style.RESET_ALL}"
                  if valid else f"{Fore.RED}✗ COMPROMISED{Style.RESET_ALL}")
        print(f"\n{Fore.YELLOW}[🔐 CHAIN OF CUSTODY]{Style.RESET_ALL}")
        print(f"  Record ID : {coc.get('record_id')}")
        print(f"  Events    : {coc.get('event_count')}")
        print(f"  Integrity : {status}")

    if verbose and metadata.get('exiftool_metadata'):
        print(f"\n{Fore.YELLOW}[EXIFTOOL METADATA]{Style.RESET_ALL}")
        print(json.dumps(metadata['exiftool_metadata'], indent=2, default=str))

    if metadata.get('extracted_metadata'):
        print(f"\n{Fore.YELLOW}[EXTRACTED METADATA]{Style.RESET_ALL}")
        print(json.dumps(metadata['extracted_metadata'], indent=2, default=str))

    if metadata.get('warnings'):
        print(f"\n{Fore.YELLOW}[WARNINGS]{Style.RESET_ALL}")
        for w in metadata['warnings']:
            print(f"  ⚠ {w}")

    if metadata.get('errors'):
        print(f"\n{Fore.RED}[ERRORS]{Style.RESET_ALL}")
        for e in metadata['errors']:
            print(f"  ❌ {e}")

    print("\n" + "=" * 80 + "\n")


def save_to_json(metadata: Dict[str, Any], output_path: str):
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(metadata, f, indent=2, default=str)
    print(f"✓ Metadata saved to: {output_path}")


def save_to_csv(data: Any, output_path: str):
    try:
        items = data if isinstance(data, list) else data.get('files', [data])
        rows  = []
        for item in items:
            flat = {}
            for section, key_prefix in [('file_info', 'file'), ('hashes', 'hash')]:
                for k, v in item.get(section, {}).items():
                    flat[f"{key_prefix}_{k}"] = v
            for k, v in item.get('forensic_info', {}).items():
                if isinstance(v, (str, int, float)):
                    flat[f"forensic_{k}"] = v
            ti = item.get('threat_intel', {})
            flat['threat_risk_score'] = ti.get('risk_score', '')
            flat['threat_risk_level'] = ti.get('risk_level', '')
            flat['threat_ioc_count']  = len(ti.get('iocs', []))
            flat['yara_match_count']  = len(ti.get('yara_matches', []))
            tl = item.get('timeline', {})
            flat['timeline_earliest'] = tl.get('earliest', '')
            flat['timeline_latest']   = tl.get('latest', '')
            flat['timeline_events']   = tl.get('total_events', '')
            rows.append(flat)

        if not rows:
            print("⚠ No data to save to CSV")
            return

        headers = set()
        for r in rows:
            headers.update(r.keys())
        with open(output_path, 'w', newline='', encoding='utf-8') as f:
            w = csv.DictWriter(f, fieldnames=sorted(headers))
            w.writeheader()
            w.writerows(rows)
        print(f"✓ CSV Report saved to: {output_path}")
    except Exception as e:
        print(f"❌ Failed to save CSV: {e}")


def save_to_html(data: Any, output_path: str):
    try:
        timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        files     = data if isinstance(data, list) else data.get('files', [data])

        html = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<title>Forensic Metadata Report – {TOOL_NAME} v{VERSION}</title>
<style>
  :root {{
    --bg: #0d1117; --surface: #161b22; --border: #30363d;
    --text: #c9d1d9; --muted: #8b949e; --accent: #58a6ff;
    --danger: #f85149; --warn: #e3b341; --success: #3fb950;
  }}
  * {{ box-sizing: border-box; margin: 0; padding: 0; }}
  body {{ background: var(--bg); color: var(--text);
          font-family: 'Segoe UI', system-ui, sans-serif; padding: 24px; }}
  header {{ background: linear-gradient(135deg,#1f2d3d,#0a1628);
            border: 1px solid var(--border); border-radius: 12px;
            padding: 28px 32px; margin-bottom: 28px; }}
  header h1 {{ font-size: 1.6rem; color: var(--accent); }}
  header p  {{ color: var(--muted); margin-top: 6px; font-size: .9rem; }}
  .card {{ background: var(--surface); border: 1px solid var(--border);
           border-radius: 10px; padding: 22px; margin-bottom: 22px; }}
  .card h2 {{ font-size: 1.1rem; color: var(--accent); margin-bottom: 14px;
              border-bottom: 1px solid var(--border); padding-bottom: 8px; }}
  .grid2 {{ display: grid; grid-template-columns: 1fr 1fr; gap: 16px; }}
  .section h3 {{ font-size: .75rem; text-transform: uppercase; letter-spacing: .08em;
                 color: var(--muted); margin-bottom: 8px; }}
  table {{ width: 100%; border-collapse: collapse; font-size: .88rem; }}
  td {{ padding: 5px 8px; border-bottom: 1px solid var(--border); color: var(--text); }}
  td:first-child {{ color: var(--muted); white-space: nowrap; width: 40%; }}
  .badge {{ display: inline-block; padding: 2px 8px; border-radius: 20px;
            font-size: .78rem; font-weight: 600; }}
  .badge-danger  {{ background: #3d1c1c; color: var(--danger); }}
  .badge-warn    {{ background: #2d2416; color: var(--warn); }}
  .badge-success {{ background: #162a1e; color: var(--success); }}
  .badge-info    {{ background: #1a2740; color: var(--accent); }}
  .risk-bar {{ height: 8px; background: var(--border); border-radius: 4px;
               margin: 6px 0; overflow: hidden; }}
  .risk-fill {{ height: 100%; border-radius: 4px; transition: width .4s; }}
  .risk-critical {{ background: var(--danger); }}
  .risk-high     {{ background: #e05252; }}
  .risk-medium   {{ background: var(--warn); }}
  .risk-low      {{ background: var(--accent); }}
  .risk-clean    {{ background: var(--success); }}
  pre {{ background: #0d1117; border: 1px solid var(--border); border-radius: 6px;
         padding: 12px; font-size: .8rem; overflow-x: auto; color: var(--text); }}
  .tl-event {{ display: flex; gap: 12px; padding: 5px 0;
               border-bottom: 1px solid var(--border); font-size: .85rem; }}
  .tl-ts    {{ color: var(--muted); white-space: nowrap; min-width: 140px; }}
  .tl-cat   {{ min-width: 90px; }}
  .ioc-item {{ padding: 4px 0; font-size: .85rem; color: var(--danger); }}
  .coc-item {{ padding: 5px 0; border-bottom: 1px solid var(--border); font-size: .82rem; }}
  .coc-act  {{ color: var(--accent); font-weight: 600; min-width: 180px; display: inline-block; }}
  .hash-val {{ font-family: monospace; font-size: .78rem; word-break: break-all;
               color: var(--success); }}
</style>
</head>
<body>
<header>
  <h1>🕵️ {TOOL_NAME} v{VERSION} – Forensic Report</h1>
  <p>Generated: {timestamp} &nbsp;|&nbsp; Files analysed: {len(files)}</p>
</header>
"""

        for f in files:
            fname    = f.get('file_info', {}).get('filename', 'Unknown')
            fpath    = f.get('file_info', {}).get('full_path', '')
            fhashes  = f.get('hashes', {})
            ti       = f.get('threat_intel', {})
            tl       = f.get('timeline', {})
            coc_data = f.get('chain_of_custody', {})
            fi       = f.get('file_info', {})
            finfo    = f.get('forensic_info', {})   # FIX: no longer inside f-string

            score      = ti.get('risk_score', 0)
            level      = ti.get('risk_level', 'CLEAN').lower()
            risk_class = f"risk-{level}"
            level_badge_class = {
                'critical': 'badge-danger',
                'high':     'badge-danger',
                'medium':   'badge-warn',
                'low':      'badge-info',
                'clean':    'badge-success',
            }.get(level, 'badge-info')

            entropy_val = finfo.get('entropy', 'N/A')   # FIX: pre-extracted

            html += f"""
<div class="card">
  <h2>📄 {fname}</h2>
  <div style="color:var(--muted);font-size:.82rem;margin-bottom:14px;">{fpath}</div>
  <div class="grid2">
    <div class="section">
      <h3>File Information</h3>
      <table>
        <tr><td>Size</td><td>{fi.get('file_size_human','')}</td></tr>
        <tr><td>Created</td><td>{fi.get('created_time','')}</td></tr>
        <tr><td>Modified</td><td>{fi.get('modified_time','')}</td></tr>
        <tr><td>Accessed</td><td>{fi.get('accessed_time','')}</td></tr>
        <tr><td>Permissions</td><td>{fi.get('permissions','')}</td></tr>
        <tr><td>Entropy</td><td>{entropy_val}</td></tr>
      </table>
    </div>
    <div class="section">
      <h3>🧠 Threat Score</h3>
      <div style="font-size:2rem;font-weight:700;color:var(--text)">{score}<span style="font-size:1rem;color:var(--muted)">/100</span></div>
      <div class="risk-bar"><div class="risk-fill {risk_class}" style="width:{score}%"></div></div>
      <span class="badge {level_badge_class}">{level.upper()}</span>
      <div style="margin-top:10px">
"""
            for ioc in ti.get('iocs', []):
                html += f'        <div class="ioc-item">⚑ {ioc}</div>\n'

            for m in ti.get('yara_matches', []):
                sev = m.get('severity', '')
                bc  = {'critical': 'badge-danger', 'high': 'badge-danger',
                        'medium': 'badge-warn', 'low': 'badge-info'}.get(sev, 'badge-info')
                html += (f'        <div style="margin-top:4px;font-size:.83rem;">'
                         f'⚡ <span class="badge {bc}">{sev}</span> '
                         f'{m["rule"]} – {m["description"]}</div>\n')

            html += "      </div>\n    </div>\n  </div>\n"

            html += """
  <div class="section" style="margin-top:16px">
    <h3>File Hashes</h3>
    <table>
"""
            for algo, digest in fhashes.items():
                html += f'      <tr><td>{algo.upper()}</td><td><span class="hash-val">{digest}</span></td></tr>\n'
            html += "    </table>\n  </div>\n"

            if tl.get('events'):
                html += """
  <div class="section" style="margin-top:16px">
    <h3>🧾 Activity Timeline</h3>
"""
                for ev in tl['events'][:20]:
                    cat_color = {
                        'filesystem': '#3fb950', 'exif': '#58a6ff',
                        'document':   '#a371f7', 'executable': '#f85149',
                    }.get(ev['category'], '#8b949e')
                    html += (f'    <div class="tl-event">'
                             f'<span class="tl-ts">{ev["timestamp"][:19]}</span>'
                             f'<span class="tl-cat" style="color:{cat_color}">[{ev["category"]}]</span>'
                             f'<span>{ev["source"]}</span>'
                             f'<span style="color:var(--muted)">{ev["description"]}</span>'
                             f'</div>\n')
                if tl['total_events'] > 20:
                    html += (f'    <div style="color:var(--muted);font-size:.82rem;padding:6px 0">'
                             f'… {tl["total_events"] - 20} more events in JSON export</div>\n')
                html += "  </div>\n"

            if coc_data.get('events'):
                valid   = coc_data.get('integrity_valid', False)
                v_style = 'color:var(--success)' if valid else 'color:var(--danger)'
                v_text  = '✓ INTACT' if valid else '✗ COMPROMISED'
                html += f"""
  <div class="section" style="margin-top:16px">
    <h3>🔐 Chain of Custody</h3>
    <div style="margin-bottom:8px;font-size:.85rem">
      Integrity: <strong style="{v_style}">{v_text}</strong>
      &nbsp;|&nbsp; Record: <span style="font-family:monospace;font-size:.78rem">{coc_data.get('record_id','')}</span>
    </div>
"""
                for ev in coc_data['events']:
                    details_str = ", ".join(
                        f"{k}={v}" for k, v in ev["details"].items()
                        if isinstance(v, (str, int, bool)) and k not in ("note", "hashes")
                    )[:80]
                    html += (
                        f'    <div class="coc-item">'
                        f'<span style="color:var(--muted);min-width:160px;display:inline-block">'
                        f'{ev["timestamp"][:19]}</span>'
                        f'<span class="coc-act">{ev["action"]}</span>'
                        f'<span style="color:var(--muted);font-size:.8rem">{details_str}</span>'
                        f'</div>\n'
                    )
                html += "  </div>\n"

            html += "</div>\n"

        html += f"""
<footer style="text-align:center;color:var(--muted);margin-top:40px;font-size:.82rem;padding:20px">
  Generated by {TOOL_NAME} v{VERSION}
</footer>
</body>
</html>"""

        with open(output_path, 'w', encoding='utf-8') as fh:
            fh.write(html)
        print(f"✓ HTML Report saved to: {output_path}")
    except Exception as e:
        print(f"❌ Failed to save HTML: {e}")


def save_output(data, file_format: str, output_file: str):
    if file_format == 'json':
        save_to_json(data, output_file)
    elif file_format == 'html':
        save_to_html(data, output_file)
    elif file_format == 'csv':
        save_to_csv(data, output_file)


# ═══════════════════════════════════════════════════════════════════════════════
#  INTERACTIVE / CLI
# ═══════════════════════════════════════════════════════════════════════════════

def print_banner():
    print(f"""{Fore.CYAN}
╔══════════════════════════════════════════════════════════════════════════════╗
║                                                                              ║
║  {Fore.YELLOW}FORENSISCAN  v{VERSION}{Fore.CYAN}                                               ║
║                                                                              ║
║  Images · Docs · Audio · Video · PE · Archives · Disk Images                ║
║  🧠 Threat Intel  🧬 YARA  🧾 Timeline  🔐 Chain of Custody               ║
║  Author: Sayan Pal  |  Collaborator: Soumit Santra                          ║
╚══════════════════════════════════════════════════════════════════════════════╝
{Style.RESET_ALL}""")


def ask_examiner() -> str:
    return input("Enter examiner name (press Enter to skip): ").strip() or "Unknown"


def ask_save_options():
    print("\n" + "=" * 80)
    print("[SAVE OPTIONS]")
    if input("\nSave the output? (y/n): ").strip().lower() != 'y':
        return None, None

    print("\n1. JSON  2. HTML  3. CSV")
    choice = input("Format (1-3): ").strip()
    fmt, ext = (
        ('html', '.html') if choice == '2' else
        ('csv',  '.csv')  if choice == '3' else
        ('json', '.json')
    )

    out = input("Output filename (Enter for auto): ").strip()
    if not out:
        ts  = datetime.now().strftime('%Y%m%d_%H%M%S')
        out = f"forensic_report_{ts}{ext}"
    elif not out.endswith(ext):
        out += ext
    return fmt, out


def interactive_mode():
    print_banner()
    print("\n[MAIN MENU]")
    print("=" * 80)
    print("1. Extract metadata from a single file")
    print("2. Process all files in a directory (non-recursive)")
    print("3. Process directory recursively")
    print("4. Forensic mode – Single file with case tracking + CoC + Intelligence")
    print("5. Forensic mode – Directory scan with case tracking")
    print("6. Analyse forensic disk image (E01, DD, AFF4)")
    print("7. Show system capabilities")
    print("8. Show supported file types")
    print("9. About this tool")
    print("0. Exit")
    print("=" * 80)

    choice = input("\nEnter your choice (0-9): ").strip()

    if choice == '0':
        print("\n✓ Thank you for using ForensiScan!")
        sys.exit(0)

    elif choice == '1':
        file_path = input("\nFile path: ").strip().strip('"').strip("'")
        if not file_path:
            print("❌ No path provided")
            return
        extractor = EnhancedMetadataExtractor(file_path)
        metadata  = extractor.extract_all()
        print_metadata(metadata, verbose=True)
        ThreatIntelligence(file_path).print_report(metadata['threat_intel'])
        extractor.timeline.print_timeline()
        fmt, out = ask_save_options()
        if out:
            save_output(metadata, fmt, out)

    elif choice in ('2', '3'):
        recursive = (choice == '3')
        dir_path  = input("\nDirectory path: ").strip().strip('"').strip("'")
        if not dir_path or not Path(dir_path).is_dir():
            print("❌ Invalid directory")
            return
        workers_in = input(f"Parallel workers (default {DEFAULT_WORKERS}): ").strip()
        workers    = int(workers_in) if workers_in.isdigit() else DEFAULT_WORKERS

        scanner = ParallelScanner(max_workers=workers)
        results = scanner.scan_directory(Path(dir_path), recursive)
        print(f"\n✓ Processed {len(results)} files  |  Failed: {scanner.failed_count}")
        fmt, out = ask_save_options()
        if out:
            save_output({'scan_info': {'directory': dir_path,
                                       'total_files': len(results),
                                       'recursive': recursive},
                         'files': results}, fmt, out)

    elif choice == '4':
        file_path = input("\nFile path: ").strip().strip('"').strip("'")
        if not file_path:
            print("❌ No path provided")
            return
        case_id  = input("Case ID (Enter=auto): ").strip() or \
                   f"CASE-{datetime.now().strftime('%Y%m%d-%H%M%S')}"
        examiner = ask_examiner()
        log_file = input("Log file path (Enter=skip): ").strip() or None
        coc_file = input("Chain of Custody output file (Enter=skip): ").strip() or None

        logger = ForensicLogger(log_file=log_file, case_id=case_id)
        coc    = ChainOfCustody(case_id=case_id, examiner=examiner, output_path=coc_file)

        extractor = EnhancedMetadataExtractor(
            file_path, case_id=case_id, logger=logger,
            chain_of_custody=coc, examiner=examiner
        )
        metadata = extractor.extract_all()
        coc.close()

        print_metadata(metadata, verbose=True)
        ThreatIntelligence(file_path).print_report(metadata['threat_intel'])
        extractor.timeline.print_timeline()
        coc.print_summary()

        fmt, out = ask_save_options()
        if out:
            coc.log_export(out, fmt)
            save_output(metadata, fmt, out)

    elif choice == '5':
        dir_path = input("\nDirectory path: ").strip().strip('"').strip("'")
        if not dir_path or not Path(dir_path).is_dir():
            print("❌ Invalid directory")
            return
        case_id      = input("Case ID (Enter=auto): ").strip() or \
                       f"CASE-{datetime.now().strftime('%Y%m%d-%H%M%S')}"
        examiner     = ask_examiner()
        log_file     = input("Log file (Enter=skip): ").strip() or None
        coc_file     = input("CoC output file (Enter=skip): ").strip() or None
        recursive_in = input("Recursive? (y/n, default y): ").strip().lower()
        recursive    = recursive_in != 'n'
        workers_in   = input(f"Workers (default {DEFAULT_WORKERS}): ").strip()
        workers      = int(workers_in) if workers_in.isdigit() else DEFAULT_WORKERS

        logger  = ForensicLogger(log_file=log_file, case_id=case_id)
        coc     = ChainOfCustody(case_id=case_id, examiner=examiner, output_path=coc_file)
        scanner = ParallelScanner(max_workers=workers, case_id=case_id,
                                  logger=logger, chain_of_custody=coc, examiner=examiner)
        results = scanner.scan_directory(Path(dir_path), recursive)
        coc.close()

        print(f"\n✓ Processed {len(results)} files")
        coc.print_summary()

        fmt, out = ask_save_options()
        if out:
            coc.log_export(out, fmt)
            save_output({'scan_info': {'case_id': case_id, 'directory': dir_path,
                                       'total_files': len(results), 'recursive': recursive,
                                       'workers': workers,
                                       'scan_timestamp': datetime.now(timezone.utc).isoformat()},
                         'files': results}, fmt, out)

    elif choice == '6':
        if not TSK_AVAILABLE:
            print(f"\n{Fore.RED}❌ pytsk3 not installed. pip install pytsk3 pyewf{Style.RESET_ALL}")
            return
        file_path = input("\nDisk image path (E01/DD/AFF4/IMG): ").strip().strip('"').strip("'")
        if not file_path or not Path(file_path).exists():
            print("❌ Invalid path")
            return
        handler = DiskImageHandler(file_path)
        if not handler.open_image():
            print("❌ Failed to open image")
            return
        size_h = EnhancedMetadataExtractor._human_readable_size(handler.img_info.get_size())
        print(f"\n✓ Opened: {size_h}")
        parts = handler.list_partitions()
        for p in parts:
            print(f"  Partition {p['id']}: {p['description']} | {p['length']} sectors")
        pid   = input("\nPartition ID to walk (default 0): ").strip()
        pid   = int(pid) if pid.isdigit() else 0
        flist = handler.walk_filesystem(pid)
        print(f"\n✓ Found {len(flist)} items")
        for ff in flist[:20]:
            print(f"  {ff['type'][0].upper()} | {ff['size']:>10} B | {ff['path']}")
        if input(f"\nSave full list ({len(flist)} items)? (y/n): ").strip().lower() == 'y':
            fmt, out = ask_save_options()
            if out:
                save_output({'image': file_path, 'partition': pid, 'files': flist}, fmt, out)

    elif choice == '7':
        print("\n[SYSTEM CAPABILITIES]")
        et = ExifToolWrapper()
        if et.available:
            r = subprocess.run(['exiftool', '-ver'], capture_output=True, text=True, timeout=5)
            print(f"✓ ExifTool v{r.stdout.strip()}")
        else:
            print("❌ ExifTool not installed – https://exiftool.org/")
        print(f"{'✓' if MAGIC_AVAILABLE   else '❌'} libmagic")
        print(f"{'✓' if YARA_AVAILABLE    else '❌'} yara-python  – pip install yara-python")
        print(f"{'✓' if TSK_AVAILABLE     else '❌'} pytsk3")
        print(f"{'✓' if EWF_AVAILABLE     else '❌'} pyewf")
        print(f"{'✓' if PIL_AVAILABLE     else '❌'} Pillow")
        print(f"{'✓' if HACHOIR_AVAILABLE else '❌'} Hachoir")
        print(f"{'✓' if PDF_AVAILABLE     else '❌'} PyPDF2")
        print(f"{'✓' if DOCX_AVAILABLE    else '❌'} python-docx")
        print(f"{'✓' if AUDIO_AVAILABLE   else '❌'} mutagen")
        print(f"{'✓' if PE_AVAILABLE      else '❌'} pefile")
        try:
            subprocess.run(['ffprobe', '-version'], capture_output=True, timeout=5)
            print("✓ ffprobe")
        except Exception:
            print("❌ ffprobe")
        input("\nPress Enter to return to main menu...")
        interactive_mode()
        return

    elif choice == '8':
        print("""
[SUPPORTED FILE TYPES]
📷 Images      : .jpg .png .gif .bmp .tiff .webp
📄 PDF         : .pdf
📝 Documents   : .docx .doc
🎵 Audio       : .mp3 .m4a .flac .ogg .wav
🎬 Video       : .mp4 .avi .mkv .mov (requires ffprobe)
📦 Archives    : .zip .jar .apk
⚙️  Executables : .exe .dll (PE files)
💽 Disk Images : .e01 .dd .img .iso .aff4

All files → MD5/SHA1/SHA256 hashes, entropy, MIME type, timestamps,
             🧠 Threat score, 🧬 YARA scan, 🧾 Timeline, 🔐 Chain of Custody
""")
        input("Press Enter to return to main menu...")
        interactive_mode()
        return

    elif choice == '9':
        print(f"""
[ABOUT  {TOOL_NAME}  v{VERSION}]
Author: Sayan Pal  |  Collaborator: Soumit Santra

🆕 v4.0.1 fixes:
  🐛 TypeError crash in GUI ({{}} f-string escaping fixed throughout)
  🐛 DeprecationWarning: datetime.utcnow() replaced with timezone-aware calls
  🐛 Duplicate logging handlers eliminated
  🐛 Pillow ≥10 getexif() compatibility
  🐛 PE timestamp comparison uses naive datetime consistently

⚠️  ETHICAL USE ONLY. Analyse only files you are authorised to examine.
""")
        input("Press Enter to return to main menu...")
        interactive_mode()
        return

    else:
        print("❌ Invalid choice")
        return

    if input("\nPerform another operation? (y/n): ").strip().lower() == 'y':
        interactive_mode()
    else:
        print("\n✓ Thank you for using ForensiScan!")


# ═══════════════════════════════════════════════════════════════════════════════
#  CLI ENTRY-POINT
# ═══════════════════════════════════════════════════════════════════════════════

def main():
    if len(sys.argv) == 1:
        interactive_mode()
        return

    parser = argparse.ArgumentParser(
        description=f'{TOOL_NAME} v{VERSION}',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python metadata_scanner.py                           # interactive mode
  python metadata_scanner.py image.jpg -v              # single file, verbose
  python metadata_scanner.py -d /evidence -r --workers 8
  python metadata_scanner.py evidence.pdf --case-id CASE-2024-001 \\
         --examiner "Jane Doe" --coc chain.json --log forensic.log
  python metadata_scanner.py samples/ -o report.html --format html
"""
    )
    parser.add_argument('input',          nargs='?', help='Input file or directory')
    parser.add_argument('-o', '--output', help='Output file path')
    parser.add_argument('-f', '--format', choices=['json', 'html', 'csv'])
    parser.add_argument('-d', '--directory', action='store_true')
    parser.add_argument('-r', '--recursive', action='store_true')
    parser.add_argument('-v', '--verbose',   action='store_true')
    parser.add_argument('--workers',  type=int, default=DEFAULT_WORKERS)
    parser.add_argument('--case-id',  help='Forensic case identifier')
    parser.add_argument('--examiner', default='Unknown', help='Examiner name')
    parser.add_argument('--log',      help='Structured log file')
    parser.add_argument('--coc',      help='Chain of Custody output JSON file')

    args    = parser.parse_args()
    print_banner()

    if not args.input:
        print("❌ No input specified")
        parser.print_help()
        sys.exit(1)

    case_id = args.case_id or f"CASE-{datetime.now().strftime('%Y%m%d-%H%M%S')}"
    logger  = ForensicLogger(log_file=args.log, case_id=case_id)
    coc     = ChainOfCustody(case_id=case_id, examiner=args.examiner,
                              output_path=args.coc)

    input_path = Path(args.input)

    if args.directory or input_path.is_dir():
        if not input_path.is_dir():
            print(f"❌ {input_path} is not a directory")
            sys.exit(1)
        scanner = ParallelScanner(max_workers=args.workers, case_id=case_id,
                                  logger=logger, chain_of_custody=coc,
                                  examiner=args.examiner)
        results = scanner.scan_directory(input_path, args.recursive)
        coc.close()
        print(f"\n✓ Processed {len(results)} files")
        coc.print_summary()
        if args.output:
            output_data = {
                'scan_info': {
                    'directory':   str(input_path.absolute()),
                    'recursive':   args.recursive,
                    'total_files': len(results),
                    'case_id':     case_id,
                    'scan_timestamp': datetime.now(timezone.utc).isoformat(),
                },
                'files': results
            }
            fmt = args.format or (
                'html' if args.output.endswith('.html') else
                'csv'  if args.output.endswith('.csv')  else 'json'
            )
            save_output(output_data, fmt, args.output)
            coc.log_export(args.output, fmt)
    else:
        try:
            extractor = EnhancedMetadataExtractor(
                str(input_path), case_id=case_id, logger=logger,
                chain_of_custody=coc, examiner=args.examiner
            )
            metadata = extractor.extract_all()
            coc.close()

            print_metadata(metadata, verbose=args.verbose)

            ti_obj = ThreatIntelligence(str(input_path))
            ti_obj.print_report(metadata['threat_intel'])
            extractor.timeline.print_timeline()
            coc.print_summary()

            if args.output:
                fmt = args.format or (
                    'html' if args.output.endswith('.html') else
                    'csv'  if args.output.endswith('.csv')  else 'json'
                )
                save_output(metadata, fmt, args.output)
                coc.log_export(args.output, fmt)

        except FileNotFoundError as e:
            print(f"❌ {e}")
            sys.exit(1)
        except Exception as e:
            print(f"❌ {e}")
            if args.verbose:
                import traceback
                traceback.print_exc()
            sys.exit(1)


# ═══════════════════════════════════════════════════════════════════════════════
#  GUI  —  Full tkinter dashboard
# ═══════════════════════════════════════════════════════════════════════════════

try:
    import tkinter as tk
    from tkinter import ttk, filedialog, messagebox, scrolledtext
    TK_AVAILABLE = True
except ImportError:
    TK_AVAILABLE = False


# ── Palette ───────────────────────────────────────────────────────────────────
_C = {
    "bg":       "#0b0f14",
    "surface":  "#111820",
    "panel":    "#161e28",
    "border":   "#1e2d3d",
    "accent":   "#00d4aa",
    "accent2":  "#0096ff",
    "danger":   "#ff4757",
    "warn":     "#ffa502",
    "success":  "#2ed573",
    "muted":    "#4a6278",
    "text":     "#c8d8e8",
    "text_dim": "#6b8299",
    "grid":     "#1a2535",
    "sel":      "#1e3a4a",
}

_FONT_MONO  = ("Courier New", 9)
_FONT_SMALL = ("Courier New", 8)
_FONT_HEAD  = ("Courier New", 11, "bold")
_FONT_TITLE = ("Courier New", 14, "bold")
_FONT_HUGE  = ("Courier New", 28, "bold")


def _hex_to_rgb(h):
    h = h.lstrip("#")
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))


def _lerp_color(c1, c2, t):
    r1, g1, b1 = _hex_to_rgb(c1)
    r2, g2, b2 = _hex_to_rgb(c2)
    r = int(r1 + (r2 - r1) * t)
    g = int(g1 + (g2 - g1) * t)
    b = int(b1 + (b2 - b1) * t)
    return f"#{r:02x}{g:02x}{b:02x}"


class _StyledScrolledText(tk.Frame):
    """Dark-themed scrollable text widget."""
    def __init__(self, parent, **kw):
        super().__init__(parent, bg=_C["panel"], bd=0)
        self.text = tk.Text(
            self, bg=_C["panel"], fg=_C["text"],
            insertbackground=_C["accent"],
            font=_FONT_MONO, relief="flat", bd=0, wrap="none",
            selectbackground=_C["sel"], selectforeground=_C["text"], **kw
        )
        vsb = tk.Scrollbar(self, orient="vertical",   command=self.text.yview,
                           bg=_C["border"], troughcolor=_C["bg"],
                           activebackground=_C["accent"], relief="flat", width=10)
        hsb = tk.Scrollbar(self, orient="horizontal", command=self.text.xview,
                           bg=_C["border"], troughcolor=_C["bg"],
                           activebackground=_C["accent"], relief="flat", width=10)
        self.text.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)
        self.text.grid(row=0, column=0, sticky="nsew")
        vsb.grid(row=0, column=1, sticky="ns")
        hsb.grid(row=1, column=0, sticky="ew")
        self.rowconfigure(0, weight=1)
        self.columnconfigure(0, weight=1)

    def set_text(self, content: str):
        self.text.configure(state="normal")
        self.text.delete("1.0", "end")
        self.text.insert("end", content)
        self.text.configure(state="disabled")

    def append(self, content: str, tag: str = ""):
        self.text.configure(state="normal")
        self.text.insert("end", content, tag)
        self.text.see("end")
        self.text.configure(state="disabled")

    def clear(self):
        self.text.configure(state="normal")
        self.text.delete("1.0", "end")
        self.text.configure(state="disabled")

    def add_tag(self, name, **kw):
        self.text.tag_configure(name, **kw)


class _CanvasGraph(tk.Canvas):
    """Reusable canvas for custom charts."""
    def __init__(self, parent, **kw):
        super().__init__(parent, bg=_C["panel"], bd=0,
                         highlightthickness=0, **kw)


# ── Tab: Dashboard ─────────────────────────────────────────────────────────────

class DashboardTab(tk.Frame):
    def __init__(self, parent, app):
        super().__init__(parent, bg=_C["bg"])
        self.app = app
        self._build()

    def _build(self):
        cards_row = tk.Frame(self, bg=_C["bg"])
        cards_row.pack(fill="x", padx=16, pady=(16, 8))

        self._card_risk = self._make_card(cards_row, "RISK SCORE",     "–", _C["muted"])
        self._card_iocs = self._make_card(cards_row, "IOCs FOUND",     "–", _C["muted"])
        self._card_yara = self._make_card(cards_row, "YARA MATCHES",   "–", _C["muted"])
        self._card_ts   = self._make_card(cards_row, "TIMELINE EVENTS","–", _C["muted"])
        self._card_hash = self._make_card(cards_row, "HASH STATUS",    "–", _C["muted"])
        for c in (self._card_risk, self._card_iocs, self._card_yara,
                  self._card_ts, self._card_hash):
            c["frame"].pack(side="left", fill="both", expand=True, padx=6)

        mid = tk.Frame(self, bg=_C["bg"])
        mid.pack(fill="both", expand=True, padx=16, pady=8)

        gauge_frame = tk.Frame(mid, bg=_C["panel"],
                               highlightthickness=1, highlightbackground=_C["border"])
        gauge_frame.pack(side="left", fill="both", expand=True, padx=(0, 8))
        tk.Label(gauge_frame, text="THREAT GAUGE", font=_FONT_HEAD,
                 bg=_C["panel"], fg=_C["accent"]).pack(pady=(12, 4))
        self._gauge_canvas = _CanvasGraph(gauge_frame, width=260, height=200)
        self._gauge_canvas.pack(padx=20, pady=(4, 16))

        info_frame = tk.Frame(mid, bg=_C["panel"],
                              highlightthickness=1, highlightbackground=_C["border"])
        info_frame.pack(side="left", fill="both", expand=True)
        tk.Label(info_frame, text="FILE INFORMATION", font=_FONT_HEAD,
                 bg=_C["panel"], fg=_C["accent"]).pack(pady=(12, 4), anchor="w", padx=14)
        self._info_text = _StyledScrolledText(info_frame, height=12)
        self._info_text.pack(fill="both", expand=True, padx=6, pady=(0, 8))

        ioc_frame = tk.Frame(self, bg=_C["panel"],
                             highlightthickness=1, highlightbackground=_C["border"])
        ioc_frame.pack(fill="both", expand=True, padx=16, pady=(0, 16))
        tk.Label(ioc_frame, text="INDICATORS OF COMPROMISE", font=_FONT_HEAD,
                 bg=_C["panel"], fg=_C["danger"]).pack(pady=(10, 4), anchor="w", padx=14)
        self._ioc_text = _StyledScrolledText(ioc_frame, height=8)
        self._ioc_text.add_tag("ioc",    foreground=_C["danger"])
        self._ioc_text.add_tag("ok",     foreground=_C["success"])
        self._ioc_text.add_tag("header", foreground=_C["accent"], font=_FONT_HEAD)
        self._ioc_text.pack(fill="both", expand=True, padx=6, pady=(0, 8))

    def _make_card(self, parent, label, value, color):
        f = tk.Frame(parent, bg=_C["panel"],
                     highlightthickness=1, highlightbackground=_C["border"])
        tk.Label(f, text=label, font=_FONT_SMALL,
                 bg=_C["panel"], fg=_C["muted"]).pack(pady=(10, 2))
        val_lbl = tk.Label(f, text=value, font=_FONT_HUGE,
                           bg=_C["panel"], fg=color)
        val_lbl.pack(pady=(0, 10))
        return {"frame": f, "label": val_lbl}

    def _update_card(self, card, value, color):
        card["label"].configure(text=str(value), fg=color)

    def _draw_gauge(self, score: int):
        c = self._gauge_canvas
        c.delete("all")
        W, H   = 260, 200
        cx, cy = W // 2, H - 30
        r      = 90

        c.create_arc(cx - r, cy - r, cx + r, cy + r,
                     start=0, extent=180, style="arc",
                     outline=_C["grid"], width=18)

        t      = score / 100
        color  = _lerp_color(_C["success"], _C["danger"], t)
        extent = 180 * t
        if extent > 0:
            c.create_arc(cx - r, cy - r, cx + r, cy + r,
                         start=180 - extent, extent=extent, style="arc",
                         outline=color, width=18)

        c.create_text(cx, cy - 20, text=str(score),
                      font=("Courier New", 34, "bold"), fill=color)
        c.create_text(cx, cy + 4, text="/ 100",
                      font=("Courier New", 11), fill=_C["muted"])

        level = (["CLEAN", "LOW", "MEDIUM", "HIGH", "CRITICAL"]
                 [min(4, score // 20)] if score < 100 else "CRITICAL")
        c.create_text(cx, cy + 26, text=level, font=_FONT_HEAD, fill=color)

    def populate(self, metadata: dict):
        ti    = metadata.get("threat_intel", {})
        fi    = metadata.get("file_info", {})
        tl    = metadata.get("timeline", {})
        h     = metadata.get("hashes", {})
        # ── FIX: pre-extract nested dicts so they are never inside an f-string ──
        finfo = metadata.get("forensic_info", {})
        ftype = metadata.get("file_type_analysis", {})

        score = ti.get("risk_score", 0)
        level = ti.get("risk_level", "CLEAN")
        iocs  = ti.get("iocs", [])
        yara  = ti.get("yara_matches", [])

        risk_color = {
            "CRITICAL": _C["danger"],
            "HIGH":     _C["danger"],
            "MEDIUM":   _C["warn"],
            "LOW":      _C["accent2"],
            "CLEAN":    _C["success"],
        }.get(level, _C["text"])

        self._update_card(self._card_risk, score,
                          risk_color)
        self._update_card(self._card_iocs, len(iocs),
                          _C["danger"] if iocs else _C["success"])
        self._update_card(self._card_yara, len(yara),
                          _C["warn"] if yara else _C["success"])
        self._update_card(self._card_ts, tl.get("total_events", 0), _C["accent2"])
        self._update_card(self._card_hash,
                          "OK" if h.get("md5") else "–",
                          _C["success"] if h.get("md5") else _C["muted"])

        self._draw_gauge(score)

        # ── FIX: use pre-extracted dicts — no nested .get() inside f-strings ──
        lines = [
            f"  Filename   : {fi.get('filename', '')}",
            f"  Path       : {fi.get('full_path', '')}",
            f"  Size       : {fi.get('file_size_human', '')}",
            f"  Created    : {fi.get('created_time', '')}",
            f"  Modified   : {fi.get('modified_time', '')}",
            f"  Accessed   : {fi.get('accessed_time', '')}",
            f"  Permissions: {fi.get('permissions', '')}",
            "",
            f"  MD5        : {h.get('md5', '')}",
            f"  SHA1       : {h.get('sha1', '')}",
            f"  SHA256     : {h.get('sha256', '')}",
            "",
            f"  Entropy    : {finfo.get('entropy', 'N/A')}",
            f"  MIME Type  : {ftype.get('mime_type', '')}",
        ]
        self._info_text.set_text("\n".join(lines))

        self._ioc_text.clear()
        if iocs:
            self._ioc_text.append("  INDICATORS OF COMPROMISE\n", "header")
            for ioc in iocs:
                self._ioc_text.append(f"  ⚑  {ioc}\n", "ioc")
        else:
            self._ioc_text.append("  ✓  No IOCs detected — file appears clean.\n", "ok")

        if yara:
            self._ioc_text.append("\n  YARA RULE MATCHES\n", "header")
            for m in yara:
                self._ioc_text.append(
                    f"  ⚡ [{m['severity'].upper()}] {m['rule']} — {m['description']}\n",
                    "ioc"
                )


# ── Tab: Timeline ──────────────────────────────────────────────────────────────

class TimelineTab(tk.Frame):
    def __init__(self, parent, app):
        super().__init__(parent, bg=_C["bg"])
        self.app = app
        self._events: list = []
        self._build()

    def _build(self):
        top = tk.Frame(self, bg=_C["bg"])
        top.pack(fill="x", padx=16, pady=(14, 6))
        tk.Label(top, text="FILE ACTIVITY TIMELINE", font=_FONT_TITLE,
                 bg=_C["bg"], fg=_C["accent"]).pack(side="left")

        canvas_frame = tk.Frame(self, bg=_C["panel"],
                                highlightthickness=1, highlightbackground=_C["border"])
        canvas_frame.pack(fill="both", expand=True, padx=16, pady=(0, 8))

        self._canvas = _CanvasGraph(canvas_frame)
        vsb = tk.Scrollbar(canvas_frame, orient="vertical",
                           command=self._canvas.yview,
                           bg=_C["border"], troughcolor=_C["bg"],
                           activebackground=_C["accent"], relief="flat", width=10)
        self._canvas.configure(yscrollcommand=vsb.set)
        self._canvas.pack(side="left", fill="both", expand=True)
        vsb.pack(side="right", fill="y")
        self._canvas.bind("<Configure>", lambda e: self._redraw())

        detail_frame = tk.Frame(self, bg=_C["panel"],
                                highlightthickness=1, highlightbackground=_C["border"])
        detail_frame.pack(fill="x", padx=16, pady=(0, 16))
        tk.Label(detail_frame, text="EVENT DETAILS", font=_FONT_HEAD,
                 bg=_C["panel"], fg=_C["accent"]).pack(anchor="w", padx=14, pady=(8, 2))
        self._detail = _StyledScrolledText(detail_frame, height=6)
        self._detail.pack(fill="x", padx=6, pady=(0, 8))

    def _cat_color(self, cat):
        return {
            "filesystem":  _C["success"],
            "exif":        _C["accent2"],
            "document":    "#a78bfa",
            "executable":  _C["danger"],
            "audio":       _C["warn"],
        }.get(cat, _C["text_dim"])

    def _redraw(self):
        c = self._canvas
        c.delete("all")
        if not self._events:
            c.create_text(c.winfo_width() // 2, 60,
                          text="No timeline data loaded",
                          fill=_C["muted"], font=_FONT_HEAD)
            return

        W       = max(c.winfo_width(), 600)
        PAD     = 160
        ROW     = 36
        DOT     = 7
        total_h = max(400, len(self._events) * ROW + 80)
        c.configure(scrollregion=(0, 0, W, total_h))

        spine_x = PAD - 24
        c.create_line(spine_x, 20, spine_x, total_h - 20,
                      fill=_C["border"], width=2)

        prev_date = None
        for i, ev in enumerate(self._events):
            y    = 50 + i * ROW
            ts   = ev["timestamp"]
            date = ts[:10]
            col  = self._cat_color(ev["category"])

            if date != prev_date:
                c.create_text(spine_x - 10, y - 14, text=date,
                              anchor="e", fill=_C["accent"], font=_FONT_HEAD)
                prev_date = date

            c.create_oval(spine_x - DOT, y - DOT, spine_x + DOT, y + DOT,
                          fill=col, outline=_C["bg"], width=2)
            c.create_line(spine_x + DOT, y, spine_x + 32, y,
                          fill=col, width=1, dash=(4, 3))

            time_str = ts[11:19] if len(ts) > 10 else ""
            c.create_text(spine_x + 36, y, text=time_str,
                          anchor="w", fill=_C["text_dim"], font=_FONT_SMALL)

            cat_x = spine_x + 100
            bw    = 80
            c.create_rectangle(cat_x, y - 10, cat_x + bw, y + 10,
                                fill=_C["grid"], outline=col, width=1)
            c.create_text(cat_x + bw // 2, y,
                          text=ev["category"][:10].upper(),
                          fill=col, font=_FONT_SMALL)

            desc_x = cat_x + bw + 14
            txt    = f"{ev['source']}  •  {ev['description']}"
            c.create_text(desc_x, y, text=txt, anchor="w",
                          fill=_C["text"], font=_FONT_MONO)

            tag = f"ev_{i}"
            c.create_rectangle(0, y - ROW // 2, W, y + ROW // 2,
                                fill="", outline="", tags=(tag,))
            c.tag_bind(tag, "<Button-1>",
                       lambda e, ev=ev: self._show_detail(ev))
            c.tag_bind(tag, "<Enter>",
                       lambda e, tag=tag: c.itemconfigure(tag, fill=_C["sel"]))
            c.tag_bind(tag, "<Leave>",
                       lambda e, tag=tag: c.itemconfigure(tag, fill=""))

    def _show_detail(self, ev):
        lines = "\n".join([
            f"  Timestamp   : {ev['timestamp']}",
            f"  Category    : {ev['category']}",
            f"  Source      : {ev['source']}",
            f"  Description : {ev['description']}",
        ])
        self._detail.set_text(lines)

    def populate(self, metadata: dict):
        tl = metadata.get("timeline", {})
        self._events = tl.get("events", [])
        self._redraw()


# ── Tab: Risk Heatmap ─────────────────────────────────────────────────────────

class HeatmapTab(tk.Frame):
    def __init__(self, parent, app):
        super().__init__(parent, bg=_C["bg"])
        self.app = app
        self._build()

    def _build(self):
        tk.Label(self, text="RISK HEATMAP & SCORE BREAKDOWN",
                 font=_FONT_TITLE, bg=_C["bg"], fg=_C["accent"]).pack(
            padx=16, pady=(14, 8), anchor="w")

        row = tk.Frame(self, bg=_C["bg"])
        row.pack(fill="both", expand=True, padx=16, pady=(0, 16))

        left = tk.Frame(row, bg=_C["panel"],
                        highlightthickness=1, highlightbackground=_C["border"])
        left.pack(side="left", fill="both", expand=True, padx=(0, 8))
        tk.Label(left, text="SCORE CONTRIBUTORS", font=_FONT_HEAD,
                 bg=_C["panel"], fg=_C["accent"]).pack(pady=(10, 4))
        self._bar_canvas = _CanvasGraph(left)
        self._bar_canvas.pack(fill="both", expand=True, padx=10, pady=(0, 10))

        right = tk.Frame(row, bg=_C["panel"],
                         highlightthickness=1, highlightbackground=_C["border"])
        right.pack(side="left", fill="both", expand=True)
        tk.Label(right, text="RISK INDICATOR MATRIX", font=_FONT_HEAD,
                 bg=_C["panel"], fg=_C["accent"]).pack(pady=(10, 4))
        self._heat_canvas = _CanvasGraph(right)
        self._heat_canvas.pack(fill="both", expand=True, padx=10, pady=(0, 10))

        bot = tk.Frame(self, bg=_C["panel"],
                       highlightthickness=1, highlightbackground=_C["border"])
        bot.pack(fill="x", padx=16, pady=(0, 16))
        tk.Label(bot, text="DETAILED SCORE REASONS", font=_FONT_HEAD,
                 bg=_C["panel"], fg=_C["accent"]).pack(anchor="w", padx=14, pady=(8, 2))
        self._reasons_text = _StyledScrolledText(bot, height=6)
        self._reasons_text.add_tag("plus",  foreground=_C["danger"])
        self._reasons_text.add_tag("clean", foreground=_C["success"])
        self._reasons_text.pack(fill="x", padx=6, pady=(0, 8))

    def _draw_bars(self, reasons: list):
        c = self._bar_canvas
        c.delete("all")
        W = max(c.winfo_width(), 300)
        H = max(c.winfo_height(), 300)

        if not reasons:
            c.create_text(W // 2, H // 2, text="No risk contributors",
                          fill=_C["muted"], font=_FONT_HEAD)
            return

        parsed = []
        for r in reasons:
            m   = re.search(r"\+(\d+)", r)
            val = int(m.group(1)) if m else 10
            label = r.split("=")[0][:22]
            parsed.append((label, val))

        max_val = max(v for _, v in parsed) or 1
        bar_h   = min(30, (H - 40) // max(len(parsed), 1))
        gap     = 4
        label_w = 160

        for i, (label, val) in enumerate(parsed):
            y     = 20 + i * (bar_h + gap)
            t     = val / max_val
            color = _lerp_color(_C["success"], _C["danger"], min(t, 1))
            bar_w = int((W - label_w - 60) * t)

            c.create_text(label_w - 6, y + bar_h // 2,
                          text=label, anchor="e",
                          fill=_C["text"], font=_FONT_SMALL)
            c.create_rectangle(label_w, y, label_w + bar_w, y + bar_h,
                                fill=color, outline="")
            c.create_text(label_w + bar_w + 6, y + bar_h // 2,
                          text=f"+{val}", anchor="w",
                          fill=color, font=_FONT_SMALL)

    def _draw_heatmap(self, metadata: dict):
        c = self._heat_canvas
        c.delete("all")
        W = max(c.winfo_width(), 260)
        H = max(c.winfo_height(), 260)

        # ── FIX: pre-extract nested dicts to avoid {{}} inside f-strings ──
        finfo    = metadata.get("forensic_info", {})
        ti       = metadata.get("threat_intel", {})
        iocs     = ti.get("iocs", [])
        reasons  = ti.get("risk_reasons", [])
        yara_m   = ti.get("yara_matches", [])
        ex_meta  = metadata.get("extracted_metadata", {})
        pdf_meta = ex_meta.get("pdf", {})

        entropy_raw = finfo.get("entropy", 0)
        try:
            entropy_val = float(entropy_raw or 0)
        except (TypeError, ValueError):
            entropy_val = 0.0

        indicators = [
            ("High Entropy",   entropy_val > 7.5),
            ("Known Hash",     any("Known-malicious" in i for i in iocs)),
            ("YARA Match",     bool(yara_m)),
            ("Packer Found",   any("Packer" in i for i in iocs)),
            ("Trailing Data",  any("trailing" in i.lower() for i in iocs)),
            ("GPS Leak",       any("GPS" in i for i in iocs)),
            ("PE Anomaly",     any("PE" in r for r in reasons)),
            ("Suspicious Flag",bool(finfo.get("suspicious_flags"))),
            ("IOC Present",    bool(iocs)),
            ("Script Pattern", any("shell" in i.lower() or "script" in i.lower() for i in iocs)),
            ("Encrypted?",     bool(pdf_meta.get("is_encrypted"))),
            ("Base64 PE",      any("Base64" in i for i in iocs)),
            ("Webshell?",      any("webshell" in m.get("rule", "").lower() for m in yara_m)),
            ("RevShell?",      any("Reverse" in m.get("rule", "") for m in yara_m)),
            ("CryptoMiner?",   any("Crypto" in m.get("rule", "") for m in yara_m)),
        ]

        COLS   = 5
        ROWS   = math.ceil(len(indicators) / COLS)
        cell_w = (W - 20) / COLS
        cell_h = (H - 20) / ROWS

        for idx, (label, active) in enumerate(indicators):
            col     = idx % COLS
            row_idx = idx // COLS
            x1      = 10 + col * cell_w
            y1      = 10 + row_idx * cell_h
            x2      = x1 + cell_w - 4
            y2      = y1 + cell_h - 4
            fill    = _C["danger"] if active else _C["grid"]
            outline = _C["warn"]   if active else _C["border"]
            c.create_rectangle(x1, y1, x2, y2, fill=fill, outline=outline, width=1)
            c.create_text((x1 + x2) // 2, (y1 + y2) // 2,
                          text=label,
                          fill=_C["text"] if active else _C["muted"],
                          font=_FONT_SMALL, width=int(cell_w - 8))

    def populate(self, metadata: dict):
        ti = metadata.get("threat_intel", {})
        self._draw_bars(ti.get("risk_reasons", []))
        self._draw_heatmap(metadata)

        reasons = ti.get("risk_reasons", [])
        self._reasons_text.clear()
        if reasons:
            for r in reasons:
                self._reasons_text.append(f"  ▸  {r}\n", "plus")
        else:
            self._reasons_text.append(
                "  ✓  Risk score: 0 — no contributors found.\n", "clean"
            )


# ── Tab: YARA ─────────────────────────────────────────────────────────────────

class YaraTab(tk.Frame):
    def __init__(self, parent, app):
        super().__init__(parent, bg=_C["bg"])
        self.app = app
        self._build()

    def _build(self):
        tk.Label(self, text="🧬  YARA / MALWARE ANALYSIS",
                 font=_FONT_TITLE, bg=_C["bg"], fg=_C["accent"]).pack(
            padx=16, pady=(14, 8), anchor="w")

        self._status_frame = tk.Frame(self, bg=_C["panel"],
                                      highlightthickness=1, highlightbackground=_C["border"])
        self._status_frame.pack(fill="x", padx=16, pady=(0, 10))
        self._status_lbl = tk.Label(self._status_frame, text="Awaiting scan...",
                                    font=_FONT_HEAD, bg=_C["panel"], fg=_C["muted"])
        self._status_lbl.pack(side="left", padx=14, pady=8)
        self._avail_lbl = tk.Label(
            self._status_frame,
            text=(f"YARA engine: {'✓ ACTIVE' if YARA_AVAILABLE else '✗ NOT INSTALLED  (pip install yara-python)'}"),
            font=_FONT_SMALL, bg=_C["panel"],
            fg=_C["success"] if YARA_AVAILABLE else _C["danger"]
        )
        self._avail_lbl.pack(side="right", padx=14)

        tree_frame = tk.Frame(self, bg=_C["panel"],
                              highlightthickness=1, highlightbackground=_C["border"])
        tree_frame.pack(fill="both", expand=True, padx=16, pady=(0, 10))
        tk.Label(tree_frame, text="RULE MATCHES", font=_FONT_HEAD,
                 bg=_C["panel"], fg=_C["accent"]).pack(anchor="w", padx=14, pady=(8, 2))

        style = ttk.Style()
        style.theme_use("default")
        style.configure("Forensic.Treeview",
                         background=_C["panel"], foreground=_C["text"],
                         fieldbackground=_C["panel"], rowheight=26, font=_FONT_MONO)
        style.configure("Forensic.Treeview.Heading",
                         background=_C["grid"], foreground=_C["accent"],
                         font=_FONT_HEAD, relief="flat")
        style.map("Forensic.Treeview",
                  background=[("selected", _C["sel"])],
                  foreground=[("selected", _C["text"])])

        cols = ("Rule", "Severity", "Description", "Score Bump")
        self._tree = ttk.Treeview(tree_frame, columns=cols, show="headings",
                                   style="Forensic.Treeview", height=8)
        for col in cols:
            self._tree.heading(col, text=col)
            self._tree.column(col, width={"Rule": 140, "Severity": 80,
                                           "Description": 300, "Score Bump": 80}[col])
        vsb = ttk.Scrollbar(tree_frame, orient="vertical", command=self._tree.yview)
        self._tree.configure(yscrollcommand=vsb.set)
        self._tree.pack(side="left", fill="both", expand=True, padx=6, pady=(0, 8))
        vsb.pack(side="right", fill="y", pady=(0, 8))
        self._tree.bind("<<TreeviewSelect>>", self._on_select)

        detail_frame = tk.Frame(self, bg=_C["panel"],
                                highlightthickness=1, highlightbackground=_C["border"])
        detail_frame.pack(fill="x", padx=16, pady=(0, 16))
        tk.Label(detail_frame, text="MATCH DETAIL", font=_FONT_HEAD,
                 bg=_C["panel"], fg=_C["accent"]).pack(anchor="w", padx=14, pady=(8, 2))
        self._detail = _StyledScrolledText(detail_frame, height=7)
        self._detail.add_tag("key", foreground=_C["accent"])
        self._detail.add_tag("val", foreground=_C["text"])
        self._detail.add_tag("off", foreground=_C["warn"])
        self._detail.pack(fill="x", padx=6, pady=(0, 8))

        self._matches: list = []

    def _on_select(self, _):
        sel = self._tree.selection()
        if not sel:
            return
        rule_name = self._tree.item(sel[0])["values"][0]
        match = next((m for m in self._matches if m["rule"] == rule_name), None)
        if not match:
            return
        self._detail.clear()
        self._detail.append("  Rule        : ", "key")
        self._detail.append(f"{match['rule']}\n", "val")
        self._detail.append("  Severity    : ", "key")
        self._detail.append(f"{match['severity'].upper()}\n", "val")
        self._detail.append("  Description : ", "key")
        self._detail.append(f"{match['description']}\n", "val")
        self._detail.append("  Score Bump  : ", "key")
        self._detail.append(f"+{match['score_bump']}\n", "val")
        if match.get("strings"):
            self._detail.append("  Matches     :\n", "key")
            for offset, ident in match["strings"]:
                self._detail.append(f"    @ {offset}  {ident}\n", "off")

    def populate(self, metadata: dict):
        ti = metadata.get("threat_intel", {})
        self._matches = ti.get("yara_matches", [])
        total = len(self._matches)

        if total:
            self._status_lbl.configure(
                text=f"⚡  {total} YARA rule(s) matched — threat indicators detected!",
                fg=_C["danger"]
            )
        else:
            self._status_lbl.configure(
                text="✓  No YARA rules matched — no known signatures found.",
                fg=_C["success"]
            )

        for row in self._tree.get_children():
            self._tree.delete(row)

        sev_order = {"critical": 0, "high": 1, "medium": 2, "low": 3}
        for m in sorted(self._matches,
                         key=lambda x: sev_order.get(x.get("severity", ""), 9)):
            self._tree.insert("", "end", values=(
                m["rule"], m["severity"].upper(),
                m["description"], f"+{m['score_bump']}"
            ))

        self._detail.clear()


# ── Tab: Chain of Custody ─────────────────────────────────────────────────────

class CocTab(tk.Frame):
    def __init__(self, parent, app):
        super().__init__(parent, bg=_C["bg"])
        self.app = app
        self._build()

    def _build(self):
        tk.Label(self, text="🔐  CHAIN OF CUSTODY",
                 font=_FONT_TITLE, bg=_C["bg"], fg=_C["accent"]).pack(
            padx=16, pady=(14, 8), anchor="w")

        top = tk.Frame(self, bg=_C["bg"])
        top.pack(fill="x", padx=16, pady=(0, 10))
        self._integrity_lbl = tk.Label(top, text="Integrity: –",
                                       font=_FONT_HEAD, bg=_C["bg"], fg=_C["muted"])
        self._integrity_lbl.pack(side="left")
        self._record_lbl = tk.Label(top, text="Record: –",
                                    font=_FONT_SMALL, bg=_C["bg"], fg=_C["muted"])
        self._record_lbl.pack(side="right")

        canvas_frame = tk.Frame(self, bg=_C["panel"],
                                highlightthickness=1, highlightbackground=_C["border"])
        canvas_frame.pack(fill="both", expand=True, padx=16, pady=(0, 10))
        tk.Label(canvas_frame, text="AUDIT TRAIL", font=_FONT_HEAD,
                 bg=_C["panel"], fg=_C["accent"]).pack(anchor="w", padx=14, pady=(8, 2))

        self._canvas = _CanvasGraph(canvas_frame)
        vsb = tk.Scrollbar(canvas_frame, orient="vertical",
                           command=self._canvas.yview,
                           bg=_C["border"], troughcolor=_C["bg"],
                           activebackground=_C["accent"], relief="flat", width=10)
        self._canvas.configure(yscrollcommand=vsb.set)
        self._canvas.pack(side="left", fill="both", expand=True, padx=6, pady=(0, 8))
        vsb.pack(side="right", fill="y", pady=(0, 8))

        hash_frame = tk.Frame(self, bg=_C["panel"],
                              highlightthickness=1, highlightbackground=_C["border"])
        hash_frame.pack(fill="x", padx=16, pady=(0, 16))
        tk.Label(hash_frame, text="HASH VERIFICATION LOG", font=_FONT_HEAD,
                 bg=_C["panel"], fg=_C["accent"]).pack(anchor="w", padx=14, pady=(8, 2))
        self._hash_text = _StyledScrolledText(hash_frame, height=5)
        self._hash_text.add_tag("ok",   foreground=_C["success"])
        self._hash_text.add_tag("fail", foreground=_C["danger"])
        self._hash_text.add_tag("dim",  foreground=_C["muted"])
        self._hash_text.pack(fill="x", padx=6, pady=(0, 8))

    def _draw_chain(self, events: list):
        c = self._canvas
        c.delete("all")
        W       = max(c.winfo_width(), 600)
        ROW     = 56
        total_h = max(400, len(events) * ROW + 60)
        c.configure(scrollregion=(0, 0, W, total_h))

        action_color = {
            "CUSTODY_OPENED":     _C["success"],
            "EVIDENCE_ACQUIRED":  _C["accent2"],
            "ANALYSIS_PERFORMED": _C["accent"],
            "HASH_VERIFICATION":  _C["warn"],
            "EVIDENCE_EXPORTED":  "#a78bfa",
            "CUSTODY_CLOSED":     _C["success"],
        }

        for i, ev in enumerate(events):
            y     = 30 + i * ROW
            color = action_color.get(ev["action"], _C["muted"])

            if i < len(events) - 1:
                c.create_line(40, y + 14, 40, y + ROW, fill=_C["border"], width=2)

            c.create_rectangle(18, y - 2, 62, y + 28,
                                fill=_C["grid"], outline=color, width=1)
            c.create_text(40, y + 13, text=str(ev["seq"]),
                          fill=color, font=_FONT_HEAD)

            c.create_text(78, y + 4, text=ev["action"],
                          anchor="w", fill=color, font=_FONT_HEAD)

            ts      = ev["timestamp"][:19]
            details = ", ".join(
                f"{k}={v}" for k, v in ev["details"].items()
                if isinstance(v, (str, int, bool)) and k not in ("note", "hashes")
            )[:80]
            c.create_text(78, y + 20,
                          text=f"{ts}   {details}",
                          anchor="w", fill=_C["text_dim"], font=_FONT_SMALL)

            h_snippet = ev.get("entry_hash", "")[:20] + "…"
            c.create_text(W - 10, y + 4,
                          text=f"hash: {h_snippet}", anchor="e",
                          fill=_C["border"], font=_FONT_SMALL)

    def populate(self, metadata: dict):
        coc = metadata.get("chain_of_custody", {})
        if not coc:
            self._integrity_lbl.configure(
                text="Chain of Custody: not recorded in this session",
                fg=_C["muted"]
            )
            return

        valid = coc.get("integrity_valid", False)
        self._integrity_lbl.configure(
            text=f"Integrity: {'✓  CHAIN INTACT' if valid else '✗  CHAIN COMPROMISED — TAMPERED?'}",
            fg=_C["success"] if valid else _C["danger"]
        )
        self._record_lbl.configure(
            text=f"Record ID: {coc.get('record_id', '')[:36]}",
            fg=_C["text_dim"]
        )

        self._draw_chain(coc.get("events", []))

        self._hash_text.clear()
        hashes = metadata.get("hashes", {})
        self._hash_text.append("  Acquisition Hashes\n", "dim")
        for algo, digest in hashes.items():
            self._hash_text.append(f"  {algo.upper():<8}  {digest}\n", "ok")

        for ev in coc.get("events", []):
            if ev["action"] == "HASH_VERIFICATION":
                ok = ev["details"].get("integrity_ok", False)
                self._hash_text.append(
                    f"\n  Verification: {'PASS ✓' if ok else 'FAIL ✗  — evidence may be altered'}\n",
                    "ok" if ok else "fail"
                )


# ── Tab: Raw JSON ─────────────────────────────────────────────────────────────

class RawTab(tk.Frame):
    def __init__(self, parent, app):
        super().__init__(parent, bg=_C["bg"])
        self.app = app
        self._build()

    def _build(self):
        tk.Label(self, text="RAW JSON OUTPUT", font=_FONT_TITLE,
                 bg=_C["bg"], fg=_C["accent"]).pack(padx=16, pady=(14, 8), anchor="w")
        self._text = _StyledScrolledText(self)
        self._text.add_tag("key",    foreground=_C["accent"])
        self._text.add_tag("string", foreground=_C["success"])
        self._text.add_tag("number", foreground=_C["warn"])
        self._text.pack(fill="both", expand=True, padx=16, pady=(0, 16))

    def populate(self, metadata: dict):
        self._text.set_text(json.dumps(metadata, indent=2, default=str))


# ── Main GUI Application ───────────────────────────────────────────────────────

class ForensicGUI(tk.Tk):
    """Main GUI window. Analysis runs on a background thread."""

    def __init__(self):
        super().__init__()
        self.title(f"{TOOL_NAME}  v{VERSION}")
        self.configure(bg=_C["bg"])
        self.geometry("1280x860")
        self.minsize(900, 600)

        self._metadata: Optional[dict] = None
        self._file_path: str = ""

        self._build_menu()
        self._build_toolbar()
        self._build_notebook()
        self._build_statusbar()

        self.protocol("WM_DELETE_WINDOW", self.destroy)

    def _build_menu(self):
        mb = tk.Menu(self, bg=_C["surface"], fg=_C["text"],
                     activebackground=_C["sel"], activeforeground=_C["accent"],
                     relief="flat", bd=0)
        self.configure(menu=mb)

        file_m = tk.Menu(mb, tearoff=0, bg=_C["surface"], fg=_C["text"],
                          activebackground=_C["sel"], activeforeground=_C["accent"])
        file_m.add_command(label="Open File…",      command=self.open_file)
        file_m.add_command(label="Open Directory…", command=self.open_directory)
        file_m.add_separator()
        file_m.add_command(label="Export JSON",     command=lambda: self._export("json"))
        file_m.add_command(label="Export HTML",     command=lambda: self._export("html"))
        file_m.add_command(label="Export CSV",      command=lambda: self._export("csv"))
        file_m.add_separator()
        file_m.add_command(label="Quit",            command=self.destroy)
        mb.add_cascade(label="File", menu=file_m)

        help_m = tk.Menu(mb, tearoff=0, bg=_C["surface"], fg=_C["text"],
                          activebackground=_C["sel"], activeforeground=_C["accent"])
        help_m.add_command(label="About", command=self._about)
        mb.add_cascade(label="Help", menu=help_m)

    def _build_toolbar(self):
        tb = tk.Frame(self, bg=_C["surface"], height=46,
                      highlightthickness=1, highlightbackground=_C["border"])
        tb.pack(fill="x")
        tb.pack_propagate(False)

        tk.Label(tb, text="⬡ FORENSISCAN", font=_FONT_HEAD,
                 bg=_C["surface"], fg=_C["accent"]).pack(side="left", padx=16)

        btn_cfg = dict(bg=_C["grid"], fg=_C["text"], relief="flat",
                       font=_FONT_SMALL, cursor="hand2",
                       activebackground=_C["sel"], activeforeground=_C["accent"],
                       padx=12, pady=4)

        tk.Button(tb, text="📂  Open File",
                  command=self.open_file, **btn_cfg).pack(side="left", padx=4, pady=6)
        tk.Button(tb, text="📁  Open Directory",
                  command=self.open_directory, **btn_cfg).pack(side="left", padx=4, pady=6)
        tk.Button(tb, text="💾  Export",
                  command=self._export_dialog, **btn_cfg).pack(side="left", padx=4, pady=6)

        self._path_lbl = tk.Label(tb, text="No file loaded",
                                  font=_FONT_SMALL, bg=_C["surface"], fg=_C["muted"])
        self._path_lbl.pack(side="left", padx=16)

        tk.Label(tb, text="Case ID:", font=_FONT_SMALL,
                 bg=_C["surface"], fg=_C["muted"]).pack(side="right", padx=(0, 4))
        self._case_var = tk.StringVar(
            value=f"CASE-{datetime.now().strftime('%Y%m%d-%H%M%S')}"
        )
        tk.Entry(tb, textvariable=self._case_var, width=22, font=_FONT_SMALL,
                 bg=_C["grid"], fg=_C["accent"], insertbackground=_C["accent"],
                 relief="flat", bd=4).pack(side="right", padx=(0, 12))
        tk.Label(tb, text="Examiner:", font=_FONT_SMALL,
                 bg=_C["surface"], fg=_C["muted"]).pack(side="right", padx=(0, 4))
        self._examiner_var = tk.StringVar(value="Unknown")
        tk.Entry(tb, textvariable=self._examiner_var, width=14, font=_FONT_SMALL,
                 bg=_C["grid"], fg=_C["text"], insertbackground=_C["accent"],
                 relief="flat", bd=4).pack(side="right", padx=(0, 8))

    def _build_notebook(self):
        style = ttk.Style()
        style.theme_use("default")
        style.configure("Forensic.TNotebook",
                         background=_C["bg"], borderwidth=0)
        style.configure("Forensic.TNotebook.Tab",
                         background=_C["surface"], foreground=_C["muted"],
                         font=_FONT_HEAD, padding=(16, 6), borderwidth=0)
        style.map("Forensic.TNotebook.Tab",
                  background=[("selected", _C["panel"])],
                  foreground=[("selected", _C["accent"])])

        nb = ttk.Notebook(self, style="Forensic.TNotebook")
        nb.pack(fill="both", expand=True)

        self._dash = DashboardTab(nb, self)
        self._tl   = TimelineTab(nb, self)
        self._heat = HeatmapTab(nb, self)
        self._yara = YaraTab(nb, self)
        self._coc  = CocTab(nb, self)
        self._raw  = RawTab(nb, self)

        nb.add(self._dash, text="  📊 Dashboard  ")
        nb.add(self._tl,   text="  🧾 Timeline   ")
        nb.add(self._heat, text="  🔥 Risk Heatmap")
        nb.add(self._yara, text="  🧬 YARA        ")
        nb.add(self._coc,  text="  🔐 Chain of Custody")
        nb.add(self._raw,  text="  { } Raw JSON  ")

        self._nb = nb

    def _build_statusbar(self):
        sb = tk.Frame(self, bg=_C["surface"], height=24,
                      highlightthickness=1, highlightbackground=_C["border"])
        sb.pack(fill="x", side="bottom")
        sb.pack_propagate(False)
        self._status_lbl = tk.Label(sb, text="Ready.", font=_FONT_SMALL,
                                    bg=_C["surface"], fg=_C["muted"])
        self._status_lbl.pack(side="left", padx=12)
        self._progress = ttk.Progressbar(sb, mode="indeterminate", length=120)
        self._progress.pack(side="right", padx=12, pady=3)

    def _set_status(self, msg: str, busy: bool = False):
        self._status_lbl.configure(text=msg)
        if busy:
            self._progress.start(10)
        else:
            self._progress.stop()

    def open_file(self):
        path = filedialog.askopenfilename(
            title="Select file to analyse",
            filetypes=[
                ("All files",     "*.*"),
                ("Images",        "*.jpg *.jpeg *.png *.gif *.bmp *.tiff *.webp"),
                ("Documents",     "*.pdf *.docx *.doc"),
                ("Archives",      "*.zip *.jar *.apk"),
                ("Executables",   "*.exe *.dll"),
                ("Audio/Video",   "*.mp3 *.mp4 *.avi *.mkv *.flac"),
            ]
        )
        if path:
            self._file_path = path
            self._path_lbl.configure(text=path[-70:], fg=_C["text"])
            self._run_extraction(path, single=True)

    def open_directory(self):
        path = filedialog.askdirectory(title="Select directory to scan")
        if path:
            self._file_path = path
            self._path_lbl.configure(text=path[-70:], fg=_C["text"])
            self._run_extraction(path, single=False)

    def _run_extraction(self, path: str, single: bool):
        import threading
        self._set_status(f"Analysing {Path(path).name}…", busy=True)

        def worker():
            try:
                case_id  = self._case_var.get()
                examiner = self._examiner_var.get() or "Unknown"
                logger   = ForensicLogger(case_id=case_id)
                coc      = ChainOfCustody(case_id=case_id, examiner=examiner)

                if single:
                    extractor = EnhancedMetadataExtractor(
                        path, case_id=case_id, logger=logger,
                        chain_of_custody=coc, examiner=examiner
                    )
                    metadata = extractor.extract_all()
                    coc.close()
                    self.after(0, self._populate_all, metadata)
                else:
                    scanner = ParallelScanner(
                        max_workers=DEFAULT_WORKERS, case_id=case_id,
                        logger=logger, chain_of_custody=coc, examiner=examiner
                    )
                    results = scanner.scan_directory(Path(path), recursive=False)
                    coc.close()
                    if results:
                        first = results[0]
                        first["chain_of_custody"] = coc.to_dict()
                        self.after(0, self._populate_all, first)
                        ts  = datetime.now().strftime("%Y%m%d_%H%M%S")
                        out = str(Path(path) / f"forensic_report_{ts}.html")
                        save_to_html(results, out)
                        self.after(0, self._set_status,
                                   f"✓ Scanned {len(results)} files — report saved: {out}")
                    else:
                        self.after(0, self._set_status, "No files found in directory.")

            except Exception as exc:
                self.after(0, messagebox.showerror, "Extraction Error", str(exc))
                self.after(0, self._set_status, f"Error: {exc}")

        threading.Thread(target=worker, daemon=True).start()

    def _populate_all(self, metadata: dict):
        self._metadata = metadata
        self._dash.populate(metadata)
        self._tl.populate(metadata)
        self._heat.populate(metadata)
        self._yara.populate(metadata)
        self._coc.populate(metadata)
        self._raw.populate(metadata)

        fi    = metadata.get("file_info", {})
        score = metadata.get("threat_intel", {}).get("risk_score", 0)
        level = metadata.get("threat_intel", {}).get("risk_level", "CLEAN")
        self._set_status(
            f"✓  {fi.get('filename', '')}  |  Risk: {score}/100 ({level})"
            f"  |  {fi.get('file_size_human', '')}",
            busy=False
        )

    def _export_dialog(self):
        if not self._metadata:
            messagebox.showwarning("No Data", "Please analyse a file first.")
            return
        win = tk.Toplevel(self, bg=_C["surface"])
        win.title("Export Report")
        win.geometry("360x180")
        win.resizable(False, False)
        tk.Label(win, text="Choose export format:", font=_FONT_HEAD,
                 bg=_C["surface"], fg=_C["text"]).pack(pady=(20, 10))
        btn_f = tk.Frame(win, bg=_C["surface"])
        btn_f.pack()
        btn_cfg = dict(bg=_C["grid"], fg=_C["text"], relief="flat",
                       font=_FONT_HEAD, cursor="hand2", padx=16, pady=8,
                       activebackground=_C["sel"], activeforeground=_C["accent"])
        for fmt, lbl in (("json", "JSON"), ("html", "HTML"), ("csv", "CSV")):
            tk.Button(btn_f, text=lbl,
                      command=lambda f=fmt, w=win: (self._export(f), w.destroy()),
                      **btn_cfg).pack(side="left", padx=8)

    def _export(self, fmt: str):
        if not self._metadata:
            messagebox.showwarning("No Data", "Please analyse a file first.")
            return
        ext  = {"json": ".json", "html": ".html", "csv": ".csv"}.get(fmt, ".json")
        path = filedialog.asksaveasfilename(
            defaultextension=ext,
            filetypes=[(fmt.upper(), f"*{ext}"), ("All", "*.*")],
            initialfile=f"forensic_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}{ext}"
        )
        if path:
            save_output(self._metadata, fmt, path)
            self._coc.populate(self._metadata)
            self._set_status(f"✓ Exported to {path}")

    def _about(self):
        messagebox.showinfo(
            f"About {TOOL_NAME}",
            f"{TOOL_NAME}  v{VERSION}\n\n"
            f"Author: Sayan Pal\nCollaborator: Soumit Santra\n\n"
            f"Features:\n"
            f"  🧠 Threat Intelligence & IOC Detection\n"
            f"  🧬 YARA Malware Scanning\n"
            f"  🧾 Timeline Reconstruction\n"
            f"  🔐 Chain of Custody\n"
            f"  🖥️  Full GUI Dashboard\n\n"
            f"⚠️  For authorised forensic use only."
        )


# ═══════════════════════════════════════════════════════════════════════════════
#  STARTUP LAUNCHER
# ═══════════════════════════════════════════════════════════════════════════════

def _launch_gui():
    if not TK_AVAILABLE:
        print("❌ tkinter is not available in this Python installation.")
        print("   On Debian/Ubuntu: sudo apt install python3-tk")
        sys.exit(1)
    app = ForensicGUI()
    app.mainloop()


def _startup_prompt() -> str:
    print_banner()
    print(f"\n{'─'*60}")
    print("  How do you want to run the tool?")
    print(f"{'─'*60}")
    print("  [1]  CLI  – Terminal / interactive menu")
    print("  [2]  GUI  – Graphical dashboard (tkinter)")
    print(f"{'─'*60}")
    while True:
        choice = input("\n  Enter 1 or 2: ").strip()
        if choice == "1":
            return "cli"
        elif choice == "2":
            return "gui"
        else:
            print("  Please enter 1 or 2.")


if __name__ == "__main__":
    if len(sys.argv) > 1:
        main()
    else:
        mode = _startup_prompt()
        if mode == "gui":
            _launch_gui()
        else:
            interactive_mode()
