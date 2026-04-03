"""
ForensiScan v6.0 — Advanced File Intelligence & Forensic Analysis Tool
=======================================================================
Author      : Sayan Pal
Collaborator: Soumit Santra
"""

# stdlib
import os, sys, json, argparse, subprocess, logging, platform, zipfile, shutil
import hashlib, math, csv, uuid, re, struct, threading, mmap, io, time
from datetime import datetime, timezone, timedelta
from pathlib import Path
from typing import Dict, Any, List, Optional, Tuple, Iterator
from concurrent.futures import ThreadPoolExecutor, as_completed
from threading import Lock

# optional third-party 
try:
    from colorama import init as _cinit, Fore, Style
    _cinit(autoreset=True)
    COLORAMA_AVAILABLE = True
except ImportError:
    COLORAMA_AVAILABLE = False
    class _MC:
        def __getattr__(self, _): return ""
    Fore = Style = _MC()

try:
    from tqdm import tqdm
except ImportError:
    def tqdm(it=None, **kw): return it or []

try:
    from PIL import Image
    from PIL.ExifTags import TAGS, GPSTAGS
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False

try:
    from hachoir.parser import createParser
    from hachoir.metadata import extractMetadata
    HACHOIR_AVAILABLE = True
except ImportError:
    HACHOIR_AVAILABLE = False

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import mutagen
    AUDIO_AVAILABLE = True
except ImportError:
    AUDIO_AVAILABLE = False

try:
    import magic
    MAGIC_AVAILABLE = True
except ImportError:
    MAGIC_AVAILABLE = False

try:
    import pefile
    PE_AVAILABLE = True
except ImportError:
    PE_AVAILABLE = False

try:
    import yara
    YARA_AVAILABLE = True
except ImportError:
    YARA_AVAILABLE = False

try:
    import pytsk3
    TSK_AVAILABLE = True
except ImportError:
    TSK_AVAILABLE = False

try:
    import pyewf
    EWF_AVAILABLE = True
except ImportError:
    EWF_AVAILABLE = False

# constants
VERSION          = "6.0.0"
TOOL_NAME        = "ForensiScan"
DEFAULT_WORKERS  = 4
EXIFTOOL_TIMEOUT = 30
CHUNK_SIZE       = 65_536          # 64 KB read chunk for streaming ops
MMAP_THRESHOLD   = 4 * 1024 * 1024  # use mmap above 4 MB
YARA_LIMIT       = 256 * 1024 * 1024  # YARA only on files ≤ 256 MB

KNOWN_MALICIOUS_HASHES: Dict[str, str] = {
    "44d88612fea8a8f36de82e1278abb02f":
        "EICAR Test File",
    "e6d290a03b70cfa5d4451da444bdea39":
        "Mirai Botnet Sample",
    "098f6bcd4621d373cade4e832627b4f6":
        "Test Hash (demo)",
    "275a021bbfb6489e54d471899f7db9d1663fc695ec2fe2a2c4538aabf651fd0f":
        "EICAR Test File (SHA256)",
}

SUSPICIOUS_PATTERNS: List[Tuple[bytes, str]] = [
    (b"This program cannot be run in DOS mode", "Embedded PE binary"),
    (b"TVqQAAMAAAAEAAAA",  "Base64-encoded PE"),
    (b"powershell",        "PowerShell reference"),
    (b"cmd.exe",           "cmd.exe reference"),
    (b"eval(base64_decode","PHP webshell pattern"),
    (b"EICAR-STANDARD-ANTIVIRUS-TEST-FILE", "EICAR test string"),
    (b"\x60\x89\xe5\x31\xc0", "Linux shellcode pattern"),
]

PACKER_SIGNATURES: List[Tuple[bytes, str]] = [
    (b"UPX0", "UPX Packer"), (b"UPX1", "UPX Packer"), (b"UPX!", "UPX Packer"),
    (b".MPRESS1", "MPRESS Packer"), (b"PECompact2", "PECompact Packer"),
    (b"ASPack", "ASPack Packer"), (b"Themida", "Themida Protector"),
    (b"ExECryptor", "ExECryptor Protector"),
]

BUILTIN_YARA_RULES = r"""
rule Detect_EICAR {
    meta: description="Detects EICAR test string" severity="low"
    strings: $a="EICAR-STANDARD-ANTIVIRUS-TEST-FILE"
    condition: $a }
rule Detect_UPX_Packed {
    meta: description="UPX-packed executable" severity="medium"
    strings: $a="UPX0" $b="UPX!"
    condition: any of them }
rule Detect_PowerShell_Download {
    meta: description="PowerShell download cradle" severity="high"
    strings: $a="DownloadString" nocase $b="WebClient" nocase
             $c="Invoke-Expression" nocase wide ascii
    condition: 2 of them }
rule Detect_PHP_Webshell {
    meta: description="PHP webshell indicators" severity="critical"
    strings: $a="eval($_POST" nocase $b="eval($_GET" nocase
             $c="system($_" nocase $d="passthru($_" nocase
    condition: any of them }
rule Detect_Base64_PE {
    meta: description="Base64-encoded PE file" severity="high"
    strings: $a="TVqQAAMAAAAEAAAA" $b="TVpAAA"
    condition: any of them }
rule Detect_Reverse_Shell {
    meta: description="Common reverse-shell patterns" severity="critical"
    strings: $a="/bin/sh" nocase $b="bash -i" nocase
             $c="nc -e" nocase $d="/dev/tcp/" nocase
    condition: 2 of them }
rule Detect_Crypto_Mining {
    meta: description="Cryptocurrency mining indicators" severity="medium"
    strings: $a="stratum+tcp://" nocase $b="xmrig" nocase
             $c="monero" nocase $d="cryptonight" nocase
    condition: any of them }
"""

# ── COMPRESSED_EXTENSIONS (used throughout) ───────────────────────────────────
COMPRESSED_EXT = frozenset({
    ".jpg",".jpeg",".png",".gif",".webp",".heic",".bmp",
    ".zip",".rar",".7z",".gz",".tar",".bz2",".xz",
    ".jar",".apk",".pdf",".docx",".xlsx",".pptx",
    ".mp3",".mp4",".avi",".mov",".mkv",".flac",".ogg",".wav",
})

#  STREAMING MEMORY HELPERS  

def stream_chunks(path: Path, chunk: int = CHUNK_SIZE) -> Iterator[bytes]:
    """Yield fixed-size chunks from a file without loading it all into RAM."""
    with open(path, "rb") as fh:
        while True:
            data = fh.read(chunk)
            if not data:
                break
            yield data


def compute_hashes_streaming(path: Path) -> Dict[str, str]:
    """Compute MD5 / SHA-1 / SHA-256 in a single streaming pass."""
    md5    = hashlib.md5()
    sha1   = hashlib.sha1()
    sha256 = hashlib.sha256()
    for chunk in stream_chunks(path):
        md5.update(chunk); sha1.update(chunk); sha256.update(chunk)
    return {"md5": md5.hexdigest(), "sha1": sha1.hexdigest(), "sha256": sha256.hexdigest()}


def compute_entropy_streaming(path: Path) -> float:
    """
    Compute Shannon entropy with O(256) RAM — only a byte-frequency table
    is kept; the file is never fully buffered.
    """
    freq  = [0] * 256
    total = 0
    for chunk in stream_chunks(path):
        total += len(chunk)
        for b in chunk:
            freq[b] += 1
    if total == 0:
        return 0.0
    entropy = 0.0
    for c in freq:
        if c:
            p = c / total
            entropy -= p * math.log2(p)
    return entropy


def open_mmap_or_bytes(path: Path, size: int) -> Tuple[Any, bool]:
    """
    Return (view, is_mmap).
    For files > MMAP_THRESHOLD use mmap (zero-copy); otherwise bytes.
    Caller must call view.close() if is_mmap.
    """
    fh = open(path, "rb")
    if size >= MMAP_THRESHOLD:
        try:
            mm = mmap.mmap(fh.fileno(), 0, access=mmap.ACCESS_READ)
            fh.close()
            return mm, True
        except Exception:
            pass
    data = fh.read()
    fh.close()
    return data, False


def scan_patterns_streaming(path: Path, patterns: List[Tuple[bytes, str]],
                             max_size: int = 200 * 1024 * 1024
                             ) -> List[str]:
    """
    Search for byte patterns using a sliding window across chunks so that
    matches spanning chunk boundaries are never missed.
    Returns list of matched description strings.
    """
    size = path.stat().st_size
    if size > max_size:
        return []
    max_pat  = max((len(p) for p, _ in patterns), default=0)
    overlap  = max_pat - 1
    found    = set()
    tail     = b""
    with open(path, "rb") as fh:
        while True:
            chunk = fh.read(CHUNK_SIZE)
            if not chunk:
                break
            window = tail + chunk
            for pat, desc in patterns:
                if pat in window:
                    found.add(desc)
            tail = window[-overlap:] if overlap > 0 else b""
    return list(found)


def read_file_tail(path: Path, n_bytes: int = 8192) -> bytes:
    """Read only the last n_bytes of a file without loading the rest."""
    size = path.stat().st_size
    if size <= n_bytes:
        return path.read_bytes()
    with open(path, "rb") as fh:
        fh.seek(-n_bytes, 2)
        return fh.read()


def human_size(n: int) -> str:
    for u in ("B","KB","MB","GB","TB"):
        if n < 1024: return f"{n:.2f} {u}"
        n /= 1024
    return f"{n:.2f} PB"

#  THEME SYSTEM

THEMES: Dict[str, Dict[str, str]] = {
    "Cyber Dark": {
        "bg":"#0a0e17","bg2":"#0f1623","surface":"#141c2b","surface2":"#1a2438",
        "border":"#1e3050","border2":"#2a4070","accent":"#00d4aa","accent2":"#0096ff",
        "accent3":"#7c3aed","danger":"#ef4444","warn":"#f59e0b","success":"#10b981",
        "muted":"#4a6278","text":"#c8d8e8","text_dim":"#5a7a8a","text_bright":"#e8f4ff",
        "sel":"#1e4060","tab_active":"#0f1e30","grid":"#111a28",
        "font_main":"Courier New","font_ui":"Segoe UI",
    },
    "Arctic": {
        "bg":"#f0f4f8","bg2":"#e8edf2","surface":"#ffffff","surface2":"#f7f9fb",
        "border":"#dde3ea","border2":"#c4cdd6","accent":"#0070f3","accent2":"#7928ca",
        "accent3":"#ff0080","danger":"#cc0000","warn":"#d97706","success":"#047857",
        "muted":"#888888","text":"#111827","text_dim":"#6b7280","text_bright":"#000000",
        "sel":"#dbeafe","tab_active":"#ffffff","grid":"#f3f4f6",
        "font_main":"Consolas","font_ui":"Segoe UI",
    },
    "Obsidian": {
        "bg":"#1a1a1a","bg2":"#222222","surface":"#2a2a2a","surface2":"#333333",
        "border":"#3a3a3a","border2":"#4a4a4a","accent":"#d4a017","accent2":"#c07830",
        "accent3":"#a855f7","danger":"#dc2626","warn":"#d97706","success":"#16a34a",
        "muted":"#666666","text":"#e8e0d0","text_dim":"#888878","text_bright":"#f5f0e8",
        "sel":"#3a3020","tab_active":"#2a2a2a","grid":"#242424",
        "font_main":"Courier New","font_ui":"Georgia",
    },
    "Matrix": {
        "bg":"#000800","bg2":"#001000","surface":"#001800","surface2":"#002000",
        "border":"#003800","border2":"#005000","accent":"#00ff41","accent2":"#00cc33",
        "accent3":"#008822","danger":"#ff3300","warn":"#ffaa00","success":"#00ff41",
        "muted":"#005500","text":"#00cc33","text_dim":"#007722","text_bright":"#00ff41",
        "sel":"#003300","tab_active":"#001a00","grid":"#000e00",
        "font_main":"Courier New","font_ui":"Courier New",
    },
    "Rose Gold": {
        "bg":"#2d1b1b","bg2":"#3a2020","surface":"#4a2828","surface2":"#5a3030",
        "border":"#7a4040","border2":"#9a5050","accent":"#f4a0a0","accent2":"#e8c0c0",
        "accent3":"#ffd4d4","danger":"#ff6b6b","warn":"#ffc875","success":"#98d8b0",
        "muted":"#8a6060","text":"#f5e0e0","text_dim":"#b08080","text_bright":"#ffe8e8",
        "sel":"#5a2a2a","tab_active":"#4a2020","grid":"#3a2020",
        "font_main":"Courier New","font_ui":"Georgia",
    },
}

_T = THEMES["Cyber Dark"]

def apply_theme(t: dict):
    global _T; _T = t

#  CHAIN OF CUSTODY

class ChainOfCustody:
    def __init__(self, case_id: str, examiner: str = "Unknown",
                 output_path: Optional[str] = None):
        self.case_id = case_id; self.examiner = examiner
        self.record_id = str(uuid.uuid4()); self.events: List[Dict] = []
        self.output_path = output_path; self._lock = Lock()
        self._record("CUSTODY_OPENED", {"tool": f"{TOOL_NAME} v{VERSION}",
                     "platform": platform.system(), "python": platform.python_version()})

    def _now(self) -> str: return datetime.now(timezone.utc).isoformat()

    def _record(self, action: str, details: Dict) -> Dict:
        prev = self.events[-1]["entry_hash"] if self.events else "GENESIS"
        e = {"seq": len(self.events)+1, "timestamp": self._now(),
             "case_id": self.case_id, "examiner": self.examiner,
             "action": action, "details": details, "prev_hash": prev}
        raw = json.dumps(e, sort_keys=True, default=str)
        e["entry_hash"] = hashlib.sha256((prev+raw).encode()).hexdigest()
        with self._lock: self.events.append(e)
        if self.output_path: self._persist()
        return e

    def _persist(self):
        try:
            with open(self.output_path, "w") as f:
                json.dump({"record_id": self.record_id, "case_id": self.case_id,
                           "events": self.events, "valid": self.verify()}, f, indent=2, default=str)
        except Exception: pass

    def log_acquisition(self, fp, hashes):
        return self._record("EVIDENCE_ACQUIRED", {"file": fp, "hashes": hashes})
    def log_analysis(self, fp, op, result):
        return self._record("ANALYSIS_PERFORMED", {"file": fp, "operation": op, "result": result})
    def log_export(self, dest, fmt):
        return self._record("EVIDENCE_EXPORTED", {"destination": dest, "format": fmt})
    def close(self):
        return self._record("CUSTODY_CLOSED", {"total_events": len(self.events)})

    def verify(self) -> bool:
        prev = "GENESIS"
        for e in self.events:
            t = dict(e); stored = t.pop("entry_hash")
            exp = hashlib.sha256((prev+json.dumps(t,sort_keys=True,default=str)).encode()).hexdigest()
            if exp != stored: return False
            prev = stored
        return True

    def print_summary(self):
        ok = self.verify()
        st = f"{Fore.GREEN}✓ INTACT{Style.RESET_ALL}" if ok else f"{Fore.RED}✗ COMPROMISED{Style.RESET_ALL}"
        print(f"\n{Fore.CYAN}{'─'*70}{Style.RESET_ALL}")
        print(f"{Fore.CYAN}  CHAIN OF CUSTODY | Case: {self.case_id} | Integrity: {st}{Style.RESET_ALL}")
        print(f"{Fore.CYAN}{'─'*70}{Style.RESET_ALL}")
        for e in self.events:
            ds = ", ".join(f"{k}={v}" for k,v in e["details"].items()
                           if isinstance(v,(str,int,bool)) and k!="note")[:80]
            print(f"  [{e['seq']:02d}] {e['timestamp']}  {Fore.YELLOW}{e['action']:<26}{Style.RESET_ALL}  {ds}")
        print(f"{Fore.CYAN}{'─'*70}{Style.RESET_ALL}")

    def to_dict(self) -> Dict:
        return {"record_id": self.record_id, "case_id": self.case_id,
                "examiner": self.examiner, "integrity_valid": self.verify(),
                "event_count": len(self.events), "events": self.events}

#  TIMELINE RECONSTRUCTOR

class TimelineReconstructor:
    def __init__(self): self._events: List[Dict] = []

    def add(self, ts: Optional[str], source: str, description: str, category: str = "general"):
        if not ts: return
        n = self._norm(ts)
        if n: self._events.append({"timestamp": n, "source": source,
                                   "description": description, "category": category})

    @staticmethod
    def _norm(raw: str) -> Optional[str]:
        raw = str(raw).strip()
        for fmt in ["%Y:%m:%d %H:%M:%S","%Y-%m-%dT%H:%M:%S","%Y-%m-%dT%H:%M:%SZ",
                    "%Y-%m-%d %H:%M:%S","%Y-%m-%d","%d/%m/%Y %H:%M:%S"]:
            try: return datetime.strptime(raw[:len(fmt)],fmt).isoformat()
            except: continue
        return raw[:19] if re.match(r"\d{4}-\d{2}-\d{2}T\d{2}:\d{2}",raw) else None

    def build(self) -> List[Dict]:
        return sorted(self._events, key=lambda e: e["timestamp"])

    def print_timeline(self):
        events = self.build()
        if not events: print("  (No timestamps found)"); return
        print(f"\n{Fore.CYAN}{'─'*70}{Style.RESET_ALL}")
        prev = None
        for ev in events:
            d = ev["timestamp"][:10]
            if d != prev: print(f"\n  {Fore.YELLOW}▶ {d}{Style.RESET_ALL}"); prev = d
            t = ev["timestamp"][11:19] if len(ev["timestamp"])>10 else ""
            cc = {"filesystem":Fore.GREEN,"exif":Fore.CYAN,"document":Fore.BLUE,
                  "audio":Fore.MAGENTA,"executable":Fore.RED}.get(ev["category"],Fore.WHITE)
            print(f"    {t}  {cc}[{ev['category']:<11}]{Style.RESET_ALL}  {ev['source']:<22}  {ev['description']}")
        print(f"\n{Fore.CYAN}{'─'*70}{Style.RESET_ALL}\n")

    def to_dict(self) -> Dict:
        ev = self.build()
        return {"total_events": len(ev),
                "earliest": ev[0]["timestamp"] if ev else None,
                "latest":   ev[-1]["timestamp"] if ev else None,
                "events":   ev}

#  THREAT INTELLIGENCE

class ThreatIntelligence:
    RISK_BANDS = [(80,"CRITICAL"),(60,"HIGH"),(40,"MEDIUM"),(20,"LOW"),(0,"CLEAN")]

    def __init__(self, file_path: str):
        self.file_path    = Path(file_path)
        self.iocs:         List[str] = []
        self.yara_matches: List[Dict] = []
        self.risk_score   = 0
        self.risk_reasons: List[str] = []
        self._rules       = self._compile_yara()

    @staticmethod
    def _compile_yara():
        if not YARA_AVAILABLE: return None
        try: return yara.compile(source=BUILTIN_YARA_RULES)
        except: return None

    def _run_yara_streaming(self) -> List[Dict]:
        """Use mmap for zero-copy YARA scan; fall back to chunked read."""
        if not self._rules: return []
        size = self.file_path.stat().st_size
        if size > YARA_LIMIT: return []
        try:
            view, is_mmap = open_mmap_or_bytes(self.file_path, size)
            try:
                matches = self._rules.match(data=bytes(view))
            finally:
                if is_mmap: view.close()
            results = []
            for m in matches:
                sev  = m.meta.get("severity","unknown")
                bump = {"low":10,"medium":25,"high":40,"critical":60}.get(sev,15)
                results.append({"rule":m.rule,"description":m.meta.get("description",""),
                                 "severity":sev,"score_bump":bump,"tags":list(m.tags),
                                 "strings":[(hex(o),i) for o,i,_ in m.strings][:5]})
                self.risk_score += bump
                self.risk_reasons.append(f"YARA:{m.rule} (+{bump})")
            return results
        except: return []

    def _check_hashes(self, hashes):
        for algo, digest in hashes.items():
            if digest.lower() in KNOWN_MALICIOUS_HASHES:
                self.iocs.append(f"Known-malicious {algo.upper()}: {digest[:16]}…")
                self.risk_score += 100; self.risk_reasons.append(f"KnownMaliciousHash=+100")

    def _check_packers_streaming(self):
        found = scan_patterns_streaming(self.file_path, PACKER_SIGNATURES)
        for desc in found:
            self.iocs.append(f"Packer/Protector: {desc}")
            self.risk_score += 30; self.risk_reasons.append(f"Packer:{desc}=+30")

    def _check_suspicious_streaming(self):
        ext = self.file_path.suffix.lower()
        safe = [(p,d) for p,d in SUSPICIOUS_PATTERNS if not (p in (b"MZ",) and ext in (".exe",".dll"))]
        found = scan_patterns_streaming(self.file_path, safe)
        for desc in found:
            self.iocs.append(f"Suspicious pattern: {desc}")
            self.risk_score += 20; self.risk_reasons.append(f"Pattern:{desc}=+20")

    def _score_entropy(self, entropy: float):
        ext = self.file_path.suffix.lower()
        if entropy > 7.95 and ext not in COMPRESSED_EXT:
            self.risk_score += 30; self.risk_reasons.append(f"HighEntropy({entropy:.2f})=+30")
            self.iocs.append(f"Very high entropy ({entropy:.2f}) – possible encryption/packing")
        elif entropy > 7.5 and ext not in COMPRESSED_EXT:
            self.risk_score += 15; self.risk_reasons.append(f"ElevatedEntropy({entropy:.2f})=+15")

    def analyse(self, metadata: Dict) -> Dict:
        hashes  = metadata.get("hashes",{})
        forensic= metadata.get("forensic_info",{})
        ex      = metadata.get("extracted_metadata",{})
        entropy = forensic.get("entropy",0.0)
        pe_meta = ex.get("pe_file",{})
        exif    = ex.get("exif",{})

        self._check_hashes(hashes)
        if isinstance(entropy,(int,float)): self._score_entropy(float(entropy))

        for flag in forensic.get("suspicious_flags",[]):
            self.iocs.append(f"Anomaly: {flag}")
            self.risk_score += 25; self.risk_reasons.append("SuspiciousFlag=+25")

        if pe_meta:
            ts = pe_meta.get("timestamp")
            if ts:
                try:
                    pet = datetime.fromisoformat(ts)
                    if pet.year < 2000:
                        self.iocs.append(f"PE timestamp very old ({pet.year}) – tampering?")
                        self.risk_score += 20; self.risk_reasons.append("PETimestampAnomaly=+20")
                    if pet > datetime.now():
                        self.iocs.append("PE timestamp in future – anti-forensic?")
                        self.risk_score += 35; self.risk_reasons.append("PEFutureTimestamp=+35")
                except: pass

        if "GPSInfo" in exif or "GPS" in str(exif):
            self.iocs.append("GPS coordinates found – location privacy leak")
            self.risk_score += 5; self.risk_reasons.append("GPSData=+5")

        # streaming ops (no full read)
        self.yara_matches = self._run_yara_streaming()
        self._check_packers_streaming()
        self._check_suspicious_streaming()

        self.risk_score = min(100, self.risk_score)
        level = "CLEAN"
        for thr, lbl in self.RISK_BANDS:
            if self.risk_score >= thr: level = lbl; break

        return {"risk_score": self.risk_score, "risk_level": level,
                "risk_reasons": self.risk_reasons,
                "iocs": list(dict.fromkeys(self.iocs)),
                "yara_matches": self.yara_matches, "yara_available": YARA_AVAILABLE}

    def print_report(self, report: Dict):
        color = {"CRITICAL":Fore.RED,"HIGH":Fore.RED,"MEDIUM":Fore.YELLOW,
                 "LOW":Fore.CYAN,"CLEAN":Fore.GREEN}.get(report["risk_level"],"")
        score = report["risk_score"]; level = report["risk_level"]
        bar   = "█"*int(40*score/100) + "░"*(40-int(40*score/100))
        print(f"\n{Fore.CYAN}{'─'*70}{Style.RESET_ALL}")
        print(f"{Fore.CYAN}  THREAT INTELLIGENCE REPORT{Style.RESET_ALL}")
        print(f"{Fore.CYAN}{'─'*70}{Style.RESET_ALL}")
        print(f"  Risk Score : {color}{score:>3}/100  [{bar}]  {level}{Style.RESET_ALL}")
        if report["iocs"]:
            print(f"\n  {Fore.RED}IOCs:{Style.RESET_ALL}")
            for ioc in report["iocs"]: print(f"    ⚑  {ioc}")
        if report["yara_matches"]:
            print(f"\n  {Fore.RED}YARA Matches:{Style.RESET_ALL}")
            for m in report["yara_matches"]:
                print(f"    ⚡ [{m['severity'].upper()}] {m['rule']} – {m['description']}")
        print(f"{Fore.CYAN}{'─'*70}{Style.RESET_ALL}\n")

#  FORENSIC LOGGER

class ForensicLogger:
    def __init__(self, log_file=None, case_id=None):
        self.case_id = case_id or f"CASE-{datetime.now().strftime('%Y%m%d-%H%M%S')}"
        self._lock   = Lock()
        self.logger  = logging.getLogger(f"{TOOL_NAME}_{self.case_id}")
        self.logger.setLevel(logging.DEBUG)
        if not self.logger.handlers:
            h = logging.StreamHandler(); h.setLevel(logging.WARNING)
            self.logger.addHandler(h)
        if log_file:
            fh = logging.FileHandler(log_file); fh.setLevel(logging.DEBUG)
            self.logger.addHandler(fh)

    def _log(self, level, msg, **kw):
        entry = {"ts": datetime.now(timezone.utc).isoformat(), "level": level,
                 "msg": msg, "version": VERSION, **kw}
        with self._lock: getattr(self.logger, level.lower(), self.logger.debug)(json.dumps(entry))

    def info(self,msg,**kw):    self._log("INFO",   msg,**kw)
    def warning(self,msg,**kw): self._log("WARNING",msg,**kw)
    def error(self,msg,**kw):   self._log("ERROR",  msg,**kw)

#  EXIFTOOL WRAPPER

class ExifToolWrapper:
    def __init__(self):
        self.available = self._check()

    def _check(self):
        try: return subprocess.run(["exiftool","-ver"],capture_output=True,text=True,timeout=5).returncode==0
        except: return False

    def extract(self, fp: str) -> Optional[Dict]:
        if not self.available: return None
        try:
            r = subprocess.run(["exiftool","-j","-G","-a","-s",str(fp)],
                               capture_output=True,text=True,timeout=EXIFTOOL_TIMEOUT)
            if r.returncode==0 and r.stdout:
                d = json.loads(r.stdout); return d[0] if d else None
        except: return None

#  FILE TYPE DETECTOR

class FileTypeDetector:
    def __init__(self): self.available = MAGIC_AVAILABLE

    def detect(self, fp: str) -> Dict:
        result = {"extension": Path(fp).suffix.lower(),
                  "mime_type": None, "description": None, "real_type": None}
        if self.available:
            try:
                result["mime_type"]   = magic.Magic(mime=True).from_file(fp)
                result["description"] = magic.Magic().from_file(fp)
            except: pass
        result["real_type"] = self._classify(result["mime_type"], result["extension"])
        return result

    def _classify(self, mime, ext=""):
        ext = (ext or "").lower()
        if ext in [".e01",".dd",".img",".iso",".aff4"]: return "disk_image"
        if not mime: return None
        if mime.startswith("image/"): return "image"
        if mime.startswith("video/"): return "video"
        if mime.startswith("audio/"): return "audio"
        if "pdf" in mime: return "pdf"
        if "word" in mime or "document" in mime: return "document"
        if "zip" in mime or "compressed" in mime: return "archive"
        if "executable" in mime or mime=="application/x-dosexec": return "executable"
        return "unknown"

#  DISK IMAGE HANDLER

class DiskImageHandler:
    def __init__(self, image_path: str, logger=None):
        self.image_path = str(image_path); self.logger = logger
        self.img_info = None; self.partitions = []

    def open_image(self):
        if not TSK_AVAILABLE: raise ImportError("pytsk3 required")
        ext = Path(self.image_path).suffix.lower()
        try:
            if ext in [".e01",".s01"] and EWF_AVAILABLE:
                fns = pyewf.get_filenames(self.image_path)
                h   = pyewf.handle(); h.open(fns)
                self.img_info = pytsk3.Img_Info(url="",type=pytsk3.TSK_IMG_TYPE_EXTERNAL,external_handle=h)
            else:
                self.img_info = pytsk3.Img_Info(self.image_path)
            return True
        except Exception as e:
            if self.logger: self.logger.error(f"Disk open failed: {e}")
            return False

    def list_partitions(self):
        parts = []
        if not self.img_info: return parts
        try:
            vol = pytsk3.Volume_Info(self.img_info)
            for p in vol:
                parts.append({"id":p.addr,"description":p.desc.decode("utf-8",errors="ignore"),
                               "start":p.start,"length":p.len,"flags":p.flags})
        except:
            parts.append({"id":0,"description":"Single Partition","start":0,
                           "length":self.img_info.get_size(),"flags":0})
        self.partitions = parts; return parts

    def walk_filesystem(self, partition_id=0) -> List[Dict]:
        files = []
        if not self.img_info: return files
        try:
            offset = next((p["start"]*512 for p in self.partitions if p["id"]==partition_id), 0)
            fs     = pytsk3.FS_Info(self.img_info, offset=offset)
            root   = fs.open_dir(path="/")
            def _walk(d, cur=""):
                for e in d:
                    if e.info.name.name in [b".",b".."]: continue
                    name = e.info.name.name.decode("utf-8",errors="ignore")
                    full = f"{cur}/{name}"
                    m    = e.info.meta
                    files.append({"name":name,"path":full,
                                  "size":m.size if m else 0,
                                  "type":"directory" if m and m.type==pytsk3.TSK_FS_META_TYPE_DIR else "file",
                                  "inode":m.addr if m else 0,
                                  "created":datetime.fromtimestamp(m.crtime).isoformat() if m and hasattr(m,"crtime") else None,
                                  "modified":datetime.fromtimestamp(m.mtime).isoformat() if m and hasattr(m,"mtime") else None})
                    if files[-1]["type"]=="directory":
                        try: _walk(e.as_directory(), full)
                        except: pass
            _walk(root)
        except Exception as e:
            if self.logger: self.logger.error(f"Filesystem walk error: {e}")
        return files
    
#  ENHANCED METADATA EXTRACTOR

class EnhancedMetadataExtractor:
    def __init__(self, file_path: str, case_id=None, logger=None,
                 chain_of_custody=None, examiner="Unknown"):
        self.file_path  = Path(file_path)
        self.case_id    = case_id
        self.logger     = logger or ForensicLogger(case_id=case_id)
        self.exiftool   = ExifToolWrapper()
        self.detector   = FileTypeDetector()
        self.coc        = chain_of_custody
        self.examiner   = examiner
        self.timeline   = TimelineReconstructor()
        self.metadata: Dict = {
            "forensic_info":{}, "file_info":{}, "file_type_analysis":{},
            "extracted_metadata":{}, "hashes":{}, "exiftool_metadata":{},
            "timeline":{}, "threat_intel":{}, "chain_of_custody":{},
            "warnings":[], "errors":[],
        }

    def extract_all(self) -> Dict:
        try:
            if not self.file_path.exists():
                raise FileNotFoundError(f"File not found: {self.file_path}")
            self._forensic_meta(); self._file_type(); self._file_info()
            self._hashes()          # always chunked
            if self.coc: self.coc.log_acquisition(str(self.file_path), self.metadata["hashes"])
            self._exiftool_meta(); self._hachoir_meta()
            self._detect_suspicious()   # streaming entropy + tail read
            self._type_specific()
            self._build_timeline()
            self.metadata["timeline"] = self.timeline.to_dict()
            ti = ThreatIntelligence(str(self.file_path))
            ti_r = ti.analyse(self.metadata)
            self.metadata["threat_intel"] = ti_r
            if self.coc:
                self.coc.log_analysis(str(self.file_path), "Full extraction",
                    f"Risk={ti_r['risk_score']}, IOCs={len(ti_r['iocs'])}")
                self.metadata["chain_of_custody"] = self.coc.to_dict()
        except Exception as e:
            self.metadata["errors"].append(str(e))
        return self.metadata

    # forensic / file info 
    def _forensic_meta(self):
        self.metadata["forensic_info"] = {
            "case_id": self.case_id, "tool_name": TOOL_NAME, "tool_version": VERSION,
            "extraction_timestamp_utc": datetime.now(timezone.utc).isoformat(),
            "extraction_timestamp_local": datetime.now().isoformat(),
            "examiner": self.examiner,
            "examiner_system": {"platform": platform.system(),
                                "release":  platform.release(),
                                "python":   platform.python_version()},
            "capabilities": {k: v for k, v in [
                ("exiftool", self.exiftool.available), ("hachoir", HACHOIR_AVAILABLE),
                ("libmagic", self.detector.available), ("pil", PIL_AVAILABLE),
                ("pypdf2", PDF_AVAILABLE), ("python_docx", DOCX_AVAILABLE),
                ("mutagen", AUDIO_AVAILABLE), ("pefile", PE_AVAILABLE),
                ("yara", YARA_AVAILABLE)]},
        }

    def _file_type(self):
        self.metadata["file_type_analysis"] = self.detector.detect(str(self.file_path))

    def _file_info(self):
        st = self.file_path.stat()
        self.metadata["file_info"] = {
            "filename":        self.file_path.name,
            "full_path":       str(self.file_path.absolute()),
            "file_size_bytes": st.st_size,
            "file_size_human": human_size(st.st_size),
            "extension":       self.file_path.suffix,
            "created_time":    datetime.fromtimestamp(st.st_ctime).isoformat(),
            "modified_time":   datetime.fromtimestamp(st.st_mtime).isoformat(),
            "accessed_time":   datetime.fromtimestamp(st.st_atime).isoformat(),
            "permissions":     oct(st.st_mode)[-3:],
        }

    def _hashes(self):
        try:
            self.metadata["hashes"] = compute_hashes_streaming(self.file_path)
        except Exception as e:
            self.metadata["errors"].append(f"Hash error: {e}")

    def _exiftool_meta(self):
        d = self.exiftool.extract(str(self.file_path))
        if d: self.metadata["exiftool_metadata"] = d
        else: self.metadata["warnings"].append("ExifTool not available or returned no data")

    def _hachoir_meta(self):
        if not HACHOIR_AVAILABLE: return
        try:
            import contextlib, io as _io
            with contextlib.redirect_stderr(_io.StringIO()):
                parser = createParser(str(self.file_path))
                if parser:
                    with parser:
                        meta = extractMetadata(parser)
                        if meta:
                            d = {}
                            for item in meta:
                                if item.values:
                                    vals = [v.text for v in item.values]
                                    d[item.key] = vals[0] if len(vals)==1 else vals
                            if d: self.metadata["extracted_metadata"]["hachoir"] = d
        except: pass

    # streaming suspicious detection
    def _detect_suspicious(self):
        suspicious = []
        try:
            size = self.file_path.stat().st_size
            # Entropy: O(256 bytes) RAM regardless of file size
            entropy = compute_entropy_streaming(self.file_path)
            self.metadata["forensic_info"]["entropy"] = round(entropy, 4)
            ext = self.file_path.suffix.lower()
            if entropy > 7.95 and ext not in COMPRESSED_EXT:
                suspicious.append(f"Extremely high entropy ({entropy:.2f}) for {ext}")
            elif entropy < 1.0 and size > 4096:
                suspicious.append(f"Very low entropy ({entropy:.2f}) – uniform data blocks")

            # Trailing data: read only tail bytes (not whole file)
            if ext in [".jpg",".jpeg"]:
                tail  = read_file_tail(self.file_path, 16384)
                eoi   = tail.rfind(b"\xff\xd9")
                if eoi != -1 and eoi+2 < len(tail):
                    extra = tail[eoi+2:]
                    if len(extra) > 512 and not all(b in b"\x00\xff\r\n\t " for b in extra):
                        suspicious.append(f"Trailing data after JPEG EOF: {len(extra)} bytes (steganography?)")
            elif ext == ".png":
                tail = read_file_tail(self.file_path, 16384)
                iend = tail.rfind(b"IEND")
                if iend != -1:
                    end = iend+8
                    if end < len(tail):
                        extra = tail[end:]
                        if len(extra) > 512 and not all(b in b"\x00\xff\r\n\t " for b in extra):
                            suspicious.append(f"Trailing data after PNG IEND: {len(extra)} bytes (steganography?)")
        except Exception as e:
            self.metadata["warnings"].append(f"Suspicious detection error: {e}")

        if suspicious:
            self.metadata["forensic_info"]["suspicious_flags"] = suspicious
            for s in suspicious: self.metadata["warnings"].append(f"SUSPICIOUS: {s}")

    # type-specific extraction and timeline building
    def _build_timeline(self):
        fi   = self.metadata.get("file_info",{})
        self.timeline.add(fi.get("created_time"), "Filesystem","File created","filesystem")
        self.timeline.add(fi.get("modified_time"),"Filesystem","File last modified","filesystem")
        self.timeline.add(fi.get("accessed_time"),"Filesystem","File last accessed","filesystem")
        exif = self.metadata.get("extracted_metadata",{}).get("exif",{})
        self.timeline.add(exif.get("DateTimeOriginal"),"EXIF","Photo taken","exif")
        self.timeline.add(exif.get("DateTime"),        "EXIF","Image date","exif")
        docx = self.metadata.get("extracted_metadata",{}).get("docx",{})
        self.timeline.add(docx.get("created"),  "DOCX","Doc created","document")
        self.timeline.add(docx.get("modified"), "DOCX","Doc modified","document")
        pdf  = self.metadata.get("extracted_metadata",{}).get("pdf",{}).get("metadata",{})
        self.timeline.add(pdf.get("CreationDate"),"PDF","PDF created","document")
        self.timeline.add(pdf.get("ModDate"),     "PDF","PDF modified","document")
        pe   = self.metadata.get("extracted_metadata",{}).get("pe_file",{})
        self.timeline.add(pe.get("timestamp"),"PE","Binary compiled","executable")
        et   = self.metadata.get("exiftool_metadata",{})
        self.timeline.add(et.get("File:FileModifyDate"),"ExifTool","File modify (ET)","filesystem")

    def _type_specific(self):
        ft  = self.metadata["file_type_analysis"].get("real_type")
        ext = self.file_path.suffix.lower()
        if ft=="image"  or ext in [".jpg",".jpeg",".png",".gif",".bmp",".tiff",".webp"]: self._image()
        elif ft=="pdf"  or ext==".pdf":    self._pdf()
        elif ft=="document" or ext in [".docx",".doc"]: self._docx()
        elif ft=="audio"    or ext in [".mp3",".mp4",".m4a",".flac",".ogg",".wav"]: self._audio()
        elif ft=="archive"  or ext in [".zip",".jar",".apk"]: self._zip()
        elif ft=="video"    or ext in [".mp4",".avi",".mkv",".mov",".wmv"]: self._video()
        elif ft=="executable" or ext in [".exe",".dll"]: self._pe()
        elif ft=="disk_image" or ext in [".e01",".dd",".img",".iso"]: self._disk()

    def _image(self):
        if not PIL_AVAILABLE: return
        try:
            img = Image.open(self.file_path)
            self.metadata["extracted_metadata"]["image"] = {
                "format":img.format,"mode":img.mode,
                "size":f"{img.width}x{img.height}","width":img.width,"height":img.height}
            raw = (img.getexif() if hasattr(img,"getexif") else
                   (img._getexif() if hasattr(img,"_getexif") else None))
            if raw:
                exif = {}
                for tid,v in raw.items():
                    tag = TAGS.get(tid,tid)
                    if tag=="GPSInfo":
                        gps = {GPSTAGS.get(gid,gid):v[gid] for gid in v}
                        try:
                            def td(x): d,m,s=x; return float(d)+float(m)/60+float(s)/3600
                            la,lr,lo,lor = gps.get("GPSLatitude"),gps.get("GPSLatitudeRef"),\
                                           gps.get("GPSLongitude"),gps.get("GPSLongitudeRef")
                            if la and lo and lr and lor:
                                lat = td(la)*(-1 if lr=="S" else 1)
                                lng = td(lo)*(-1 if lor=="W" else 1)
                                gps["coordinates"]={"latitude":lat,"longitude":lng,
                                    "google_maps_url":f"https://www.google.com/maps?q={lat},{lng}"}
                        except: pass
                        exif[tag]=gps
                    else:
                        if isinstance(v,bytes):
                            try: v=v.decode("utf-8",errors="ignore")
                            except: v=str(v)
                        exif[tag]=v
                if exif: self.metadata["extracted_metadata"]["exif"]=exif
        except Exception as e: self.metadata["errors"].append(f"Image error: {e}")

    def _pdf(self):
        if not PDF_AVAILABLE: return
        try:
            with open(self.file_path,"rb") as f:
                r = PyPDF2.PdfReader(f)
                info = {"num_pages":len(r.pages),"is_encrypted":r.is_encrypted}
                if r.metadata: info["metadata"]={k.lstrip("/"):v for k,v in r.metadata.items()}
                self.metadata["extracted_metadata"]["pdf"]=info
        except Exception as e: self.metadata["errors"].append(f"PDF error: {e}")

    def _docx(self):
        if not DOCX_AVAILABLE: return
        try:
            doc = DocxDocument(self.file_path); cp = doc.core_properties
            self.metadata["extracted_metadata"]["docx"]={
                "author":cp.author,"title":cp.title,"subject":cp.subject,
                "created":cp.created.isoformat() if cp.created else None,
                "modified":cp.modified.isoformat() if cp.modified else None,
                "last_modified_by":cp.last_modified_by,"revision":cp.revision,
                "num_paragraphs":len(doc.paragraphs),"num_tables":len(doc.tables)}
        except Exception as e: self.metadata["errors"].append(f"DOCX error: {e}")

    def _audio(self):
        if not AUDIO_AVAILABLE: return
        try:
            af = mutagen.File(self.file_path)
            if af:
                info={"length_seconds":getattr(af.info,"length",None),
                      "bitrate":getattr(af.info,"bitrate",None),
                      "sample_rate":getattr(af.info,"sample_rate",None)}
                if af.tags: info["tags"]={k:str(v) for k,v in af.tags.items()}
                self.metadata["extracted_metadata"]["audio"]=info
        except Exception as e: self.metadata["errors"].append(f"Audio error: {e}")

    def _zip(self):
        try:
            if not zipfile.is_zipfile(self.file_path): return
            with zipfile.ZipFile(self.file_path,"r") as zf:
                info={"file_count":len(zf.namelist()),"files":[],"total_uncompressed":0,"total_compressed":0}
                for zi in zf.infolist():
                    info["files"].append({"filename":zi.filename,"compressed":zi.compress_size,
                        "uncompressed":zi.file_size,"date":datetime(*zi.date_time).isoformat(),"crc":hex(zi.CRC)})
                    info["total_uncompressed"]+=zi.file_size; info["total_compressed"]+=zi.compress_size
                info["total_uncompressed_human"]=human_size(info["total_uncompressed"])
                info["total_compressed_human"]=human_size(info["total_compressed"])
                self.metadata["extracted_metadata"]["zip"]=info
        except Exception as e: self.metadata["errors"].append(f"ZIP error: {e}")

    def _video(self):
        try:
            r=subprocess.run(["ffprobe","-v","quiet","-print_format","json",
                              "-show_format","-show_streams",str(self.file_path)],
                             capture_output=True,text=True,timeout=30)
            if r.returncode==0 and r.stdout:
                vd=json.loads(r.stdout); fmt=vd.get("format",{})
                self.metadata["extracted_metadata"]["video"]={
                    "format":fmt,"streams":vd.get("streams",[]),
                    "duration":fmt.get("duration"),"bit_rate":fmt.get("bit_rate"),
                    "format_name":fmt.get("format_name")}
        except Exception as e: self.metadata["warnings"].append(f"Video error: {e}")

    def _pe(self):
        if not PE_AVAILABLE: return
        try:
            pe=pefile.PE(str(self.file_path))
            info={"machine":hex(pe.FILE_HEADER.Machine),
                  "timestamp":datetime.fromtimestamp(pe.FILE_HEADER.TimeDateStamp).isoformat(),
                  "number_of_sections":pe.FILE_HEADER.NumberOfSections,
                  "characteristics":hex(pe.FILE_HEADER.Characteristics),
                  "sections":[{"name":s.Name.decode("utf-8",errors="ignore").strip("\x00"),
                                "virtual_address":hex(s.VirtualAddress),
                                "virtual_size":s.Misc_VirtualSize,"raw_size":s.SizeOfRawData}
                               for s in pe.sections]}
            if hasattr(pe,"DIRECTORY_ENTRY_IMPORT"):
                info["imports"]=[{"dll":e.dll.decode("utf-8",errors="ignore"),
                                   "import_count":len(e.imports)}
                                  for e in pe.DIRECTORY_ENTRY_IMPORT]
            self.metadata["extracted_metadata"]["pe_file"]=info
        except Exception as e: self.metadata["errors"].append(f"PE error: {e}")

    def _disk(self):
        if not TSK_AVAILABLE: self.metadata["warnings"].append("pytsk3 not installed"); return
        try:
            h=DiskImageHandler(str(self.file_path),self.logger)
            if h.open_image():
                parts=h.list_partitions()
                info={"total_size":h.img_info.get_size(),
                      "total_size_human":human_size(h.img_info.get_size()),
                      "partition_count":len(parts),"partitions":parts}
                if parts:
                    dp = next((p for p in parts if p.get("flags")==getattr(pytsk3,"TSK_VS_PART_FLAG_ALLOC",1)),parts[0])
                    info["sample_files"]=h.walk_filesystem(dp["id"])[:50]
                self.metadata["extracted_metadata"]["disk_image"]=info
        except Exception as e: self.metadata["errors"].append(f"Disk image error: {e}")

#  PARALLEL SCANNER

class ParallelScanner:
    def __init__(self, max_workers=DEFAULT_WORKERS, case_id=None, logger=None,
                 chain_of_custody=None, examiner="Unknown"):
        self.max_workers = max_workers; self.case_id = case_id
        self.logger      = logger or ForensicLogger(case_id=case_id)
        self.coc         = chain_of_custody; self.examiner = examiner
        self.processed   = 0; self.failed = 0; self._lock = Lock()

    def scan_directory(self, directory: Path, recursive=False) -> List[Dict]:
        pat   = "**/*" if recursive else "*"
        files = [f for f in directory.glob(pat) if f.is_file()]
        self.logger.info("Starting scan", directory=str(directory), total=len(files))
        results = []
        with tqdm(total=len(files), desc=f"{Fore.CYAN}Scanning{Style.RESET_ALL}", unit="file") as pbar:
            with ThreadPoolExecutor(max_workers=self.max_workers) as ex:
                futs = {ex.submit(self._process, fp): fp for fp in files}
                for fut in as_completed(futs):
                    try:
                        results.append(fut.result())
                        with self._lock: self.processed += 1
                    except Exception as e:
                        with self._lock: self.failed += 1; self.processed += 1
                        self.logger.error(f"Failed: {futs[fut]}: {e}")
                    finally:
                        pbar.update(1)
        return results

    def _process(self, fp: Path) -> Dict:
        ext = EnhancedMetadataExtractor(str(fp), case_id=self.case_id,
                                        logger=self.logger, chain_of_custody=self.coc,
                                        examiner=self.examiner)
        return ext.extract_all()

#  OUTPUT / REPORTING

def print_metadata(metadata: Dict, verbose=False):
    print("\n" + Fore.BLUE + "="*80 + Style.RESET_ALL)
    print(f"{Fore.GREEN}FORENSISCAN v{VERSION} — EXTRACTION REPORT{Style.RESET_ALL}")
    print(Fore.BLUE + "="*80 + Style.RESET_ALL)

    fi = metadata.get("file_info",{})
    if fi:
        print(f"\n{Fore.YELLOW}[FILE INFORMATION]{Style.RESET_ALL}")
        for k,v in fi.items(): print(f"  {k.replace('_',' ').title()}: {v}")

    h = metadata.get("hashes",{})
    if h:
        print(f"\n{Fore.YELLOW}[HASHES]{Style.RESET_ALL}")
        for k,v in h.items(): print(f"  {k.upper()}: {v}")

    ti = metadata.get("threat_intel",{})
    if ti:
        score = ti.get("risk_score",0); level = ti.get("risk_level","CLEAN")
        bar   = "█"*int(40*score/100) + "░"*(40-int(40*score/100))
        color = {"CRITICAL":Fore.RED,"HIGH":Fore.RED,"MEDIUM":Fore.YELLOW,
                 "LOW":Fore.CYAN,"CLEAN":Fore.GREEN}.get(level,"")
        print(f"\n{Fore.YELLOW}[THREAT INTELLIGENCE]{Style.RESET_ALL}")
        print(f"  Risk: {color}{score:>3}/100  [{bar}]  {level}{Style.RESET_ALL}")
        for ioc in ti.get("iocs",[]): print(f"  ⚑  {ioc}")

    tl = metadata.get("timeline",{})
    if tl.get("events"):
        print(f"\n{Fore.YELLOW}[TIMELINE — {tl['total_events']} events]{Style.RESET_ALL}")
        for ev in tl["events"][:15]:
            print(f"  {ev['timestamp'][:19]}  [{ev['category']:<11}]  {ev['source']:<18}  {ev['description']}")
        if tl["total_events"]>15: print(f"  … {tl['total_events']-15} more")

    coc = metadata.get("chain_of_custody",{})
    if coc:
        ok = coc.get("integrity_valid",False)
        st = f"{Fore.GREEN}✓ INTACT{Style.RESET_ALL}" if ok else f"{Fore.RED}✗ COMPROMISED{Style.RESET_ALL}"
        print(f"\n{Fore.YELLOW}[CHAIN OF CUSTODY]{Style.RESET_ALL}")
        print(f"  Record: {coc.get('record_id')}  |  Events: {coc.get('event_count')}  |  Integrity: {st}")

    if verbose and metadata.get("exiftool_metadata"):
        print(f"\n{Fore.YELLOW}[EXIFTOOL]{Style.RESET_ALL}")
        print(json.dumps(metadata["exiftool_metadata"],indent=2,default=str))

    if metadata.get("extracted_metadata"):
        print(f"\n{Fore.YELLOW}[EXTRACTED METADATA]{Style.RESET_ALL}")
        print(json.dumps(metadata["extracted_metadata"],indent=2,default=str))

    for w in metadata.get("warnings",[]): print(f"  {Fore.YELLOW}⚠{Style.RESET_ALL} {w}")
    for e in metadata.get("errors",[]):   print(f"  {Fore.RED}❌{Style.RESET_ALL} {e}")
    print("\n" + "="*80 + "\n")


def save_to_json(data, path):
    with open(path,"w",encoding="utf-8") as f: json.dump(data,f,indent=2,default=str)
    print(f"✓ JSON saved: {path}")

def save_to_csv(data, path):
    items = data if isinstance(data,list) else data.get("files",[data])
    rows  = []
    for item in items:
        flat = {}
        for k,v in item.get("file_info",{}).items(): flat[f"file_{k}"]=v
        for k,v in item.get("hashes",{}).items():    flat[f"hash_{k}"]=v
        ti = item.get("threat_intel",{})
        flat["risk_score"]=ti.get("risk_score",""); flat["risk_level"]=ti.get("risk_level","")
        flat["ioc_count"]=len(ti.get("iocs",[]));   flat["yara_count"]=len(ti.get("yara_matches",[]))
        tl = item.get("timeline",{})
        flat["tl_earliest"]=tl.get("earliest",""); flat["tl_latest"]=tl.get("latest","")
        flat["tl_events"]=tl.get("total_events","")
        rows.append(flat)
    if rows:
        hdrs = sorted(set().union(*rows))
        with open(path,"w",newline="",encoding="utf-8") as f:
            w=csv.DictWriter(f,fieldnames=hdrs); w.writeheader(); w.writerows(rows)
        print(f"✓ CSV saved: {path}")

def save_to_html(data, path):
    ts    = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    files = data if isinstance(data,list) else data.get("files",[data])
    rows  = ""
    for f in files:
        fi    = f.get("file_info",{})
        ti    = f.get("threat_intel",{})
        score = ti.get("risk_score",0)
        level = ti.get("risk_level","CLEAN").lower()
        color = {"critical":"#ef4444","high":"#e05252","medium":"#f59e0b",
                 "low":"#0096ff","clean":"#10b981"}.get(level,"#888")
        ioc_s = " | ".join(ti.get("iocs",[])[:3])
        rows += f"""<tr>
          <td>{fi.get("filename","")}</td>
          <td style="color:{color};font-weight:700">{score}</td>
          <td style="color:{color}">{level.upper()}</td>
          <td>{fi.get("file_size_human","")}</td>
          <td style="font-size:.75rem;color:#8b949e">{ioc_s[:80]}</td>
          <td style="font-family:monospace;font-size:.72rem">{f.get("hashes",{}).get("md5","")[:16]}…</td>
        </tr>"""
    html = f"""<!DOCTYPE html><html><head><meta charset="UTF-8">
<title>ForensiScan v{VERSION} Report</title>
<style>
body{{background:#0d1117;color:#c9d1d9;font-family:'Segoe UI',sans-serif;padding:24px}}
h1{{color:#58a6ff;font-size:1.6rem}} p{{color:#8b949e;margin-top:6px}}
table{{width:100%;border-collapse:collapse;font-size:.88rem;margin-top:20px}}
th{{background:#161b22;color:#58a6ff;padding:10px 12px;text-align:left;border-bottom:2px solid #30363d}}
td{{padding:8px 12px;border-bottom:1px solid #21262d}}
tr:hover td{{background:#161b22}}
</style></head><body>
<h1>ForensiScan v{VERSION} — Forensic Report</h1>
<p>Generated: {ts} | Files: {len(files)}</p>
<table><thead><tr><th>Filename</th><th>Risk Score</th><th>Level</th>
<th>Size</th><th>IOCs</th><th>MD5 (prefix)</th></tr></thead>
<tbody>{rows}</tbody></table>
<pre style="margin-top:32px;background:#161b22;padding:16px;border-radius:8px;
font-size:.78rem;overflow-x:auto">{json.dumps(files,indent=2,default=str)[:40000]}</pre>
</body></html>"""
    with open(path,"w",encoding="utf-8") as fh: fh.write(html)
    print(f"✓ HTML saved: {path}")

def save_output(data, fmt, path):
    if fmt=="json":  save_to_json(data,path)
    elif fmt=="html":save_to_html(data,path)
    elif fmt=="csv": save_to_csv(data,path)

#  CLI / INTERACTIVE MODE

def print_banner():
    print(f"""{Fore.CYAN}
╔══════════════════════════════════════════════════════════════════════════════╗
║  {Fore.YELLOW}FORENSISCAN  v{VERSION}{Fore.CYAN}                                               ║
║  Images · Docs · Audio · Video · PE · Archives · Disk Images                ║
║  Threat Intel · YARA · Timeline · Chain of Custody                          ║
║  v6: Streaming Memory · Timeline Graph · Multi-File Dashboard · Search      ║
║  Author: Sayan Pal  |  Collaborator: Soumit Santra                          ║
╚══════════════════════════════════════════════════════════════════════════════╝{Style.RESET_ALL}""")

def ask_examiner(): return input("Examiner name (Enter=skip): ").strip() or "Unknown"

def ask_save():
    if input("\nSave output? (y/n): ").strip().lower()!="y": return None,None
    print("1. JSON  2. HTML  3. CSV")
    c = input("Format (1-3): ").strip()
    fmt,ext = ("html",".html") if c=="2" else ("csv",".csv") if c=="3" else ("json",".json")
    out = input("Filename (Enter=auto): ").strip()
    if not out: out = f"forensic_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}{ext}"
    elif not out.endswith(ext): out+=ext
    return fmt,out

def interactive_mode():
    print_banner()
    print("\n[MAIN MENU]")
    print("="*80)
    opts = ["1. Extract metadata – single file",
            "2. Process directory (non-recursive)",
            "3. Process directory (recursive)",
            "4. Forensic mode – single file + case tracking + CoC",
            "5. Forensic mode – directory scan",
            "6. Analyse disk image (E01/DD/AFF4)",
            "7. System capabilities",
            "8. Supported file types",
            "9. About",
            "0. Exit"]
    for o in opts: print(o)
    print("="*80)
    choice = input("\nChoice (0-9): ").strip()

    if choice=="0": print("\n✓ Thank you for using ForensiScan!"); sys.exit(0)

    elif choice=="1":
        fp = input("\nFile path: ").strip().strip("'\"")
        if not fp: print("❌ No path"); return
        ext = EnhancedMetadataExtractor(fp)
        md  = ext.extract_all()
        print_metadata(md,verbose=True)
        ThreatIntelligence(fp).print_report(md["threat_intel"])
        ext.timeline.print_timeline()
        fmt,out = ask_save()
        if out: save_output(md,fmt,out)

    elif choice in ("2","3"):
        recursive = (choice=="3")
        dp = input("\nDirectory: ").strip().strip("'\"")
        if not dp or not Path(dp).is_dir(): print("❌ Invalid directory"); return
        w  = input(f"Workers (default {DEFAULT_WORKERS}): ").strip()
        scanner = ParallelScanner(max_workers=int(w) if w.isdigit() else DEFAULT_WORKERS)
        results = scanner.scan_directory(Path(dp),recursive)
        print(f"\n✓ Processed {len(results)} files  |  Failed: {scanner.failed}")
        fmt,out = ask_save()
        if out: save_output({"scan_info":{"directory":dp,"total":len(results)},"files":results},fmt,out)

    elif choice=="4":
        fp = input("\nFile path: ").strip().strip("'\"")
        if not fp: print("❌ No path"); return
        cid = input("Case ID (Enter=auto): ").strip() or f"CASE-{datetime.now().strftime('%Y%m%d-%H%M%S')}"
        examiner = ask_examiner()
        log_f = input("Log file (Enter=skip): ").strip() or None
        coc_f = input("CoC output file (Enter=skip): ").strip() or None
        logger = ForensicLogger(log_file=log_f, case_id=cid)
        coc    = ChainOfCustody(case_id=cid, examiner=examiner, output_path=coc_f)
        ext    = EnhancedMetadataExtractor(fp, case_id=cid, logger=logger,
                                           chain_of_custody=coc, examiner=examiner)
        md = ext.extract_all(); coc.close()
        print_metadata(md,verbose=True)
        ThreatIntelligence(fp).print_report(md["threat_intel"])
        ext.timeline.print_timeline(); coc.print_summary()
        fmt,out = ask_save()
        if out: coc.log_export(out,fmt); save_output(md,fmt,out)

    elif choice=="5":
        dp = input("\nDirectory: ").strip().strip("'\"")
        if not dp or not Path(dp).is_dir(): print("❌ Invalid directory"); return
        cid  = input("Case ID (Enter=auto): ").strip() or f"CASE-{datetime.now().strftime('%Y%m%d-%H%M%S')}"
        examiner = ask_examiner()
        log_f = input("Log file (Enter=skip): ").strip() or None
        coc_f = input("CoC file (Enter=skip): ").strip() or None
        rec_in= input("Recursive? (y/n, default y): ").strip().lower()
        w     = input(f"Workers (default {DEFAULT_WORKERS}): ").strip()
        logger  = ForensicLogger(log_file=log_f, case_id=cid)
        coc     = ChainOfCustody(case_id=cid, examiner=examiner, output_path=coc_f)
        scanner = ParallelScanner(max_workers=int(w) if w.isdigit() else DEFAULT_WORKERS,
                                  case_id=cid, logger=logger, chain_of_custody=coc, examiner=examiner)
        results = scanner.scan_directory(Path(dp), rec_in!="n")
        coc.close(); print(f"\n✓ Processed {len(results)} files"); coc.print_summary()
        fmt,out = ask_save()
        if out: coc.log_export(out,fmt); save_output({"files":results},fmt,out)

    elif choice=="6":
        if not TSK_AVAILABLE: print(f"{Fore.RED}❌ pytsk3 not installed{Style.RESET_ALL}"); return
        fp = input("\nDisk image path: ").strip().strip("'\"")
        if not fp or not Path(fp).exists(): print("❌ Invalid path"); return
        h = DiskImageHandler(fp)
        if not h.open_image(): print("❌ Failed to open"); return
        print(f"\n✓ Opened: {human_size(h.img_info.get_size())}")
        parts = h.list_partitions()
        for p in parts: print(f"  Partition {p['id']}: {p['description']} | {p['length']} sectors")
        pid   = input("\nPartition ID (default 0): ").strip()
        flist = h.walk_filesystem(int(pid) if pid.isdigit() else 0)
        print(f"\n✓ Found {len(flist)} items")
        for ff in flist[:20]: print(f"  {ff['type'][0].upper()} | {ff['size']:>10} B | {ff['path']}")
        fmt,out = ask_save()
        if out: save_output({"image":fp,"files":flist},fmt,out)

    elif choice=="7":
        print("\n[SYSTEM CAPABILITIES]")
        et = ExifToolWrapper()
        print(f"{'✓' if et.available else '❌'} ExifTool")
        for ok,name in [(MAGIC_AVAILABLE,"libmagic"),(YARA_AVAILABLE,"yara-python"),
                        (TSK_AVAILABLE,"pytsk3"),(EWF_AVAILABLE,"pyewf"),
                        (PIL_AVAILABLE,"Pillow"),(HACHOIR_AVAILABLE,"Hachoir"),
                        (PDF_AVAILABLE,"PyPDF2"),(DOCX_AVAILABLE,"python-docx"),
                        (AUDIO_AVAILABLE,"mutagen"),(PE_AVAILABLE,"pefile")]:
            print(f"{'✓' if ok else '❌'} {name}")
        try: subprocess.run(["ffprobe","-version"],capture_output=True,timeout=5); print("✓ ffprobe")
        except: print("❌ ffprobe")
        input("\nPress Enter to continue...")
        interactive_mode(); return

    elif choice=="8":
        print("""
[SUPPORTED FILE TYPES]
Images      : .jpg .png .gif .bmp .tiff .webp
PDF         : .pdf
Documents   : .docx .doc
Audio       : .mp3 .m4a .flac .ogg .wav
Video       : .mp4 .avi .mkv .mov (requires ffprobe)
Archives    : .zip .jar .apk
Executables : .exe .dll (PE files)
Disk Images : .e01 .dd .img .iso .aff4

All → MD5/SHA1/SHA256 (streamed), entropy (O(256) RAM),
      YARA (mmap), Timeline, Chain of Custody, Threat Score
""")
        input("Press Enter to continue..."); interactive_mode(); return

    elif choice=="9":
        print(f"""
[ABOUT  {TOOL_NAME}  v{VERSION}]
Author: Sayan Pal  |  Collaborator: Soumit Santra

v6.0 upgrades:
  • Streaming memory — entropy O(256) RAM, mmap YARA, chunked patterns
  • Timeline visualization graph in GUI (canvas, lanes, tooltips, zoom)
  • Multi-file risk dashboard (sortable table, sparklines, click-to-detail)
  • Live search + risk/category/date filters across all views
  • 5 GUI themes (Cyber Dark, Arctic, Obsidian, Matrix, Rose Gold)
""")
        input("Press Enter to continue..."); interactive_mode(); return
    else: print("❌ Invalid choice"); return

    if input("\nAnother operation? (y/n): ").strip().lower()=="y": interactive_mode()
    else: print("\n✓ Thank you for using ForensiScan!")


def main():
    if len(sys.argv)==1: interactive_mode(); return
    p = argparse.ArgumentParser(description=f"{TOOL_NAME} v{VERSION}")
    p.add_argument("input",nargs="?")
    p.add_argument("-o","--output"); p.add_argument("-f","--format",choices=["json","html","csv"])
    p.add_argument("-d","--directory",action="store_true")
    p.add_argument("-r","--recursive",action="store_true")
    p.add_argument("-v","--verbose",action="store_true")
    p.add_argument("--workers",type=int,default=DEFAULT_WORKERS)
    p.add_argument("--case-id"); p.add_argument("--examiner",default="Unknown")
    p.add_argument("--log"); p.add_argument("--coc")
    args = p.parse_args(); print_banner()
    if not args.input: print("❌ No input specified"); p.print_help(); sys.exit(1)
    cid    = args.case_id or f"CASE-{datetime.now().strftime('%Y%m%d-%H%M%S')}"
    logger = ForensicLogger(log_file=args.log, case_id=cid)
    coc    = ChainOfCustody(case_id=cid, examiner=args.examiner, output_path=args.coc)
    inp    = Path(args.input)
    if args.directory or inp.is_dir():
        scanner = ParallelScanner(max_workers=args.workers, case_id=cid,
                                  logger=logger, chain_of_custody=coc, examiner=args.examiner)
        results = scanner.scan_directory(inp, args.recursive)
        coc.close(); print(f"\n✓ Processed {len(results)} files"); coc.print_summary()
        if args.output:
            fmt = args.format or ("html" if args.output.endswith(".html") else
                                  "csv"  if args.output.endswith(".csv") else "json")
            save_output({"files":results},fmt,args.output); coc.log_export(args.output,fmt)
    else:
        try:
            ext = EnhancedMetadataExtractor(str(inp), case_id=cid, logger=logger,
                                            chain_of_custody=coc, examiner=args.examiner)
            md = ext.extract_all(); coc.close()
            print_metadata(md,verbose=args.verbose)
            ThreatIntelligence(str(inp)).print_report(md["threat_intel"])
            ext.timeline.print_timeline(); coc.print_summary()
            if args.output:
                fmt = args.format or ("html" if args.output.endswith(".html") else
                                      "csv"  if args.output.endswith(".csv") else "json")
                save_output(md,fmt,args.output); coc.log_export(args.output,fmt)
        except FileNotFoundError as e: print(f"❌ {e}"); sys.exit(1)
        except Exception as e:
            print(f"❌ {e}")
            if args.verbose: import traceback; traceback.print_exc()
            sys.exit(1)
#  GUI

try:
    import tkinter as tk
    from tkinter import ttk, filedialog, messagebox
    TK_AVAILABLE = True
except ImportError:
    TK_AVAILABLE = False


class ThemeManager:
    _widgets: list = []
    @classmethod
    def register(cls,w): cls._widgets.append(w)
    @classmethod
    def apply(cls, name):
        apply_theme(THEMES[name])
        for w in cls._widgets:
            try: w._apply_theme()
            except: pass


class AnimatedGauge(tk.Canvas):
    def __init__(self, parent, size=220, **kw):
        super().__init__(parent, width=size, height=size, bd=0, highlightthickness=0, **kw)
        self._size=size; self._score=0.0; self._target=0; self._anim=False
        self._draw(0); ThemeManager.register(self)
    def _apply_theme(self): self.configure(bg=_T["surface"]); self._draw(self._score)
    def set_score(self, s):
        self._target=s
        if not self._anim: self._anim=True; self._step()
    def _step(self):
        if abs(self._score-self._target)<0.8:
            self._score=self._target; self._anim=False; self._draw(self._score); return
        self._score+=(self._target-self._score)*0.12; self._draw(self._score); self.after(16,self._step)
    def _col(self,s):
        if s>=80: return _T["danger"]
        if s>=60: return "#e05252"
        if s>=40: return _T["warn"]
        if s>=20: return _T["accent2"]
        return _T["success"]
    def _draw(self, score):
        self.delete("all"); S=self._size; cx=cy=S//2; R=S//2-18; r2=R-14
        self.create_arc(cx-R,cy-R,cx+R,cy+R,start=220,extent=-260,style="arc",outline=_T["grid"],width=12)
        c=self._col(score); ext=-260*(score/100)
        if abs(ext)>0.5:
            self.create_arc(cx-R,cy-R,cx+R,cy+R,start=220,extent=ext,style="arc",outline=c,width=12)
            self.create_arc(cx-R+2,cy-R+2,cx+R-2,cy+R-2,start=220,extent=ext,style="arc",outline=c,width=4)
        self.create_oval(cx-r2,cy-r2,cx+r2,cy+r2,fill=_T["bg2"],outline=_T["border"],width=1)
        self.create_text(cx,cy-14,text=f"{int(score)}",font=(_T["font_main"],34,"bold"),fill=c)
        self.create_text(cx,cy+10,text="/ 100",font=(_T["font_main"],11),fill=_T["muted"])
        lvls={0:"CLEAN",20:"LOW",40:"MEDIUM",60:"HIGH",80:"CRITICAL"}; lbl="CLEAN"
        for t,n in sorted(lvls.items()):
            if score>=t: lbl=n
        self.create_text(cx,cy+28,text=lbl,font=(_T["font_main"],10,"bold"),fill=c)
        self.configure(bg=_T["surface"])


class MetricCard(tk.Frame):
    def __init__(self, parent, label, value="–", color=None, **kw):
        super().__init__(parent,bd=0,**kw); self._color=color
        self.configure(bg=_T["surface2"],highlightthickness=1,highlightbackground=_T["border"])
        tk.Label(self,text=label,font=(_T["font_ui"],8),bg=_T["surface2"],fg=_T["muted"]).pack(pady=(10,2),padx=12)
        self._val=tk.Label(self,text=str(value),font=(_T["font_main"],22,"bold"),bg=_T["surface2"],fg=color or _T["muted"])
        self._val.pack(pady=(0,10),padx=12); ThemeManager.register(self)
    def _apply_theme(self):
        self.configure(bg=_T["surface2"],highlightbackground=_T["border"])
        for w in self.winfo_children(): w.configure(bg=_T["surface2"])
        self._val.configure(fg=self._color or _T["muted"])
    def update_value(self, v, color=None):
        self._color=color or self._color; self._val.configure(text=str(v),fg=self._color or _T["text"])


class StyledText(tk.Frame):
    def __init__(self, parent, height=10, **kw):
        super().__init__(parent,bd=0); self.configure(bg=_T["surface"])
        self.text=tk.Text(self,bg=_T["surface"],fg=_T["text"],insertbackground=_T["accent"],
                          font=(_T["font_main"],9),relief="flat",bd=0,wrap="none",
                          selectbackground=_T["sel"],selectforeground=_T["text"],
                          height=height,padx=8,pady=6,**kw)
        vsb=tk.Scrollbar(self,orient="vertical",command=self.text.yview,
                         bg=_T["border"],troughcolor=_T["bg"],activebackground=_T["accent"],relief="flat",width=8)
        hsb=tk.Scrollbar(self,orient="horizontal",command=self.text.xview,
                         bg=_T["border"],troughcolor=_T["bg"],activebackground=_T["accent"],relief="flat",width=8)
        self.text.configure(yscrollcommand=vsb.set,xscrollcommand=hsb.set)
        self.text.grid(row=0,column=0,sticky="nsew"); vsb.grid(row=0,column=1,sticky="ns")
        hsb.grid(row=1,column=0,sticky="ew"); self.rowconfigure(0,weight=1); self.columnconfigure(0,weight=1)
        ThemeManager.register(self)
    def _apply_theme(self):
        self.configure(bg=_T["surface"])
        self.text.configure(bg=_T["surface"],fg=_T["text"],selectbackground=_T["sel"])
    def set_text(self,c):
        self.text.configure(state="normal"); self.text.delete("1.0","end")
        self.text.insert("end",c); self.text.configure(state="disabled")
    def append(self,c,tag=""):
        self.text.configure(state="normal"); self.text.insert("end",c,tag)
        self.text.see("end"); self.text.configure(state="disabled")
    def clear(self):
        self.text.configure(state="normal"); self.text.delete("1.0","end")
        self.text.configure(state="disabled")
    def add_tag(self,name,**kw): self.text.tag_configure(name,**kw)


def panel(parent, title=None):
    f=tk.Frame(parent,bg=_T["surface"],highlightthickness=1,highlightbackground=_T["border"])
    if title:
        h=tk.Frame(f,bg=_T["surface2"]); h.pack(fill="x")
        tk.Label(h,text=title,font=(_T["font_ui"],9,"bold"),bg=_T["surface2"],fg=_T["accent"],padx=14,pady=7).pack(side="left")
        tk.Frame(f,bg=_T["border"],height=1).pack(fill="x")
    return f


#  Dashboard Tab
class DashboardTab(tk.Frame):
    def __init__(self, parent, app):
        super().__init__(parent,bg=_T["bg"]); self.app=app; self._build(); ThemeManager.register(self)
    def _apply_theme(self): self.configure(bg=_T["bg"])

    def _build(self):
        row=tk.Frame(self,bg=_T["bg"]); row.pack(fill="x",padx=16,pady=(16,8))
        self._c_risk=MetricCard(row,"RISK SCORE","–",_T["muted"])
        self._c_ioc =MetricCard(row,"IOCs","–",_T["muted"])
        self._c_yara=MetricCard(row,"YARA HITS","–",_T["muted"])
        self._c_tl  =MetricCard(row,"TIMELINE","–",_T["muted"])
        self._c_hash=MetricCard(row,"HASH STATUS","–",_T["muted"])
        self._c_size=MetricCard(row,"FILE SIZE","–",_T["muted"])
        for c in [self._c_risk,self._c_ioc,self._c_yara,self._c_tl,self._c_hash,self._c_size]:
            c.pack(side="left",fill="both",expand=True,padx=4)

        mid=tk.Frame(self,bg=_T["bg"]); mid.pack(fill="both",expand=True,padx=16,pady=8)
        gp=panel(mid,"  THREAT GAUGE"); gp.pack(side="left",fill="both",expand=False,ipadx=10,padx=(0,8))
        self._gauge=AnimatedGauge(gp,size=220,bg=_T["surface"]); self._gauge.pack(padx=20,pady=(8,16))
        ip=panel(mid,"  FILE INFORMATION"); ip.pack(side="left",fill="both",expand=True)
        self._info=StyledText(ip,height=10)
        self._info.add_tag("key",foreground=_T["accent"]); self._info.add_tag("val",foreground=_T["text"])
        self._info.add_tag("hash",foreground=_T["success"]); self._info.add_tag("dim",foreground=_T["muted"])
        self._info.pack(fill="both",expand=True,padx=6,pady=(0,8))

        ioc_p=panel(self,"  INDICATORS OF COMPROMISE"); ioc_p.pack(fill="both",expand=True,padx=16,pady=(0,16))
        self._ioc=StyledText(ioc_p,height=7)
        self._ioc.add_tag("ioc",foreground=_T["danger"]); self._ioc.add_tag("ok",foreground=_T["success"])
        self._ioc.add_tag("yara",foreground=_T["warn"]); self._ioc.add_tag("hdr",foreground=_T["accent"],font=(_T["font_main"],9,"bold"))
        self._ioc.pack(fill="both",expand=True,padx=6,pady=(0,8))

    def populate(self, md: dict):
        ti=md.get("threat_intel",{}); fi=md.get("file_info",{})
        tl=md.get("timeline",{}); h=md.get("hashes",{})
        finfo=md.get("forensic_info",{}); ftype=md.get("file_type_analysis",{})
        score=ti.get("risk_score",0); level=ti.get("risk_level","CLEAN")
        iocs=ti.get("iocs",[]); yara=ti.get("yara_matches",[])
        rc={"CRITICAL":_T["danger"],"HIGH":_T["danger"],"MEDIUM":_T["warn"],
            "LOW":_T["accent2"],"CLEAN":_T["success"]}.get(level,_T["text"])
        self._c_risk.update_value(score,rc)
        self._c_ioc.update_value(len(iocs),_T["danger"] if iocs else _T["success"])
        self._c_yara.update_value(len(yara),_T["warn"] if yara else _T["success"])
        self._c_tl.update_value(tl.get("total_events",0),_T["accent2"])
        self._c_hash.update_value("OK" if h.get("md5") else "–",_T["success"] if h.get("md5") else _T["muted"])
        self._c_size.update_value(fi.get("file_size_human","–"),_T["accent"])
        self._gauge.set_score(score)
        self._info.clear()
        for k,v in [("Filename  ",fi.get("filename","")),("Path      ",fi.get("full_path","")),
                    ("Extension ",fi.get("extension","")),("MIME Type ",ftype.get("mime_type","–")),
                    ("Size      ",fi.get("file_size_human","")),("Created   ",fi.get("created_time","")),
                    ("Modified  ",fi.get("modified_time","")),("Perms     ",fi.get("permissions","")),
                    ("Entropy   ",str(finfo.get("entropy","N/A")))]:
            self._info.append(f"  {k}: ","key"); self._info.append(f"{v}\n","val")
        if h:
            self._info.append("\n  ──── Hashes ────\n","dim")
            for algo,digest in h.items():
                self._info.append(f"  {algo.upper():<6}: ","key"); self._info.append(f"{digest}\n","hash")
        self._ioc.clear()
        if iocs:
            self._ioc.append("  INDICATORS OF COMPROMISE\n\n","hdr")
            for ioc in iocs: self._ioc.append(f"  ⚑  {ioc}\n","ioc")
        else:
            self._ioc.append("  ✓  No IOCs detected — file appears clean.\n","ok")
        if yara:
            self._ioc.append("\n  YARA MATCHES\n\n","hdr")
            for m in yara: self._ioc.append(f"  ⚡  [{m['severity'].upper()}] {m['rule']} — {m['description']}\n","yara")


# Timeline Graph Tab
class TimelineGraphTab(tk.Frame):
    """
    Interactive timeline visualizer with:
    • Vertical lane per category (colour-coded)
    • Dots scaled by category with hover tooltip
    • Horizontal zoom (slider)
    • Click-to-detail side panel
    • Density heatband at top
    """
    CAT_COLORS = {"filesystem":"#10b981","exif":"#0096ff","document":"#a78bfa",
                  "executable":"#ef4444","audio":"#f59e0b","general":"#8b949e"}

    def __init__(self, parent, app):
        super().__init__(parent,bg=_T["bg"]); self.app=app
        self._events=[]; self._zoom=1.0; self._tooltip=None
        self._build(); ThemeManager.register(self)

    def _apply_theme(self): self.configure(bg=_T["bg"])

    def _build(self):
        # ── top toolbar
        tb=tk.Frame(self,bg=_T["bg"]); tb.pack(fill="x",padx=16,pady=(12,4))
        tk.Label(tb,text="ACTIVITY TIMELINE",font=(_T["font_ui"],13,"bold"),bg=_T["bg"],fg=_T["accent"]).pack(side="left")

        # Zoom controls
        tk.Label(tb,text="Zoom:",font=(_T["font_ui"],9),bg=_T["bg"],fg=_T["muted"]).pack(side="right",padx=(8,2))
        self._zoom_var=tk.DoubleVar(value=1.0)
        zsl=tk.Scale(tb,from_=0.3,to=5.0,resolution=0.1,orient="horizontal",
                     variable=self._zoom_var,command=self._on_zoom,
                     bg=_T["bg"],fg=_T["text"],troughcolor=_T["border"],
                     activebackground=_T["accent"],highlightthickness=0,bd=0,length=160)
        zsl.pack(side="right")

        # ── legend
        leg=tk.Frame(self,bg=_T["bg"]); leg.pack(fill="x",padx=16,pady=(0,4))
        for cat,col in self.CAT_COLORS.items():
            f=tk.Frame(leg,bg=_T["bg"]); f.pack(side="left",padx=8)
            tk.Frame(f,bg=col,width=12,height=12).pack(side="left",padx=(0,4))
            tk.Label(f,text=cat,font=(_T["font_main"],8),bg=_T["bg"],fg=col).pack(side="left")

        # ── main paned view: graph + detail
        pw=tk.PanedWindow(self,orient="horizontal",sashwidth=6,bg=_T["border"],bd=0)
        pw.pack(fill="both",expand=True,padx=16,pady=(0,16))

        # Graph canvas with scrollbar
        graph_f=tk.Frame(pw,bg=_T["surface"]); pw.add(graph_f,minsize=500)
        self._canvas=tk.Canvas(graph_f,bg=_T["surface"],bd=0,highlightthickness=0)
        hsc=tk.Scrollbar(graph_f,orient="horizontal",command=self._canvas.xview,
                         bg=_T["border"],troughcolor=_T["bg"],activebackground=_T["accent"],relief="flat",width=8)
        vsc=tk.Scrollbar(graph_f,orient="vertical",command=self._canvas.yview,
                         bg=_T["border"],troughcolor=_T["bg"],activebackground=_T["accent"],relief="flat",width=8)
        self._canvas.configure(xscrollcommand=hsc.set,yscrollcommand=vsc.set)
        self._canvas.grid(row=0,column=0,sticky="nsew")
        vsc.grid(row=0,column=1,sticky="ns"); hsc.grid(row=1,column=0,sticky="ew")
        graph_f.rowconfigure(0,weight=1); graph_f.columnconfigure(0,weight=1)
        self._canvas.bind("<Configure>",lambda e: self._redraw())
        self._canvas.bind("<Motion>",self._on_hover)
        self._canvas.bind("<Leave>",self._hide_tooltip)
        self._canvas.bind("<Button-1>",self._on_click)

        # Detail panel
        det_f=panel(pw,"  EVENT DETAIL"); pw.add(det_f,minsize=240)
        self._detail=StyledText(det_f,height=20)
        self._detail.add_tag("key",foreground=_T["accent"]); self._detail.add_tag("val",foreground=_T["text"])
        self._detail.pack(fill="both",expand=True,padx=6,pady=(0,8))

        # Summary bar
        self._summary=tk.Label(self,text="No events loaded.",font=(_T["font_main"],8),
                                bg=_T["bg"],fg=_T["muted"])
        self._summary.pack(pady=(0,4))

        self._dot_map: List[Tuple[int,int,int,int,Dict]] = []   # x1,y1,x2,y2,event

    def _on_zoom(self, _=None):
        self._zoom = self._zoom_var.get()
        self._redraw()

    def _category_lane(self, cats: List[str], cat: str) -> int:
        return cats.index(cat) if cat in cats else 0

    def _redraw(self):
        c=self._canvas; c.delete("all"); self._dot_map=[]
        if not self._events:
            W=max(c.winfo_width(),600)
            c.create_text(W//2,80,text="No timeline events",fill=_T["muted"],font=(_T["font_main"],12))
            return

        cats  = sorted(set(e["category"] for e in self._events))
        n_cat = max(len(cats),1)
        LANE_H= 80
        HEADER= 60   # density heatband + date labels at top
        DOT_R = 7
        PAD_L = 90   # left padding for category labels

        # time extents
        ts_list= [datetime.fromisoformat(e["timestamp"]) for e in self._events]
        t_min  = min(ts_list); t_max=max(ts_list)
        span   = max((t_max-t_min).total_seconds(), 1)

        W      = max(c.winfo_width()-20, 800)
        PLOT_W = max(int((W - PAD_L) * self._zoom), W - PAD_L)
        H      = HEADER + n_cat * LANE_H + 40

        c.configure(scrollregion=(0,0, PAD_L+PLOT_W, H))

        # density heatband 
        N_BANDS = 60
        counts  = [0]*N_BANDS
        for ts in ts_list:
            idx = int((ts-t_min).total_seconds()/span*(N_BANDS-1))
            counts[idx] = min(counts[idx]+1, 20)
        bw = PLOT_W / N_BANDS
        for i,cnt in enumerate(counts):
            if cnt==0: continue
            alpha = cnt/20
            def lerp_hex(a,b,t):
                def h2r(h): return tuple(int(h.lstrip("#")[j:j+2],16) for j in (0,2,4))
                r1,g1,b1=h2r(a); r2,g2,b2=h2r(b); t=max(0,min(1,t))
                return f"#{int(r1+(r2-r1)*t):02x}{int(g1+(g2-g1)*t):02x}{int(b1+(b2-b1)*t):02x}"
            fill=lerp_hex(_T["surface"],_T["accent"],alpha)
            x1=PAD_L+i*bw; x2=x1+bw
            c.create_rectangle(x1,4,x2,28,fill=fill,outline="")

        c.create_text(PAD_L//2,16,text="Density",font=(_T["font_main"],7),fill=_T["muted"])

        # date labels 
        n_ticks=min(10, int(PLOT_W//80))
        for i in range(n_ticks+1):
            t_frac = i/n_ticks
            ts     = t_min + timedelta(seconds=span*t_frac)
            x      = PAD_L + int(PLOT_W*t_frac)
            c.create_line(x,32,x,H-20,fill=_T["border"],dash=(3,5))
            lbl = ts.strftime("%m-%d %H:%M") if span>86400 else ts.strftime("%H:%M:%S")
            c.create_text(x,44,text=lbl,font=(_T["font_main"],7),fill=_T["text_dim"],angle=30,anchor="w")

        # category lanes 
        for li,cat in enumerate(cats):
            y_center = HEADER + li*LANE_H + LANE_H//2
            col = self.CAT_COLORS.get(cat, _T["muted"])
            # lane stripe
            c.create_rectangle(PAD_L,HEADER+li*LANE_H, PAD_L+PLOT_W, HEADER+(li+1)*LANE_H,
                                fill=_T["grid"],outline="")
            # lane label
            c.create_rectangle(0,HEADER+li*LANE_H, PAD_L-4, HEADER+(li+1)*LANE_H,
                                fill=_T["surface2"],outline=_T["border"])
            c.create_text(PAD_L//2,y_center,text=cat[:10],font=(_T["font_main"],8),fill=col)
            # baseline
            c.create_line(PAD_L,y_center, PAD_L+PLOT_W,y_center, fill=_T["border"],width=1)

        # ── events 
        for ev in self._events:
            ts  = datetime.fromisoformat(ev["timestamp"])
            li  = cats.index(ev["category"]) if ev["category"] in cats else 0
            col = self.CAT_COLORS.get(ev["category"],_T["muted"])
            t_frac  = (ts-t_min).total_seconds()/span
            x       = PAD_L + int(PLOT_W*t_frac)
            y_center= HEADER + li*LANE_H + LANE_H//2
            # vertical drop line
            c.create_line(x, HEADER+li*LANE_H+4, x, y_center-DOT_R, fill=col, width=1, dash=(2,3))
            # glow ring
            c.create_oval(x-DOT_R-3,y_center-DOT_R-3,x+DOT_R+3,y_center+DOT_R+3,
                          fill="",outline=col,width=1)
            # dot
            tag=f"dot_{id(ev)}"
            c.create_oval(x-DOT_R,y_center-DOT_R,x+DOT_R,y_center+DOT_R,
                          fill=col,outline=_T["bg"],width=2,tags=(tag,))
            self._dot_map.append((x-DOT_R-4,y_center-DOT_R-4,x+DOT_R+4,y_center+DOT_R+4,ev))

        self._summary.configure(
            text=f"  {len(self._events)} events  |  "
                 f"Earliest: {self._events[0]['timestamp'][:19]}  |  "
                 f"Latest: {self._events[-1]['timestamp'][:19]}  |  "
                 f"Zoom: {self._zoom:.1f}x")

    def _find_event(self, cx, cy) -> Optional[Dict]:
        for x1,y1,x2,y2,ev in self._dot_map:
            if x1<=cx<=x2 and y1<=cy<=y2: return ev
        return None

    def _on_hover(self, event):
        cx=self._canvas.canvasx(event.x); cy=self._canvas.canvasy(event.y)
        ev=self._find_event(cx,cy)
        if ev: self._show_tooltip(event.x_root,event.y_root,ev)
        else:  self._hide_tooltip()

    def _show_tooltip(self, rx, ry, ev):
        self._hide_tooltip()
        tip=tk.Toplevel(self); tip.wm_overrideredirect(True)
        tip.wm_geometry(f"+{rx+12}+{ry-28}")
        tip.configure(bg=_T["surface2"])
        tk.Label(tip,text=f"{ev['timestamp'][:19]}\n{ev['source']} · {ev['description']}",
                 font=(_T["font_main"],8),bg=_T["surface2"],fg=_T["text"],padx=8,pady=4,
                 justify="left").pack()
        self._tooltip=tip

    def _hide_tooltip(self, _=None):
        if self._tooltip:
            try: self._tooltip.destroy()
            except: pass
            self._tooltip=None

    def _on_click(self, event):
        cx=self._canvas.canvasx(event.x); cy=self._canvas.canvasy(event.y)
        ev=self._find_event(cx,cy)
        if not ev: return
        self._detail.clear()
        for k,v in [("Timestamp  ",ev["timestamp"]),("Category   ",ev["category"]),
                    ("Source     ",ev["source"]),("Description",ev["description"])]:
            self._detail.append(f"  {k}: ","key"); self._detail.append(f"{v}\n","val")

    def populate(self, md: dict):
        tl=md.get("timeline",{}); self._events=tl.get("events",[])
        self._redraw()


# ── Multi-File Dashboard Tab  (NEW in v6) ─────────────────────────────────────
class MultiFileDashTab(tk.Frame):
    """
    Sortable risk table for batch directory scans.
    • Per-column colour coding (risk level → colour)
    • Mini bar sparklines in the Risk column
    • Click row → populates single-file tabs
    • Live search + risk-level filter + category filter
    """
    COLS = ("Filename","Risk","Level","Size","IOCs","YARA","Entropy","Type","Modified")

    def __init__(self, parent, app):
        super().__init__(parent,bg=_T["bg"]); self.app=app
        self._all_results: List[Dict]=[]
        self._filtered:   List[Dict]=[]
        self._sort_col    = "Risk"
        self._sort_rev    = True
        self._build(); ThemeManager.register(self)

    def _apply_theme(self): self.configure(bg=_T["bg"])

    def _build(self):
        # ── filter bar
        fb=tk.Frame(self,bg=_T["bg"]); fb.pack(fill="x",padx=16,pady=(12,6))
        tk.Label(fb,text="MULTI-FILE RISK DASHBOARD",font=(_T["font_ui"],13,"bold"),
                 bg=_T["bg"],fg=_T["accent"]).pack(side="left")

        # Right controls
        right=tk.Frame(fb,bg=_T["bg"]); right.pack(side="right")

        tk.Label(right,text="Search:",font=(_T["font_ui"],9),bg=_T["bg"],fg=_T["muted"]).pack(side="left",padx=(0,4))
        self._search_var=tk.StringVar()
        self._search_var.trace_add("write",lambda *_: self._apply_filter())
        tk.Entry(right,textvariable=self._search_var,width=18,font=(_T["font_main"],9),
                 bg=_T["grid"],fg=_T["text"],insertbackground=_T["accent"],relief="flat",bd=4
                 ).pack(side="left",padx=(0,10))

        tk.Label(right,text="Risk:",font=(_T["font_ui"],9),bg=_T["bg"],fg=_T["muted"]).pack(side="left",padx=(0,4))
        self._risk_var=tk.StringVar(value="ALL")
        risk_combo=ttk.Combobox(right,textvariable=self._risk_var,
                                values=["ALL","CRITICAL","HIGH","MEDIUM","LOW","CLEAN"],
                                state="readonly",width=9,font=(_T["font_main"],9))
        risk_combo.pack(side="left",padx=(0,10))
        risk_combo.bind("<<ComboboxSelected>>",lambda _: self._apply_filter())

        tk.Label(right,text="Type:",font=(_T["font_ui"],9),bg=_T["bg"],fg=_T["muted"]).pack(side="left",padx=(0,4))
        self._type_var=tk.StringVar(value="ALL")
        self._type_combo=ttk.Combobox(right,textvariable=self._type_var,
                                      values=["ALL"],state="readonly",width=12,font=(_T["font_main"],9))
        self._type_combo.pack(side="left",padx=(0,6))
        self._type_combo.bind("<<ComboboxSelected>>",lambda _: self._apply_filter())

        # ── summary stats row
        stats=tk.Frame(self,bg=_T["bg"]); stats.pack(fill="x",padx=16,pady=(0,6))
        self._stat_lbl=tk.Label(stats,text="Load a directory scan to populate this view.",
                                font=(_T["font_main"],9),bg=_T["bg"],fg=_T["muted"])
        self._stat_lbl.pack(side="left")

        # ── treeview
        tv_p=panel(self,None); tv_p.pack(fill="both",expand=True,padx=16,pady=(0,6))
        style=ttk.Style(); style.theme_use("default")
        style.configure("MF.Treeview",background=_T["surface"],foreground=_T["text"],
                        fieldbackground=_T["surface"],rowheight=26,font=(_T["font_main"],9))
        style.configure("MF.Treeview.Heading",background=_T["surface2"],foreground=_T["accent"],
                        font=(_T["font_main"],8,"bold"),relief="flat")
        style.map("MF.Treeview",background=[("selected",_T["sel"])],
                  foreground=[("selected",_T["text"])])
        self._tree=ttk.Treeview(tv_p,columns=self.COLS,show="headings",
                                style="MF.Treeview",height=16)
        widths={"Filename":200,"Risk":80,"Level":80,"Size":80,"IOCs":50,
                "YARA":50,"Entropy":70,"Type":90,"Modified":150}
        for col in self.COLS:
            self._tree.heading(col,text=col,command=lambda c=col: self._sort(c))
            self._tree.column(col,width=widths.get(col,100),anchor="w")
        vsb=ttk.Scrollbar(tv_p,orient="vertical",command=self._tree.yview)
        hsb=ttk.Scrollbar(tv_p,orient="horizontal",command=self._tree.xview)
        self._tree.configure(yscrollcommand=vsb.set,xscrollcommand=hsb.set)
        self._tree.grid(row=0,column=0,sticky="nsew"); vsb.grid(row=0,column=1,sticky="ns")
        hsb.grid(row=1,column=0,sticky="ew"); tv_p.rowconfigure(0,weight=1); tv_p.columnconfigure(0,weight=1)
        self._tree.bind("<<TreeviewSelect>>",self._on_select)

        # ── detail panel below
        det=panel(self,"  SELECTED FILE — QUICK DETAIL"); det.pack(fill="x",padx=16,pady=(0,16))
        self._det_text=StyledText(det,height=6)
        self._det_text.add_tag("k",foreground=_T["accent"]); self._det_text.add_tag("v",foreground=_T["text"])
        self._det_text.add_tag("ioc",foreground=_T["danger"]); self._det_text.add_tag("ok",foreground=_T["success"])
        self._det_text.pack(fill="x",padx=6,pady=(0,8))

        # tag colours for risk rows
        for tag,col in [("CRITICAL",_T["danger"]),("HIGH","#e05252"),
                        ("MEDIUM",_T["warn"]),("LOW",_T["accent2"]),("CLEAN",_T["success"])]:
            self._tree.tag_configure(tag,foreground=col)

    def _row_data(self, r: Dict) -> tuple:
        fi   = r.get("file_info",{})
        ti   = r.get("threat_intel",{})
        forg = r.get("forensic_info",{})
        ftyp = r.get("file_type_analysis",{})
        return (fi.get("filename",""),
                ti.get("risk_score",0),
                ti.get("risk_level","CLEAN"),
                fi.get("file_size_human",""),
                len(ti.get("iocs",[])),
                len(ti.get("yara_matches",[])),
                round(forg.get("entropy",0.0) or 0.0, 2),
                ftyp.get("real_type","–") or "–",
                fi.get("modified_time","")[:19])

    def _apply_filter(self):
        search   = self._search_var.get().lower()
        risk_f   = self._risk_var.get()
        type_f   = self._type_var.get()
        self._filtered = [r for r in self._all_results
                          if (not search or search in r.get("file_info",{}).get("filename","").lower()
                                        or search in r.get("file_info",{}).get("full_path","").lower())
                          and (risk_f=="ALL" or r.get("threat_intel",{}).get("risk_level","")==risk_f)
                          and (type_f=="ALL" or (r.get("file_type_analysis",{}).get("real_type","") or "")==type_f)]
        self._render()

    def _sort(self, col: str):
        idx = self.COLS.index(col)
        self._sort_rev = not self._sort_rev if self._sort_col==col else True
        self._sort_col = col
        self._filtered.sort(key=lambda r: self._row_data(r)[idx], reverse=self._sort_rev)
        self._render()

    def _render(self):
        for row in self._tree.get_children(): self._tree.delete(row)
        for r in self._filtered:
            vals = self._row_data(r)
            lvl  = str(vals[2])
            self._tree.insert("","end",values=vals,tags=(lvl,))
        total  = len(self._all_results); shown  = len(self._filtered)
        crits  = sum(1 for r in self._filtered if r.get("threat_intel",{}).get("risk_level")=="CRITICAL")
        highs  = sum(1 for r in self._filtered if r.get("threat_intel",{}).get("risk_level")=="HIGH")
        avg_s  = (sum(r.get("threat_intel",{}).get("risk_score",0) for r in self._filtered)//max(shown,1))
        self._stat_lbl.configure(
            text=f"Showing {shown}/{total}  |  Critical: {crits}  High: {highs}  |  Avg Risk: {avg_s}/100")

    def _on_select(self, _):
        sel=self._tree.selection()
        if not sel: return
        vals=self._tree.item(sel[0])["values"]
        fname=str(vals[0])
        r=next((x for x in self._filtered if x.get("file_info",{}).get("filename","")==fname), None)
        if not r: return
        self._det_text.clear()
        ti=r.get("threat_intel",{}); fi=r.get("file_info",{})
        score=ti.get("risk_score",0); level=ti.get("risk_level","CLEAN")
        self._det_text.append(f"  File      : ","k"); self._det_text.append(f"{fi.get('full_path','')}\n","v")
        self._det_text.append(f"  Risk      : ","k"); self._det_text.append(f"{score}/100 ({level})\n","v")
        self._det_text.append(f"  MD5       : ","k"); self._det_text.append(f"{r.get('hashes',{}).get('md5','–')}\n","v")
        self._det_text.append(f"  SHA256    : ","k"); self._det_text.append(f"{r.get('hashes',{}).get('sha256','–')}\n","v")
        iocs=ti.get("iocs",[])
        if iocs:
            self._det_text.append(f"\n  IOCs ({len(iocs)}):\n","k")
            for ioc in iocs[:5]: self._det_text.append(f"    ⚑ {ioc}\n","ioc")
        else:
            self._det_text.append("\n  ✓  No IOCs detected.\n","ok")
        # bubble click-to-detail to main app
        if hasattr(self.app,"_populate_single_file"): self.app._populate_single_file(r)

    def load_results(self, results: List[Dict]):
        self._all_results = results
        # update type filter choices
        types = sorted(set((r.get("file_type_analysis",{}).get("real_type","") or "unknown")
                           for r in results))
        self._type_combo.configure(values=["ALL"]+types)
        self._apply_filter()


# YARA Tab
class YaraTab(tk.Frame):
    def __init__(self, parent, app):
        super().__init__(parent,bg=_T["bg"]); self.app=app; self._matches=[]; self._build(); ThemeManager.register(self)
    def _apply_theme(self): self.configure(bg=_T["bg"])
    def _build(self):
        tk.Label(self,text="YARA / MALWARE ANALYSIS",font=(_T["font_ui"],13,"bold"),
                 bg=_T["bg"],fg=_T["accent"]).pack(padx=16,pady=(14,8),anchor="w")
        sp=panel(self,None); sp.pack(fill="x",padx=16,pady=(0,10))
        inner=tk.Frame(sp,bg=_T["surface"]); inner.pack(fill="x",padx=14,pady=10)
        self._stat=tk.Label(inner,text="Awaiting scan…",font=(_T["font_ui"],10,"bold"),bg=_T["surface"],fg=_T["muted"])
        self._stat.pack(side="left")
        tk.Label(inner,text=f"YARA: {'✓ ACTIVE' if YARA_AVAILABLE else '✗ pip install yara-python'}",
                 font=(_T["font_main"],8),bg=_T["surface"],fg=_T["success"] if YARA_AVAILABLE else _T["danger"]).pack(side="right")
        tp=panel(self,"  RULE MATCHES"); tp.pack(fill="both",expand=True,padx=16,pady=(0,10))
        cols=("Rule","Severity","Description","Score Bump")
        self._tree=ttk.Treeview(tp,columns=cols,show="headings",height=8)
        for col in cols: self._tree.heading(col,text=col); self._tree.column(col,width={"Rule":140,"Severity":80,"Description":300,"Score Bump":80}[col])
        vsb=ttk.Scrollbar(tp,orient="vertical",command=self._tree.yview)
        self._tree.configure(yscrollcommand=vsb.set)
        self._tree.pack(side="left",fill="both",expand=True,padx=6,pady=(0,8))
        vsb.pack(side="right",fill="y",pady=(0,8))
        self._tree.bind("<<TreeviewSelect>>",self._on_sel)
        dp=panel(self,"  MATCH DETAIL"); dp.pack(fill="x",padx=16,pady=(0,16))
        self._det=StyledText(dp,height=7)
        self._det.add_tag("k",foreground=_T["accent"]); self._det.add_tag("v",foreground=_T["text"])
        self._det.add_tag("off",foreground=_T["warn"])
        self._det.pack(fill="x",padx=6,pady=(0,8))
    def _on_sel(self,_):
        sel=self._tree.selection()
        if not sel: return
        rule=self._tree.item(sel[0])["values"][0]
        m=next((x for x in self._matches if x["rule"]==rule),None)
        if not m: return
        self._det.clear()
        for k,v in [("Rule",m["rule"]),("Severity",m["severity"].upper()),
                    ("Description",m["description"]),("Score Bump",f"+{m['score_bump']}")]:
            self._det.append(f"  {k:<14}: ","k"); self._det.append(f"{v}\n","v")
        if m.get("strings"):
            self._det.append("  Matches     :\n","k")
            for off,ident in m["strings"]: self._det.append(f"    @ {off}  {ident}\n","off")
    def populate(self,md:dict):
        ti=md.get("threat_intel",{}); self._matches=ti.get("yara_matches",[])
        self._stat.configure(
            text=(f"⚡ {len(self._matches)} YARA rule(s) matched!" if self._matches
                  else "✓ No YARA rules matched."),
            fg=(_T["danger"] if self._matches else _T["success"]))
        for r in self._tree.get_children(): self._tree.delete(r)
        for m in sorted(self._matches,key=lambda x:{"critical":0,"high":1,"medium":2,"low":3}.get(x.get("severity",""),9)):
            self._tree.insert("","end",values=(m["rule"],m["severity"].upper(),m["description"],f"+{m['score_bump']}"))
        self._det.clear()


# Risk Heatmap Tab
class HeatmapTab(tk.Frame):
    def __init__(self,parent,app):
        super().__init__(parent,bg=_T["bg"]); self.app=app; self._build(); ThemeManager.register(self)
    def _apply_theme(self): self.configure(bg=_T["bg"])
    def _build(self):
        tk.Label(self,text="RISK HEATMAP & SCORE BREAKDOWN",font=(_T["font_ui"],13,"bold"),
                 bg=_T["bg"],fg=_T["accent"]).pack(padx=16,pady=(14,8),anchor="w")
        row=tk.Frame(self,bg=_T["bg"]); row.pack(fill="both",expand=True,padx=16,pady=(0,16))
        left=panel(row,"  SCORE CONTRIBUTORS"); left.pack(side="left",fill="both",expand=True,padx=(0,8))
        self._bar=tk.Canvas(left,bg=_T["surface"],bd=0,highlightthickness=0)
        self._bar.pack(fill="both",expand=True,padx=10,pady=(0,10))
        self._bar.bind("<Configure>",lambda e: self._draw_bars([]))
        right=panel(row,"  INDICATOR MATRIX"); right.pack(side="left",fill="both",expand=True)
        self._heat=tk.Canvas(right,bg=_T["surface"],bd=0,highlightthickness=0)
        self._heat.pack(fill="both",expand=True,padx=10,pady=(0,10))
        bot=panel(self,"  RISK REASONS"); bot.pack(fill="x",padx=16,pady=(0,16))
        self._reasons=StyledText(bot,height=6)
        self._reasons.add_tag("plus",foreground=_T["danger"]); self._reasons.add_tag("ok",foreground=_T["success"])
        self._reasons.pack(fill="x",padx=6,pady=(0,8))
    @staticmethod
    def _lerp(c1,c2,t):
        def h2r(h): return tuple(int(h.lstrip("#")[i:i+2],16) for i in (0,2,4))
        r1,g1,b1=h2r(c1); r2,g2,b2=h2r(c2); t=min(1,max(0,t))
        return f"#{int(r1+(r2-r1)*t):02x}{int(g1+(g2-g1)*t):02x}{int(b1+(b2-b1)*t):02x}"
    def _draw_bars(self, reasons):
        c=self._bar; c.delete("all")
        W=max(c.winfo_width(),300); H=max(c.winfo_height(),300)
        if not reasons: c.create_text(W//2,H//2,text="No score contributors",fill=_T["muted"],font=(_T["font_main"],10)); return
        parsed=[(r.split("=")[0][:24], int(m.group(1)) if (m:=re.search(r"\+(\d+)",r)) else 10) for r in reasons]
        mv=max(v for _,v in parsed) or 1; bar_h=min(28,(H-40)//max(len(parsed),1)); gap=4; lw=170
        for i,(lbl,val) in enumerate(parsed):
            y=20+i*(bar_h+gap); t=val/mv; col=self._lerp(_T["success"],_T["danger"],min(t,1))
            bw=int((W-lw-60)*t)
            c.create_text(lw-6,y+bar_h//2,text=lbl,anchor="e",fill=_T["text"],font=(_T["font_main"],8))
            if bw>0: c.create_rectangle(lw,y,lw+bw,y+bar_h,fill=col,outline="")
            c.create_text(lw+bw+6,y+bar_h//2,text=f"+{val}",anchor="w",fill=col,font=(_T["font_main"],8))
    def _draw_heatmap(self, md):
        c=self._heat; c.delete("all")
        W=max(c.winfo_width(),260); H=max(c.winfo_height(),260)
        ti=md.get("threat_intel",{}); iocs=ti.get("iocs",[]); reasons=ti.get("risk_reasons",[])
        yara_m=ti.get("yara_matches",[]); finfo=md.get("forensic_info",{})
        try: ev=float(finfo.get("entropy",0) or 0)
        except: ev=0.0
        pdf_m=md.get("extracted_metadata",{}).get("pdf",{})
        inds=[("High Entropy",ev>7.5),("Known Hash",any("Known-malicious" in i for i in iocs)),
              ("YARA Match",bool(yara_m)),("Packer Found",any("Packer" in i for i in iocs)),
              ("Trailing Data",any("trailing" in i.lower() for i in iocs)),
              ("GPS Leak",any("GPS" in i for i in iocs)),("PE Anomaly",any("PE" in r for r in reasons)),
              ("Suspicious",bool(finfo.get("suspicious_flags"))),("IOC Present",bool(iocs)),
              ("Script",any("shell" in i.lower() for i in iocs)),("Encrypted?",bool(pdf_m.get("is_encrypted"))),
              ("Base64 PE",any("Base64" in i for i in iocs)),
              ("Webshell?",any("webshell" in m.get("rule","").lower() for m in yara_m)),
              ("RevShell?",any("Reverse" in m.get("rule","") for m in yara_m)),
              ("CryptoMiner?",any("Crypto" in m.get("rule","") for m in yara_m))]
        COLS=5; ROWS=math.ceil(len(inds)/COLS); cw=(W-20)/COLS; ch=(H-20)/ROWS
        for idx,(lbl,active) in enumerate(inds):
            col=idx%COLS; ri=idx//COLS; x1=10+col*cw; y1=10+ri*ch; x2=x1+cw-4; y2=y1+ch-4
            c.create_rectangle(x1,y1,x2,y2,fill=_T["danger"] if active else _T["grid"],
                               outline=_T["warn"] if active else _T["border"],width=1)
            c.create_text((x1+x2)//2,(y1+y2)//2,text=lbl,
                          fill=_T["text"] if active else _T["muted"],font=(_T["font_main"],7),width=int(cw-8))
    def populate(self, md:dict):
        ti=md.get("threat_intel",{})
        self._draw_bars(ti.get("risk_reasons",[])); self._draw_heatmap(md)
        self._reasons.clear(); reasons=ti.get("risk_reasons",[])
        if reasons:
            for r in reasons: self._reasons.append(f"  ▸  {r}\n","plus")
        else: self._reasons.append("  ✓  Risk score: 0 — no contributors.\n","ok")


# Chain of Custody Tab 
class CocTab(tk.Frame):
    def __init__(self,parent,app):
        super().__init__(parent,bg=_T["bg"]); self.app=app; self._build(); ThemeManager.register(self)
    def _apply_theme(self): self.configure(bg=_T["bg"])
    def _build(self):
        tk.Label(self,text="CHAIN OF CUSTODY",font=(_T["font_ui"],13,"bold"),
                 bg=_T["bg"],fg=_T["accent"]).pack(padx=16,pady=(14,8),anchor="w")
        top=tk.Frame(self,bg=_T["bg"]); top.pack(fill="x",padx=16,pady=(0,10))
        self._int_lbl=tk.Label(top,text="Integrity: –",font=(_T["font_ui"],11,"bold"),bg=_T["bg"],fg=_T["muted"])
        self._int_lbl.pack(side="left")
        self._rec_lbl=tk.Label(top,text="Record: –",font=(_T["font_main"],8),bg=_T["bg"],fg=_T["muted"])
        self._rec_lbl.pack(side="right")
        cp=panel(self,"  AUDIT TRAIL"); cp.pack(fill="both",expand=True,padx=16,pady=(0,10))
        self._c=tk.Canvas(cp,bg=_T["surface"],bd=0,highlightthickness=0)
        vsb=tk.Scrollbar(cp,orient="vertical",command=self._c.yview,bg=_T["border"],
                         troughcolor=_T["bg"],activebackground=_T["accent"],relief="flat",width=8)
        self._c.configure(yscrollcommand=vsb.set)
        self._c.pack(side="left",fill="both",expand=True,padx=6,pady=(0,8))
        vsb.pack(side="right",fill="y",pady=(0,8))
        hp=panel(self,"  HASH LOG"); hp.pack(fill="x",padx=16,pady=(0,16))
        self._ht=StyledText(hp,height=5)
        self._ht.add_tag("ok",foreground=_T["success"]); self._ht.add_tag("dim",foreground=_T["muted"])
        self._ht.pack(fill="x",padx=6,pady=(0,8))
    def _draw(self,events):
        c=self._c; c.delete("all")
        W=max(c.winfo_width(),700); ROW=54; total_h=max(400,len(events)*ROW+60)
        c.configure(scrollregion=(0,0,W,total_h))
        AC={"CUSTODY_OPENED":_T["success"],"EVIDENCE_ACQUIRED":_T["accent2"],
            "ANALYSIS_PERFORMED":_T["accent"],"HASH_VERIFICATION":_T["warn"],
            "EVIDENCE_EXPORTED":"#a78bfa","CUSTODY_CLOSED":_T["success"]}
        for i,ev in enumerate(events):
            y=30+i*ROW; col=AC.get(ev["action"],_T["muted"])
            if i<len(events)-1: c.create_line(42,y+16,42,y+ROW,fill=_T["border2"],width=2,dash=(3,3))
            c.create_rectangle(22,y,62,y+28,fill=_T["surface2"],outline=col,width=2)
            c.create_text(42,y+14,text=str(ev["seq"]),fill=col,font=(_T["font_main"],10,"bold"))
            c.create_text(76,y+8,text=ev["action"],anchor="w",fill=col,font=(_T["font_main"],9,"bold"))
            ds=", ".join(f"{k}={v}" for k,v in ev["details"].items()
                         if isinstance(v,(str,int,bool)) and k not in ("note","hashes"))[:90]
            c.create_text(76,y+22,text=f"{ev['timestamp'][:19]}   {ds}",anchor="w",fill=_T["text_dim"],font=(_T["font_main"],8))
            c.create_text(W-10,y+8,text=f"#{ev.get('entry_hash','')[:24]}…",anchor="e",fill=_T["border2"],font=(_T["font_main"],7))
    def populate(self,md:dict):
        coc=md.get("chain_of_custody",{})
        if not coc: self._int_lbl.configure(text="Chain of Custody: not recorded.",fg=_T["muted"]); return
        valid=coc.get("integrity_valid",False)
        self._int_lbl.configure(text=f"Integrity: {'✓  CHAIN INTACT' if valid else '✗  COMPROMISED'}",
                                fg=_T["success"] if valid else _T["danger"])
        self._rec_lbl.configure(text=f"Record ID: {coc.get('record_id','')[:36]}",fg=_T["text_dim"])
        self._draw(coc.get("events",[]))
        self._ht.clear()
        self._ht.append("  Acquisition Hashes\n\n","dim")
        for algo,digest in md.get("hashes",{}).items():
            self._ht.append(f"  {algo.upper():<8}  {digest}\n","ok")


# Raw JSON Tab
class RawTab(tk.Frame):
    def __init__(self,parent,app):
        super().__init__(parent,bg=_T["bg"]); self.app=app; self._build(); ThemeManager.register(self)
    def _apply_theme(self): self.configure(bg=_T["bg"])
    def _build(self):
        top=tk.Frame(self,bg=_T["bg"]); top.pack(fill="x",padx=16,pady=(14,8))
        tk.Label(top,text="RAW JSON OUTPUT",font=(_T["font_ui"],13,"bold"),bg=_T["bg"],fg=_T["accent"]).pack(side="left")
        p=panel(self,None); p.pack(fill="both",expand=True,padx=16,pady=(0,16))
        self._t=StyledText(p); self._t.pack(fill="both",expand=True,padx=6,pady=(0,8))
    def populate(self,md:dict): self._t.set_text(json.dumps(md,indent=2,default=str))


# Theme Picker 
class ThemePicker(tk.Toplevel):
    def __init__(self, parent, current, on_select):
        super().__init__(parent); self.title("Choose Theme")
        self.configure(bg=_T["bg"]); self.resizable(False,False)
        self.on_select=on_select
        tk.Label(self,text="SELECT THEME",font=(_T["font_main"],12,"bold"),
                 bg=_T["bg"],fg=_T["accent"]).pack(pady=(20,12),padx=30)
        for name,theme in THEMES.items():
            row=tk.Frame(self,bg=theme["bg"],highlightthickness=2,
                         highlightbackground=theme["accent"] if name==current else theme["border"],cursor="hand2")
            row.pack(fill="x",padx=20,pady=4)
            sw=tk.Canvas(row,width=200,height=40,bg=theme["bg"],bd=0,highlightthickness=0)
            sw.pack(side="left",padx=8,pady=8)
            for i,col in enumerate([theme["accent"],theme["accent2"],theme["success"],
                                    theme["warn"],theme["danger"],theme["muted"]]):
                sw.create_rectangle(i*33,5,i*33+30,35,fill=col,outline="")
            info=tk.Frame(row,bg=theme["bg"]); info.pack(side="left",padx=8)
            tk.Label(info,text=name,font=(_T["font_ui"],11,"bold"),bg=theme["bg"],fg=theme["text_bright"]).pack(anchor="w")
            if name==current:
                tk.Label(info,text="● ACTIVE",font=(_T["font_main"],8),bg=theme["bg"],fg=theme["accent"]).pack(anchor="w")
            for w in [row,sw,info]+list(info.winfo_children()):
                try: w.bind("<Button-1>",lambda e,n=name: self._pick(n))
                except: pass
        tk.Button(self,text="  Close  ",command=self.destroy,
                  bg=_T["surface2"],fg=_T["text"],relief="flat",font=(_T["font_ui"],9),
                  padx=12,pady=6,cursor="hand2").pack(pady=16)
        self.grab_set()
    def _pick(self,name): self.on_select(name); self.destroy()


# Main GUI Window
class ForensicGUI(tk.Tk):
    def __init__(self):
        super().__init__(); self.title(f"{TOOL_NAME}  v{VERSION}")
        self.configure(bg=_T["bg"]); self.geometry("1400x940"); self.minsize(1050,700)
        self._metadata:Optional[dict]=None; self._results:List[dict]=[]
        self._file_path=""; self._current_theme="Cyber Dark"
        self._build_titlebar(); self._build_toolbar()
        self._build_notebook(); self._build_statusbar()
        self.protocol("WM_DELETE_WINDOW",self.destroy)

    def _build_titlebar(self):
        bar=tk.Frame(self,bg=_T["bg2"],height=48); bar.pack(fill="x"); bar.pack_propagate(False)
        brand=tk.Frame(bar,bg=_T["bg2"]); brand.pack(side="left",padx=16)
        tk.Label(brand,text="◈",font=(_T["font_main"],18,"bold"),bg=_T["bg2"],fg=_T["accent"]).pack(side="left")
        tk.Label(brand,text=f" {TOOL_NAME}",font=(_T["font_main"],14,"bold"),bg=_T["bg2"],fg=_T["text_bright"]).pack(side="left")
        tk.Label(brand,text=f"  v{VERSION}",font=(_T["font_main"],9),bg=_T["bg2"],fg=_T["muted"]).pack(side="left")
        tk.Frame(bar,bg=_T["border2"],width=1).pack(side="left",fill="y",padx=12,pady=8)
        tk.Label(bar,text="Digital Forensics  ·  Streaming Memory  ·  Timeline Graph  ·  Multi-File Dashboard",
                 font=(_T["font_ui"],9),bg=_T["bg2"],fg=_T["text_dim"]).pack(side="left")
        tk.Button(bar,text="◐  Theme",command=self._theme_picker,
                  bg=_T["surface2"],fg=_T["accent"],relief="flat",font=(_T["font_ui"],9),
                  cursor="hand2",padx=14,pady=6,activebackground=_T["sel"],activeforeground=_T["accent"],
                  highlightthickness=1,highlightbackground=_T["border2"]).pack(side="right",padx=16,pady=8)

    def _build_toolbar(self):
        tb=tk.Frame(self,bg=_T["surface"],height=52,highlightthickness=1,highlightbackground=_T["border"])
        tb.pack(fill="x"); tb.pack_propagate(False)
        def btn(text,cmd,prim=False):
            b=tk.Button(tb,text=text,command=cmd,
                        bg=_T["accent"] if prim else _T["grid"],
                        fg=_T["bg"] if prim else _T["text"],relief="flat",bd=0,
                        font=(_T["font_ui"],9),cursor="hand2",
                        activebackground=_T["sel"],activeforeground=_T["accent"],
                        padx=14,pady=8,highlightthickness=1,
                        highlightbackground=_T["accent"] if prim else _T["border"])
            b.pack(side="left",padx=4,pady=7); return b
        btn("  Open File",     self.open_file, prim=True)
        btn("  Scan Directory",self.open_directory)
        btn("  Export JSON",   lambda: self._export("json"))
        btn("  Export HTML",   lambda: self._export("html"))
        btn("  Export CSV",    lambda: self._export("csv"))
        tk.Frame(tb,bg=_T["border2"],width=1).pack(side="left",fill="y",padx=8,pady=8)
        tk.Label(tb,text="Case ID:",font=(_T["font_ui"],8),bg=_T["surface"],fg=_T["muted"]).pack(side="left",padx=(0,4))
        self._case_var=tk.StringVar(value=f"CASE-{datetime.now().strftime('%Y%m%d-%H%M%S')}")
        tk.Entry(tb,textvariable=self._case_var,width=22,font=(_T["font_main"],9),
                 bg=_T["grid"],fg=_T["accent"],insertbackground=_T["accent"],relief="flat",bd=4).pack(side="left",padx=(0,12))
        tk.Label(tb,text="Examiner:",font=(_T["font_ui"],8),bg=_T["surface"],fg=_T["muted"]).pack(side="left",padx=(0,4))
        self._exam_var=tk.StringVar(value="Unknown")
        tk.Entry(tb,textvariable=self._exam_var,width=14,font=(_T["font_main"],9),
                 bg=_T["grid"],fg=_T["text"],insertbackground=_T["accent"],relief="flat",bd=4).pack(side="left",padx=(0,12))
        self._path_lbl=tk.Label(tb,text="No file loaded",font=(_T["font_ui"],9),bg=_T["surface"],fg=_T["muted"])
        self._path_lbl.pack(side="left",padx=8)

    def _build_notebook(self):
        s=ttk.Style(); s.theme_use("default")
        s.configure("F.TNotebook",background=_T["bg"],borderwidth=0)
        s.configure("F.TNotebook.Tab",background=_T["surface"],foreground=_T["muted"],
                    font=(_T["font_ui"],9,"bold"),padding=(18,8),borderwidth=0)
        s.map("F.TNotebook.Tab",background=[("selected",_T["tab_active"])],
              foreground=[("selected",_T["accent"])])
        nb=ttk.Notebook(self,style="F.TNotebook"); nb.pack(fill="both",expand=True)
        self._dash=DashboardTab(nb,self)
        self._tl  =TimelineGraphTab(nb,self)
        self._mfd =MultiFileDashTab(nb,self)
        self._heat=HeatmapTab(nb,self)
        self._yara=YaraTab(nb,self)
        self._coc =CocTab(nb,self)
        self._raw =RawTab(nb,self)
        nb.add(self._dash, text="  Dashboard  ")
        nb.add(self._tl,   text="  Timeline Graph  ")
        nb.add(self._mfd,  text="  Multi-File Risk  ")
        nb.add(self._heat, text="  Risk Heatmap  ")
        nb.add(self._yara, text="  YARA  ")
        nb.add(self._coc,  text="  Chain of Custody  ")
        nb.add(self._raw,  text="  Raw JSON  ")
        self._nb=nb

    def _build_statusbar(self):
        sb=tk.Frame(self,bg=_T["surface2"],height=28,highlightthickness=1,highlightbackground=_T["border"])
        sb.pack(fill="x",side="bottom"); sb.pack_propagate(False)
        self._stat=tk.Label(sb,text="Ready.",font=(_T["font_main"],8),bg=_T["surface2"],fg=_T["muted"])
        self._stat.pack(side="left",padx=12)
        self._prog=ttk.Progressbar(sb,mode="indeterminate",length=120)
        self._prog.pack(side="right",padx=12,pady=4)
        tk.Label(sb,text=f"{TOOL_NAME} v{VERSION}  |  Python {platform.python_version()}  |  v6: Streaming Memory",
                 font=(_T["font_main"],7),bg=_T["surface2"],fg=_T["muted"]).pack(side="right",padx=16)

    def _set_status(self, msg, busy=False):
        self._stat.configure(text=msg)
        if busy: self._prog.start(10)
        else:    self._prog.stop()

    def _theme_picker(self):
        ThemePicker(self, self._current_theme, self._change_theme)

    def _change_theme(self, name):
        self._current_theme=name; ThemeManager.apply(name)
        s=ttk.Style(); s.theme_use("default")
        s.configure("F.TNotebook",background=_T["bg"])
        s.configure("F.TNotebook.Tab",background=_T["surface"],foreground=_T["muted"],
                    font=(_T["font_ui"],9,"bold"),padding=(18,8))
        s.map("F.TNotebook.Tab",background=[("selected",_T["tab_active"])],
              foreground=[("selected",_T["accent"])])
        self.configure(bg=_T["bg"])
        self._set_status(f"Theme: {name}")

    def open_file(self):
        path=filedialog.askopenfilename(title="Select file",
            filetypes=[("All","*.*"),("Images","*.jpg *.jpeg *.png *.gif *.bmp *.tiff *.webp"),
                       ("Documents","*.pdf *.docx *.doc"),("Archives","*.zip *.jar *.apk"),
                       ("Executables","*.exe *.dll"),("Audio/Video","*.mp3 *.mp4 *.avi *.mkv")])
        if path:
            self._file_path=path; self._path_lbl.configure(text=Path(path).name[-60:],fg=_T["text"])
            self._run(path,single=True)

    def open_directory(self):
        path=filedialog.askdirectory(title="Select directory to scan")
        if path:
            self._file_path=path; self._path_lbl.configure(text=str(path)[-60:],fg=_T["text"])
            self._run(path,single=False)

    def _run(self, path, single):
        self._set_status(f"Analysing {Path(path).name}…", busy=True)
        def worker():
            try:
                cid  = self._case_var.get(); examiner = self._exam_var.get() or "Unknown"
                logger=ForensicLogger(case_id=cid); coc=ChainOfCustody(case_id=cid,examiner=examiner)
                if single:
                    ext=EnhancedMetadataExtractor(path,case_id=cid,logger=logger,
                                                  chain_of_custody=coc,examiner=examiner)
                    md=ext.extract_all(); coc.close()
                    self.after(0,self._populate_single_file,md)
                else:
                    scanner=ParallelScanner(max_workers=DEFAULT_WORKERS,case_id=cid,
                                            logger=logger,chain_of_custody=coc,examiner=examiner)
                    results=scanner.scan_directory(Path(path),recursive=False)
                    coc.close()
                    self.after(0,self._populate_multi,results,coc)
            except Exception as ex:
                self.after(0,messagebox.showerror,"Error",str(ex))
                self.after(0,self._set_status,f"Error: {ex}")
        threading.Thread(target=worker,daemon=True).start()

    def _populate_single_file(self, md: dict):
        self._metadata=md
        self._dash.populate(md); self._tl.populate(md); self._heat.populate(md)
        self._yara.populate(md); self._coc.populate(md); self._raw.populate(md)
        fi=md.get("file_info",{}); score=md.get("threat_intel",{}).get("risk_score",0)
        level=md.get("threat_intel",{}).get("risk_level","CLEAN")
        self._set_status(f"✓  {fi.get('filename','')}  |  Risk: {score}/100 ({level})  |  {fi.get('file_size_human','')}",busy=False)

    def _populate_multi(self, results: List[dict], coc):
        self._results=results
        self._mfd.load_results(results)
        if results:
            best=min(results,key=lambda r: r.get("threat_intel",{}).get("risk_score",0))
            best["chain_of_custody"]=coc.to_dict()
            self._populate_single_file(best)
        self._set_status(f"✓ Scanned {len(results)} files  |  Multi-File Risk tab populated",busy=False)
        # switch to multi-file tab
        self._nb.select(2)

    def _export(self, fmt):
        if not self._metadata and not self._results:
            messagebox.showwarning("No Data","Please analyse a file first."); return
        ext={"json":".json","html":".html","csv":".csv"}.get(fmt,".json")
        path=filedialog.asksaveasfilename(
            defaultextension=ext, filetypes=[(fmt.upper(),f"*{ext}"),("All","*.*")],
            initialfile=f"forensic_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}{ext}")
        if path:
            data={"files":self._results} if self._results else self._metadata
            save_output(data,fmt,path); self._set_status(f"✓ Exported: {path}")


# Startup

def _launch_gui():
    if not TK_AVAILABLE:
        print("❌ tkinter not available. On Debian/Ubuntu: sudo apt install python3-tk")
        sys.exit(1)
    ForensicGUI().mainloop()

def _startup_prompt():
    print_banner()
    print(f"\n{'─'*62}")
    print("  How do you want to run the tool?")
    print(f"{'─'*62}")
    print("  [1]  CLI  – Terminal / interactive menu")
    print("  [2]  GUI  – Graphical dashboard with all v6 features")
    print(f"{'─'*62}")
    while True:
        c=input("\n  Enter 1 or 2: ").strip()
        if c=="1": return "cli"
        if c=="2": return "gui"
        print("  Please enter 1 or 2.")

if __name__=="__main__":
    if len(sys.argv)>1: main()
    else:
        mode=_startup_prompt()
        if mode=="gui": _launch_gui()
        else: interactive_mode()